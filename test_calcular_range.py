"""Teste para função _calcular_range_respeitando_secao."""
from __future__ import annotations

import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.services.servico_merge_docx import _calcular_range_respeitando_secao


def criar_docx_teste(caminho: str) -> None:
    """Cria um DOCX de teste com múltiplas seções e headings."""
    doc = Document()
    
    # Seção 1
    doc.add_heading('Capítulo 1: Introdução', level=1)
    doc.add_paragraph('Conteúdo do capítulo 1.')
    doc.add_heading('1.1 Subseção', level=2)
    doc.add_paragraph('Conteúdo da subseção 1.1.')
    
    # Adicionar quebra de seção (nova página)
    doc.add_page_break()
    
    # Seção 2  
    doc.add_heading('Capítulo 2: Metodologia', level=1)
    doc.add_paragraph('Conteúdo do capítulo 2.')
    doc.add_heading('2.1 Subseção A', level=2)
    doc.add_paragraph('Conteúdo da subseção 2.1.')
    doc.add_heading('2.2 Subseção B', level=2)
    doc.add_paragraph('Conteúdo da subseção 2.2.')
    
    # Adicionar quebra de seção (contínua)
    # Nota: python-docx não tem método direto para sectPr contínuo
    # Vamos adicionar um parágrafo marcador
    p = doc.add_paragraph('--- QUEBRA DE SEÇÃO ---')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Seção 3
    doc.add_heading('Anexo A: Dados Complementares', level=1)
    doc.add_paragraph('Conteúdo do anexo A.')
    
    doc.save(caminho)
    print(f'DOCX de teste criado: {caminho}')


def testar_funcao():
    """Testa a função _calcular_range_respeitando_secao."""
    caminho_teste = 'teste_secoes.docx'
    
    try:
        # Criar documento de teste
        criar_docx_teste(caminho_teste)
        
        # Carregar documento
        doc = Document(caminho_teste)
        
        print(f'Documento carregado. Total de elementos no body: {len(doc.element.body)}')
        
        # Encontrar índices dos headings
        body = doc.element.body
        indices_headings = []
        for i, child in enumerate(body):
            if child.tag.endswith('}p'):
                # Verificar se tem texto
                texto = ''
                for t in child.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if t.text:
                        texto += t.text
                if 'Capítulo 1' in texto:
                    print(f'Heading "Capítulo 1" encontrado no índice {i}')
                    # Testar função com este heading
                    resultado = _calcular_range_respeitando_secao(doc, i, 1)
                    print(f'Resultado para Capítulo 1 (nível 1):')
                    print(f'  inicio: {resultado["inicio"]}')
                    print(f'  fim: {resultado["fim"]}')
                    print(f'  secao_inicio: {resultado["secao_inicio"]}')
                    print(f'  secao_fim: {resultado["secao_fim"]}')
                    print(f'  encontrou_limite_secao: {resultado["encontrou_limite_secao"]}')
                
                if 'Capítulo 2' in texto:
                    print(f'\nHeading "Capítulo 2" encontrado no índice {i}')
                    resultado = _calcular_range_respeitando_secao(doc, i, 1)
                    print(f'Resultado para Capítulo 2 (nível 1):')
                    print(f'  inicio: {resultado["inicio"]}')
                    print(f'  fim: {resultado["fim"]}')
                    print(f'  secao_inicio: {resultado["secao_inicio"]}')
                    print(f'  secao_fim: {resultado["secao_fim"]}')
                    print(f'  encontrou_limite_secao: {resultado["encontrou_limite_secao"]}')
        
        # Verificar elementos sectPr
        print('\nProcurando elementos sectPr no documento:')
        for i, child in enumerate(body):
            if child.tag.endswith('}sectPr'):
                print(f'  sectPr encontrado no índice {i}')
        
    finally:
        # Limpar arquivo de teste
        if os.path.exists(caminho_teste):
            os.remove(caminho_teste)
            print(f'\nArquivo de teste removido: {caminho_teste}')


if __name__ == '__main__':
    testar_funcao()