"""Teste manual de fumaça para `servico_merge_docx.substituir_capitulo`.

Uso (relatório id=7 com capítulo "Coordenação" como alvo):
    .\\venv\\Scripts\\python test_merge_docx.py
"""
from __future__ import annotations

import os
import shutil
import sys

from docx import Document

from app import create_app
from app.models.capitulo_documento import CapituloDocumento
from app.models.relatorio_producao import RelatorioProducao
from app import db
from app.services.servico_merge_docx import (
    localizar_range_capitulo,
    sincronizar_subcapitulos,
    substituir_capitulo,
)


def gerar_docx_autor_temporario(caminho_saida: str, titulo: str) -> None:
    """Gera um DOCX de teste simulando o upload de um autor."""
    doc = Document()
    doc.add_heading(titulo, level=2)
    doc.add_paragraph(
        'Este é um parágrafo gerado pelo teste de merge. '
        'Ele substitui o conteúdo original do capítulo no DOCX em '
        'produção, preservando o cabeçalho do capítulo.'
    )
    doc.add_paragraph(
        'Segundo parágrafo: validar que múltiplos parágrafos são '
        'inseridos corretamente após o heading.'
    )
    doc.add_heading('Subseção criada pelo autor', level=3)
    doc.add_paragraph(
        'Conteúdo da subseção. Quando captioning estiver pronto, '
        'esta subseção receberá numeração contínua do capítulo pai.'
    )
    doc.save(caminho_saida)


def main():
    app = create_app()
    with app.app_context():
        id_rel = int(sys.argv[1]) if len(sys.argv) > 1 else 7
        rel = RelatorioProducao.query.get(id_rel)
        if not rel or not rel.caminho_template:
            print(f'Relatório {id_rel} não encontrado ou sem template.')
            sys.exit(1)

        master_orig = rel.caminho_template
        master_bkp = master_orig + '.bkp_test_merge'
        master_test = master_orig + '.test_merge.docx'

        # Backup do original e criar cópia para teste
        if not os.path.exists(master_bkp):
            shutil.copy(master_orig, master_bkp)
            print(f'Backup criado: {master_bkp}')
        shutil.copy(master_bkp, master_test)
        print(f'Copia para teste: {master_test}')

        # Capítulo alvo: pegar um nivel 2 conhecido para teste
        cap = (
            CapituloDocumento.query
            .filter_by(id_relatorio=id_rel)
            .filter(CapituloDocumento.nivel_capitulo == 2)
            .first()
        )
        if not cap:
            print('Sem capítulo de nível 2 para testar.')
            sys.exit(1)
        print(f'Capítulo alvo: [{cap.indice_capitulo}] '
              f'{cap.titulo_capitulo} (nivel={cap.nivel_capitulo})')

        # Localizar range
        doc_master = Document(master_test)
        rng = localizar_range_capitulo(doc_master, cap)
        if rng is None:
            print('FALHA: range do capítulo não localizado.')
            sys.exit(1)
        inicio, fim = rng
        print(f'Range no body: [{inicio}..{fim}] '
              f'({fim - inicio + 1} elementos)')

        # Gerar DOCX do autor
        autor_path = master_orig + '.autor_teste.docx'
        gerar_docx_autor_temporario(autor_path, cap.titulo_capitulo)
        print(f'DOCX autor gerado: {autor_path}')

        # Mesclar
        ok = substituir_capitulo(master_test, cap, autor_path)
        print(f'Substituição ok: {ok}')

        # Sincronizar subcapítulos no banco
        sync = sincronizar_subcapitulos(db.session, cap, master_test)
        print(f'Sync subcapítulos: {sync}')
        # Reverter qualquer alteração no banco (esse é só teste!)
        db.session.rollback()

        # Verificar resultado
        doc_apos = Document(master_test)
        rng2 = localizar_range_capitulo(doc_apos, cap)
        if rng2:
            inicio2, fim2 = rng2
            n_elementos = fim2 - inicio2 + 1
            print(f'Range após merge: [{inicio2}..{fim2}] '
                  f'({n_elementos} elementos)')
            print('Conteúdo do capítulo após merge:')
            body = doc_apos.element.body
            for i in range(inicio2, fim2 + 1):
                el = body[i]
                if el.tag.endswith('}p'):
                    txt = ''.join(t.text or '' for t in el.iter(
                        '{http://schemas.openxmlformats.org/'
                        'wordprocessingml/2006/main}t'
                    ))
                    print(f'  [{i}] p: {txt[:90]}')
                else:
                    print(f'  [{i}] {el.tag.split("}")[-1]}')
        else:
            print('FALHA: capítulo não localizável após merge.')

        print()
        print(f'Arquivo de teste preservado em: {master_test}')
        print(f'Original intocado: {master_orig}')
        print(f'Para restaurar:    copy "{master_bkp}" "{master_orig}"')


if __name__ == '__main__':
    main()
