"""Configuração global de fixtures para testes."""
from datetime import date
import os
import tempfile

import pytest
from docx import Document
import hypothesis.strategies as st
from hypothesis import assume

from app import create_app, db
from app.models.dominio import Dominio
from app.models.usuario import Usuario
from app.models.relatorio_producao import RelatorioProducao
from app.config import Config


class TestConfig(Config):
    """Configuração para testes."""
    TESTING = True
    SQLALCHEMY_DATABASE_URI = 'sqlite:///:memory:'
    WTF_CSRF_ENABLED = False
    # Desabilitar proteção CSRF em testes
    TESTING = True


@pytest.fixture(name='app', scope='function')
def fixture_app():
    """Cria app Flask com BD em memória para cada teste."""
    flask_app = create_app(TestConfig)
    flask_app.config['TESTING'] = True
    flask_app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///:memory:'
    flask_app.config['WTF_CSRF_ENABLED'] = False

    with flask_app.app_context():
        db.create_all()

        # Seed mínimo de domínios necessários
        dominios_esperados = [
            ('perfil_usuario', 'autor', 'Autor'),
            ('perfil_usuario', 'coordenador', 'Coordenador'),
            ('status_relatorio', 'em_producao', 'Em Produção'),
            ('status_capitulo', 'em_edicao', 'Em Edição'),
            ('status_capitulo', 'enviado_revisao', 'Enviado para Revisão'),
            ('status_capitulo', 'aprovado', 'Aprovado'),
            ('status_capitulo', 'rejeitado', 'Rejeitado'),
        ]

        for tipo, valor, descricao in dominios_esperados:
            db.session.add(Dominio(
                tipo=tipo, valor=valor, descricao=descricao
            ))

        db.session.commit()

        yield flask_app

        db.session.remove()
        db.drop_all()


@pytest.fixture(name='app_context')
def fixture_app_context(app):
    """Fornece app context já criado."""
    with app.app_context():
        yield app


# =====================================================================
# Estratégias Hypothesis para Property-Based Testing
# =====================================================================

@st.composite
def estilo_heading_strategy(draw) -> str:
    """Gera estilos de heading válidos em DOCX."""
    estilos = [
        'Heading 1', 'Heading 2', 'Heading 3',
        'Título 1', 'Título 2', 'Título 3',
        'Title', 'Subtitle',
        'Heading1', 'Heading2',
    ]
    return draw(st.sampled_from(estilos))


@st.composite
def estilo_anexo_apendice_strategy(draw) -> tuple:
    """Gera estilos específicos para anexos e apêndices.

    Retorna (estilo_docx, classificacao_esperada, prefixo_esperado)
    """
    tipo = draw(st.sampled_from(['anexo', 'apendice']))

    if tipo == 'anexo':
        estilos = [
            ('Anexo', 'anexo', 'ANEXO_'),
            ('ANEXO', 'anexo', 'ANEXO_'),
            ('Anexo A', 'anexo', 'ANEXO_'),
        ]
    else:  # apendice
        estilos = [
            ('Apêndice', 'apendice', 'APENDICE_'),
            ('APÊNDICE', 'apendice', 'APENDICE_'),
            ('Apêndice I', 'apendice', 'APENDICE_'),
        ]

    return draw(st.sampled_from(estilos))


@st.composite
def titulo_capitulo_com_variacao_strategy(draw) -> str:
    """Gera títulos de capítulos com variações reais.

    Inclui espaços extras, acentos, maiúsculas/minúsculas
    """
    palavras = [
        'Introdução', 'Metodologia', 'Resultados', 'Discussão',
        'Conclusão', 'Recomendações', 'Referências', 'Anexos',
        'Apêndices', 'Prefácio', 'Dedicatória', 'Agradecimentos',
        'Resumo', 'Abstract', 'Lista de Figuras', 'Lista de Tabelas'
    ]

    titulo = draw(st.sampled_from(palavras))

    # Aplicar variações ocasionais
    variacao = draw(st.sampled_from(['original', 'maiuscula', 'minuscula', 'espacos']))

    if variacao == 'maiuscula':
        titulo = titulo.upper()
    elif variacao == 'minuscula':
        titulo = titulo.lower()
    elif variacao == 'espacos':
        titulo = '  ' + titulo + '  '

    return titulo


@st.composite
def estrutura_capitulos_strategy(draw) -> list:
    """Gera estrutura de capítulos com mistura de tipos e classificações.

    Retorna lista de dicts: [{'titulo', 'estilo', 'nivel', 'classificacao_esperada'}, ...]
    """
    capitulos = []

    # 40% de chances de ter capítulos de cada tipo
    tem_textual = draw(st.integers(min_value=1, max_value=10)) <= 7
    tem_anexos = draw(st.integers(min_value=1, max_value=10)) <= 5
    tem_apendices = draw(st.integers(min_value=1, max_value=10)) <= 4

    # Capítulos textuais
    if tem_textual:
        num_textual = draw(st.integers(min_value=1, max_value=3))
        for i in range(num_textual):
            titulo = f"Capítulo {i+1}: " + draw(titulo_capitulo_com_variacao_strategy())
            capitulos.append({
                'titulo': titulo,
                'estilo': draw(estilo_heading_strategy()),
                'nivel': draw(st.integers(min_value=1, max_value=2)),
                'classificacao_esperada': None,  # Textual não tem classificação específica
            })

    # Anexos
    if tem_anexos:
        num_anexos = draw(st.integers(min_value=1, max_value=2))
        for i in range(num_anexos):
            estilo, classif, prefixo = draw(st.sampled_from([
                ('Anexo', 'anexo', 'ANEXO_'),
                ('ANEXO', 'anexo', 'ANEXO_'),
            ]))
            capitulos.append({
                'titulo': f"Anexo {chr(65+i)}: Dados Adicionais",
                'estilo': estilo,
                'nivel': 1,
                'classificacao_esperada': classif,
                'prefixo_esperado': prefixo,
            })

    # Apêndices
    if tem_apendices:
        num_apendices = draw(st.integers(min_value=1, max_value=2))
        for i in range(num_apendices):
            estilo, classif, prefixo = draw(st.sampled_from([
                ('Apêndice', 'apendice', 'APENDICE_'),
                ('APÊNDICE', 'apendice', 'APENDICE_'),
            ]))
            capitulos.append({
                'titulo': f"Apêndice {chr(73+i)}: Formulário",
                'estilo': estilo,
                'nivel': 1,
                'classificacao_esperada': classif,
                'prefixo_esperado': prefixo,
            })

    assume(len(capitulos) > 0)  # Garantir que há pelo menos um capítulo
    return capitulos


@st.composite
def relatorio_com_template_strategy(draw, app_context) -> RelatorioProducao:
    """Gera RelatorioProducao com template DOCX válido contendo capítulos.

    Retorna instância de RelatorioProducao pronta para usar.
    """
    capitulos_estrutura = draw(estrutura_capitulos_strategy())

    # Criar DOCX
    doc = Document()
    for cap_info in capitulos_estrutura:
        heading_level = cap_info.get('nivel', 1)
        titulo = cap_info['titulo']
        doc.add_heading(titulo, level=heading_level)
        doc.add_paragraph(f"Conteúdo de {titulo}")

    # Salvar DOCX temporário
    tmpdir = tempfile.mkdtemp()
    docx_path = os.path.join(tmpdir, 'template_teste.docx')
    doc.save(docx_path)

    # Criar usuário
    perfil = Dominio.query.filter_by(tipo='perfil_usuario', valor='coordenador').first()
    usuario = Usuario(
        nome='Teste User',
        email=f'teste_{draw(st.uuids())}@test.com',
        nome_de_usuario=f'user_{draw(st.just(1)).bit_length()}',
        senha_hash='x',
        perfil_id=perfil.id if perfil else None,
        ativo=True,
    )
    db.session.add(usuario)
    db.session.flush()

    # Criar relatório
    status = Dominio.query.filter_by(tipo='status_relatorio').first()
    rel = RelatorioProducao(
        codigo_d20=f'D-20-TEST-{draw(st.uuids())}',
        numero_medicao=1,
        mes_referencia=date.today(),
        periodo_inicio=date.today(),
        periodo_fim=date.today(),
        titulo_curto='Template Teste PBT',
        status_id=status.id if status else None,
        criado_por=usuario.id,
        versao_atual='R00',
        caminho_template=docx_path,
    )
    db.session.add(rel)
    db.session.flush()

    # Retornar com estrutura esperada para validação
    setattr(rel, '_expected_capitulos', capitulos_estrutura)

    return rel


@st.composite
def capitulo_com_classificacao_strategy(draw) -> dict:
    """Gera dicionário representando um capítulo com classificação esperada.

    Retorna: {'titulo', 'estilo', 'classificacao_esperada', 'prefixo_esperado'}
    """
    tipo = draw(st.sampled_from(['textual', 'anexo', 'apendice']))

    if tipo == 'textual':
        return {
            'titulo': draw(titulo_capitulo_com_variacao_strategy()),
            'estilo': draw(estilo_heading_strategy()),
            'classificacao_esperada': None,
            'prefixo_esperado': None,
        }
    estilo, classif, prefixo = draw(estilo_anexo_apendice_strategy())
    return {
        'titulo': f"{estilo} {chr(65)}: Dados",
        'estilo': estilo,
        'classificacao_esperada': classif,
        'prefixo_esperado': prefixo,
    }


@st.composite
def lista_capitulos_variados_strategy(draw, min_size=1, max_size=5) -> list:
    """Gera lista de capítulos com classificações variadas."""
    num = draw(st.integers(min_value=min_size, max_value=max_size))
    return [draw(capitulo_com_classificacao_strategy()) for _ in range(num)]
