"""Testes property-based para determinismo de localização.

Feature: automacao-montagem-relatorios, Property 2: Determinismo e Idempotência de Localização

Valida que para qualquer capítulo com título consistente, múltiplas execuções 
de `localizar_range_capitulo_robusto()` retornam exatamente o mesmo resultado.

**Validates: Requirements 2.1, 2.3, 2.5, Determinismo**
"""
from __future__ import annotations

import hypothesis.strategies as st
from hypothesis import given, settings, assume
import re
import unicodedata
from typing import Any, Dict, List, Optional
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.servico_merge_docx import (
    localizar_range_capitulo_robusto,
    _normalizar,
    _eh_paragrafo_heading,
    _texto_paragrafo
)


# ----------------------------------------------------------------------
# Estratégias de geração de dados para testes property-based
# ----------------------------------------------------------------------

def _normalizar_texto(texto: str) -> str:
    """Normaliza texto para comparação (cópia da função do servico_merge_docx)."""
    if not texto:
        return ''
    s = unicodedata.normalize('NFD', texto)
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    s = re.sub(r'\s+', ' ', s).strip().lower()
    s = re.sub(r'^\s*(?:\d+(?:\.\d+)*|[ivx]+|[a-z])[\.\)]?\s+', '', s)
    return s


@st.composite
def titulo_capitulo_strategy(draw) -> str:
    """Gera títulos de capítulos variados para testes.
    
    Inclui variações com:
    - Números (ex: "1. INTRODUÇÃO", "2.1 METODOLOGIA")
    - Letras (ex: "A. ANEXO", "B APÊNDICE")
    - Números romanos (ex: "I. SUMÁRIO")
    - Sem numeração (ex: "INTRODUÇÃO")
    - Com acentos e caracteres especiais
    - Com espaços variados
    - Com typos leves (para testar fuzzy matching)
    """
    # Partes do título
    prefixos = draw(st.one_of(
        st.just(''),
        st.integers(min_value=1, max_value=10).map(lambda x: f"{x}."),
        st.integers(min_value=1, max_value=10).map(lambda x: f"{x}"),
        st.sampled_from(['I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.', 'IX.', 'X.']),
        st.sampled_from(['A.', 'B.', 'C.', 'D.', 'E.', 'F.', 'G.', 'H.']),
        st.sampled_from(['ANEXO', 'APÊNDICE', 'ANEXO A', 'APÊNDICE B']),
    ))
    
    palavras = draw(st.lists(
        st.sampled_from([
            'INTRODUÇÃO', 'METODOLOGIA', 'RESULTADOS', 'DISCUSSÃO', 'CONCLUSÃO',
            'REFERÊNCIAS', 'SUMÁRIO', 'RESUMO', 'ABSTRACT', 'LISTA', 'FIGURAS',
            'TABELAS', 'ABREVIATURAS', 'GLOSSÁRIO', 'APÊNDICE', 'ANEXO',
            'REVISÃO', 'FUNDAMENTAÇÃO', 'ANÁLISE', 'AVALIAÇÃO', 'PROPOSTA'
        ]),
        min_size=1,
        max_size=3
    ))
    
    # Juntar palavras com espaços
    titulo_base = ' '.join(palavras)
    
    # Aplicar variações
    variacao = draw(st.sampled_from([
        lambda x: x,  # Sem alteração
        lambda x: x.lower(),  # Minúsculas
        lambda x: x.upper(),  # Maiúsculas
        lambda x: x.title(),  # Title case
        lambda x: re.sub(r'[ÁÀÂÃ]', 'A', x),  # Remover acentos (simplificado)
        lambda x: re.sub(r'[ÉÈÊ]', 'E', x),
        lambda x: re.sub(r'[ÍÌÎ]', 'I', x),
        lambda x: re.sub(r'[ÓÒÔÕ]', 'O', x),
        lambda x: re.sub(r'[ÚÙÛ]', 'U', x),
        lambda x: re.sub(r'[Ç]', 'C', x),
    ]))
    
    titulo_modificado = variacao(titulo_base)
    
    # Adicionar espaços extras aleatoriamente
    if draw(st.booleans()):
        titulo_modificado = '  ' + titulo_modificado + '  '
    
    # Combinar prefixo e título
    if prefixos:
        if prefixos.endswith('.') or prefixos in ['ANEXO', 'APÊNDICE', 'ANEXO A', 'APÊNDICE B']:
            separador = ' ' if draw(st.booleans()) else ''
            resultado = f"{prefixos}{separador}{titulo_modificado}"
        else:
            separador = draw(st.sampled_from([' ', '. ', ') ', '- ', ' - ']))
            resultado = f"{prefixos}{separador}{titulo_modificado}"
    else:
        resultado = titulo_modificado
    
    # Garantir que o resultado é XML seguro
    return texto_xml_seguro(resultado)


@st.composite
def nivel_capitulo_strategy(draw) -> int:
    """Gera níveis de capítulo (1-3 para testes)."""
    return draw(st.integers(min_value=1, max_value=3))


@st.composite  
def indice_capitulo_strategy(draw) -> Optional[str]:
    """Gera índices de capítulo (opcional)."""
    return draw(st.one_of(
        st.none(),
        st.integers(min_value=1, max_value=10).map(str),
        st.integers(min_value=1, max_value=10).map(lambda x: f"{x}.1"),
        st.integers(min_value=1, max_value=10).map(lambda x: f"{x}.2.1"),
    ))


@st.composite
def classificacao_capitulo_strategy(draw) -> Optional[str]:
    """Gera classificação de capítulo."""
    return draw(st.one_of(
        st.none(),
        st.sampled_from(['textual', 'pre_textual', 'pos_textual', 'anexo', 'apendice'])
    ))


class MockCapituloDocumento:
    """Mock de CapituloDocumento para testes property-based."""
    def __init__(self, titulo: str, indice: Optional[str] = None, 
                 nivel: int = 1, classificacao: Optional[str] = None):
        self.titulo_capitulo = titulo
        self.indice_capitulo = indice
        self.nivel_capitulo = nivel
        self.classificacao = classificacao
        self.tipo_elemento = 'textual'
        if classificacao == 'anexo' or classificacao == 'apendice':
            self.tipo_elemento = 'pos_textual'
        elif classificacao == 'pre_textual':
            self.tipo_elemento = 'pre_textual'
        
        # Para match por contexto
        self.numero_capitulo_esperado = None
        if indice:
            match = re.match(r'^(\d+)', indice)
            if match:
                self.numero_capitulo_esperado = int(match.group(1))
        
        # Campos adicionais para compatibilidade
        self.id_capitulo_documento = 1
        self.id_relatorio = 1
        self.ordem_capitulo = 1
        self.nome_capitulo = titulo
        self.status_capitulo = 'em_edicao'
        self.ativo = True


@st.composite
def capitulo_fixture(draw) -> MockCapituloDocumento:
    """Gera um objeto MockCapituloDocumento para testes."""
    titulo = draw(titulo_capitulo_strategy())
    indice = draw(indice_capitulo_strategy())
    nivel = draw(nivel_capitulo_strategy())
    classificacao = draw(classificacao_capitulo_strategy())
    
    return MockCapituloDocumento(titulo, indice, nivel, classificacao)


def texto_xml_seguro(texto: str) -> str:
    """Remove caracteres não compatíveis com XML de uma string."""
    # Remove NULL bytes e caracteres de controle (exceto tab, newline, carriage return)
    import re
    # Mantém: tab (\t), newline (\n), carriage return (\r)
    # Remove: outros caracteres de controle (0x00-0x08, 0x0B-0x0C, 0x0E-0x1F)
    texto = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', texto)
    return texto


@st.composite
def documento_com_headings_fixture(draw, min_headings: int = 1, max_headings: int = 5) -> Document:
    """Gera um documento DOCX com headings variados.
    
    Cria um documento com múltiplos headings de diferentes níveis
    e conteúdo entre eles.
    """
    doc = Document()
    
    num_headings = draw(st.integers(min_value=min_headings, max_value=max_headings))
    
    for i in range(num_headings):
        # Gerar título para este heading
        titulo = draw(titulo_capitulo_strategy())
        titulo = texto_xml_seguro(titulo)
        
        # Nível do heading (1-3)
        nivel = draw(st.integers(min_value=1, max_value=3))
        
        # Adicionar heading (se título não estiver vazio após limpeza)
        if titulo.strip():
            doc.add_heading(titulo, level=nivel)
        
        # Adicionar algum conteúdo (1-3 parágrafos)
        num_paragrafos = draw(st.integers(min_value=1, max_value=3))
        for _ in range(num_paragrafos):
            texto_paragrafo = draw(st.text(min_size=10, max_size=100, alphabet=st.characters(blacklist_categories=['Cc', 'Cs'])))
            texto_paragrafo = texto_xml_seguro(texto_paragrafo)
            if texto_paragrafo.strip():
                doc.add_paragraph(texto_paragrafo)
    
    # Adicionar possíveis anexos/apêndices
    if draw(st.booleans()):
        tipos_anexo = draw(st.lists(
            st.sampled_from(['ANEXO A - Dados', 'ANEXO B - Código', 
                           'APÊNDICE I - Metodologia', 'APÊNDICE II - Resultados']),
            min_size=0,
            max_size=2
        ))
        
        for anexo in tipos_anexo:
            doc.add_heading(anexo, level=1)
            doc.add_paragraph(f"Conteúdo do {anexo}")
    
    return doc


# ----------------------------------------------------------------------
# Testes Property-Based
# ----------------------------------------------------------------------

# Feature: automacao-montagem-relatorios, Property 2: Determinismo e Idempotência de Localização
@given(
    capitulo=capitulo_fixture(),
    doc=documento_com_headings_fixture(min_headings=2, max_headings=5),
    estrategia=st.sampled_from(['multi_niveis', 'exato', 'fuzzy', 'contexto'])
)
@settings(max_examples=100, deadline=5000)
def test_property_2_determinismo_localizacao(
    capitulo: MockCapituloDocumento,
    doc: Document,
    estrategia: str
):
    """Testa que múltiplas execuções retornam exatamente o mesmo resultado.
    
    Propriedade invariante: Para qualquer capítulo e documento,
    executar `localizar_range_capitulo_robusto()` múltiplas vezes
    com os mesmos parâmetros deve retornar resultados idênticos.
    """
    # Executar a função múltiplas vezes
    resultados = []
    num_execucoes = 3  # Número suficiente para verificar determinismo
    
    for _ in range(num_execucoes):
        resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia=estrategia)
        resultados.append(resultado)
    
    # Verificar que todos os resultados são iguais
    # Comparar campos-chave que devem ser determinísticos
    for i in range(1, len(resultados)):
        resultado_atual = resultados[i]
        resultado_anterior = resultados[i-1]
        
        # Campo 'encontrado' deve ser igual
        assert resultado_atual['encontrado'] == resultado_anterior['encontrado'], \
            f"Campo 'encontrado' divergiu entre execuções: {resultado_atual['encontrado']} vs {resultado_anterior['encontrado']}"
        
        # Campo 'estrategia_usada' deve ser igual
        assert resultado_atual['estrategia_usada'] == resultado_anterior['estrategia_usada'], \
            f"Campo 'estrategia_usada' divergiu: {resultado_atual['estrategia_usada']} vs {resultado_anterior['estrategia_usada']}"
        
        # Se encontrado, campos de posição devem ser iguais
        if resultado_atual['encontrado']:
            assert resultado_atual['inicio'] == resultado_anterior['inicio'], \
                f"Campo 'inicio' divergiu: {resultado_atual['inicio']} vs {resultado_anterior['inicio']}"
            
            assert resultado_atual['fim'] == resultado_anterior['fim'], \
                f"Campo 'fim' divergiu: {resultado_atual['fim']} vs {resultado_anterior['fim']}"
            
            # Confiança deve ser igual (com tolerância de ponto flutuante)
            assert abs(resultado_atual['confianca'] - resultado_anterior['confianca']) < 0.0001, \
                f"Campo 'confianca' divergiu: {resultado_atual['confianca']} vs {resultado_anterior['confianca']}"
            
            # Título encontrado deve ser igual
            assert resultado_atual['titulo_encontrado'] == resultado_anterior['titulo_encontrado'], \
                f"Campo 'titulo_encontrado' divergiu: {resultado_atual['titulo_encontrado']} vs {resultado_anterior['titulo_encontrado']}"
        
        # Número de alternativas deve ser igual
        assert len(resultado_atual['alternativas']) == len(resultado_anterior['alternativas']), \
            f"Número de alternativas divergiu: {len(resultado_atual['alternativas'])} vs {len(resultado_anterior['alternativas'])}"


@given(
    capitulo=capitulo_fixture(),
    doc=documento_com_headings_fixture(min_headings=1, max_headings=5)
)
@settings(max_examples=50, deadline=5000)
def test_property_2_ordem_estrategias_deterministica(
    capitulo: MockCapituloDocumento,
    doc: Document
):
    """Testa que a ordem de tentativa de estratégias é sempre a mesma.
    
    Propriedade invariante: A cascata de estratégias (exato → fuzzy → contexto)
    é sempre tentada na mesma ordem, sem aleatoriedade.
    
    **Validates: Property 10: Determinismo de Match Multi-Nível**
    """
    # Executar com estratégia 'multi_niveis' (padrão)
    resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
    
    # Verificar que a estratégia usada é uma das esperadas na ordem correta
    estrategias_validas = ['exato', 'fuzzy', 'contexto', 'nenhuma']
    assert resultado['estrategia_usada'] in estrategias_validas, \
        f"Estratégia usada inválida: {resultado['estrategia_usada']}"
    
    # Se encontrado, verificar que foi usada a primeira estratégia que funcionou
    # na ordem: exato → fuzzy → contexto
    if resultado['encontrado']:
        # Podemos inferir a ordem tentada baseado no resultado
        # Se estratégia usada é 'contexto', significa que 'exato' e 'fuzzy' falharam
        # Se estratégia usada é 'fuzzy', significa que 'exato' falhou
        # Se estratégia usada é 'exato', significa que foi a primeira a funcionar
        
        # Para validar isso mais rigorosamente, executar cada estratégia individualmente
        # e verificar a consistência
        resultado_exato = localizar_range_capitulo_robusto(doc, capitulo, estrategia='exato')
        resultado_fuzzy = localizar_range_capitulo_robusto(doc, capitulo, estrategia='fuzzy')
        resultado_contexto = localizar_range_capitulo_robusto(doc, capitulo, estrategia='contexto')
        
        # Determinar qual estratégia deveria ter sido usada baseado na ordem
        estrategia_esperada = 'nenhuma'
        if resultado_exato['encontrado']:
            estrategia_esperada = 'exato'
        elif resultado_fuzzy['encontrado']:
            estrategia_esperada = 'fuzzy'
        elif resultado_contexto['encontrado']:
            estrategia_esperada = 'contexto'
        
        # Verificar consistência
        if resultado['encontrado']:
            assert resultado['estrategia_usada'] == estrategia_esperada, \
                f"Estratégia usada ({resultado['estrategia_usada']}) não corresponde à ordem esperada ({estrategia_esperada})"
        else:
            assert estrategia_esperada == 'nenhuma', \
                f"Capítulo não encontrado, mas alguma estratégia individual funcionou: {estrategia_esperada}"


@given(
    capitulo=capitulo_fixture(),
    doc=documento_com_headings_fixture(min_headings=2, max_headings=5)
)
@settings(max_examples=30, deadline=5000)
def test_property_2_estrutura_resultado_completa(
    capitulo: MockCapituloDocumento,
    doc: Document
):
    """Testa que o resultado sempre tem estrutura completa.
    
    Propriedade invariante: O dict retornado por 
    `localizar_range_capitulo_robusto()` sempre contém todos os campos
    esperados, independente do sucesso ou falha.
    
    **Validates: Property 3: Coerência de Estrutura Retornada**
    """
    resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
    
    # Campos obrigatórios que devem sempre estar presentes
    campos_obrigatorios = [
        'encontrado',           # bool
        'inicio',               # int ou None
        'fim',                  # int ou None
        'secao_inicio',         # int ou None
        'secao_fim',            # int ou None
        'titulo_encontrado',    # str ou None
        'confianca',            # float
        'estrategia_usada',     # str
        'diagnostico',          # str
        'alternativas',         # list
    ]
    
    for campo in campos_obrigatorios:
        assert campo in resultado, f"Campo '{campo}' faltando no resultado"
    
    # Validações de tipo
    assert isinstance(resultado['encontrado'], bool), "'encontrado' deve ser bool"
    assert isinstance(resultado['confianca'], float), "'confianca' deve ser float"
    assert 0.0 <= resultado['confianca'] <= 1.0, f"'confianca' fora do intervalo [0,1]: {resultado['confianca']}"
    assert isinstance(resultado['estrategia_usada'], str), "'estrategia_usada' deve ser str"
    assert isinstance(resultado['diagnostico'], str), "'diagnostico' deve ser str"
    assert isinstance(resultado['alternativas'], list), "'alternativas' deve ser list"
    
    # Validações condicionais
    if resultado['encontrado']:
        assert resultado['inicio'] is not None, "'inicio' não pode ser None quando encontrado=True"
        assert resultado['fim'] is not None, "'fim' não pode ser None quando encontrado=True"
        assert resultado['titulo_encontrado'] is not None, "'titulo_encontrado' não pode ser None quando encontrado=True"
        assert resultado['confianca'] > 0.0, "'confianca' deve ser > 0 quando encontrado=True"
    else:
        # Quando não encontrado, início e fim podem ser None
        # Mas alternativas deve estar presente (pode ser lista vazia)
        assert resultado['alternativas'] is not None, "'alternativas' não pode ser None"


@given(
    capitulo=capitulo_fixture(),
    doc=documento_com_headings_fixture(min_headings=1, max_headings=3)
)
@settings(max_examples=20, deadline=5000)
def test_property_2_confianca_consistente(
    capitulo: MockCapituloDocumento,
    doc: Document
):
    """Testa que a confiança é consistente com a estratégia usada.
    
    Propriedade invariante: A confiança retornada deve ser consistente
    com a estratégia usada:
    - Exato: confiança ~0.95 (com nível correspondente) ou ~0.9 (sem nível)
    - Fuzzy: confiança entre 0.5 e 0.9
    - Contexto: confiança entre 0.6 e 0.8
    """
    resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
    
    if not resultado['encontrado']:
        return  # Não aplicável quando não encontrado
    
    estrategia = resultado['estrategia_usada']
    confianca = resultado['confianca']
    
    if estrategia == 'exato':
        # Exato deve ter confiança alta
        assert confianca >= 0.9, f"Confiança para exato deve ser ≥ 0.9, mas é {confianca}"
        assert confianca <= 1.0, f"Confiança para exato deve ser ≤ 1.0, mas é {confianca}"
    
    elif estrategia == 'fuzzy':
        # Fuzzy deve ter confiança moderada
        assert confianca >= 0.5, f"Confiança para fuzzy deve ser ≥ 0.5, mas é {confianca}"
        assert confianca <= 0.9, f"Confiança para fuzzy deve ser ≤ 0.9, mas é {confianca}"
    
    elif estrategia == 'contexto':
        # Contexto deve ter confiança baixa-moderada
        assert confianca >= 0.6, f"Confiança para contexto deve ser ≥ 0.6, mas é {confianca}"
        assert confianca <= 0.8, f"Confiança para contexto deve ser ≤ 0.8, mas é {confianca}"


# ----------------------------------------------------------------------
# Testes auxiliares (não property-based)
# ----------------------------------------------------------------------

def test_exemplo_determinismo_manual():
    """Teste manual de exemplo para verificar determinismo."""
    print("=== Teste Manual de Determinismo ===")
    
    # Criar documento simples
    doc = Document()
    doc.add_heading('3. RESULTADOS', level=1)
    doc.add_paragraph('Texto dos resultados...')
    doc.add_heading('4. DISCUSSÃO', level=1)
    doc.add_paragraph('Texto da discussão...')
    
    # Capítulo para buscar
    capitulo = MockCapituloDocumento('RESULTADOS', indice='3', nivel=1)
    
    # Executar múltiplas vezes
    resultados = []
    for i in range(5):
        resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
        resultados.append(resultado)
        print(f"Execução {i+1}: encontrado={resultado['encontrado']}, "
              f"estrategia={resultado['estrategia_usada']}, "
              f"confiança={resultado['confianca']:.2f}")
    
    # Verificar consistência
    primeiro = resultados[0]
    for i, resultado in enumerate(resultados[1:], 1):
        assert resultado['encontrado'] == primeiro['encontrado']
        assert resultado['estrategia_usada'] == primeiro['estrategia_usada']
        if primeiro['encontrado']:
            assert resultado['inicio'] == primeiro['inicio']
            assert resultado['fim'] == primeiro['fim']
            assert abs(resultado['confianca'] - primeiro['confianca']) < 0.01
    
    print("✅ Teste manual de determinismo passou")
    return True


if __name__ == '__main__':
    # Execução direta para debugging
    import sys
    import pytest
    
    # Executar teste manual primeiro
    try:
        test_exemplo_determinismo_manual()
        print("\n✅ Teste manual executado com sucesso")
    except Exception as e:
        print(f"\n❌ Teste manual falhou: {e}")
        import traceback
        traceback.print_exc()
    
    # Executar testes property-based com pytest
    print("\n" + "="*60)
    print("Executando testes property-based...")
    print("="*60)
    
    sys.exit(pytest.main([__file__, '-v', '--tb=short']))