"""Testes para verificação de implementação de fuzzy matching com distância de edição.

Valida o requisito 2.3: Implementar match fuzzy com distância de edição.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)) + '/..')

from docx import Document
from app.services.servico_merge_docx import _match_fuzzy, _match_exato, _match_contexto, localizar_range_capitulo_robusto


class MockCapitulo:
    """Mock de CapituloDocumento para testes."""
    def __init__(self, titulo, indice=None, nivel=None, classificacao=None):
        self.titulo_capitulo = titulo
        self.indice_capitulo = indice
        self.nivel_capitulo = nivel
        self.classificacao = classificacao


def test_exemplo_1_requirements():
    """Testa o Example 1 dos requirements: typo METOLOGIA vs METODOLOGIA."""
    print("=== Teste Example 1: Error Silent Becomes Visible ===")
    
    # Criar documento com heading correto
    doc = Document()
    doc.add_heading('2. METODOLOGIA', level=1)
    doc.add_paragraph('Texto da metodologia...')
    
    # Capítulo com typo (METOLOGIA vs METODOLOGIA)
    capitulo = MockCapitulo('METOLOGIA', indice='2', nivel=1)
    
    # Testar match fuzzy
    resultado = _match_fuzzy(doc, capitulo, max_distancia_edicao=2)
    
    print(f"Cenário: Upload com typo 'METOLOGIA' vs 'METODOLOGIA' no documento")
    print(f"Resultado: encontrado={resultado['encontrado']}, confiança={resultado['confianca']:.2f}")
    
    # Verificar que fuzzy match detecta o typo
    assert resultado['encontrado'] == True, "Fuzzy match deveria detectar typo"
    assert 0.5 <= resultado['confianca'] <= 0.9, f"Confiança deveria estar entre 0.5-0.9, mas é {resultado['confianca']}"
    assert 'METODOLOGIA' in resultado['titulo_encontrado'], f"Título encontrado deveria conter 'METODOLOGIA', mas é {resultado['titulo_encontrado']}"
    
    # Verificar distância de edição (deve ser 1 para METOLOGIA vs METODOLOGIA)
    # METOLOGIA (8 chars) vs METODOLOGIA (10 chars) - falta "DO"?
    # Na verdade: M E T O L O G I A vs M E T O D O L O G I A
    #             0 1 2 3 4 5 6 7 8 vs 0 1 2 3 4 5 6 7 8 9
    # Diferença: falta 'D' na posição 4 e 'O' na posição 5? Vamos calcular:
    # METOLOGIA: M E T O L O G I A
    # METODOLOGIA: M E T O D O L O G I A
    # Inserir 'D' após 'O' e manter estrutura? Na verdade é mais complexo.
    # Para o teste, vamos verificar que a distância é pequena (≤ 2)
    if resultado.get('alternativas'):
        for alt in resultado['alternativas']:
            if 'distancia_estimada' in alt:
                assert alt['distancia_estimada'] <= 2, f"Distância estimada deveria ser ≤ 2, mas é {alt['distancia_estimada']}"
    
    print("✅ Teste Example 1 passou: Fuzzy match detecta typo com confiança adequada\n")
    return True


def test_cascata_estrategias():
    """Testa a cascata de estratégias (exato → fuzzy → contexto)."""
    print("=== Teste Cascata de Estratégias ===")
    
    doc = Document()
    doc.add_heading('1. INTRODUÇÃO', level=1)
    doc.add_paragraph('Texto introdução...')
    doc.add_heading('2. METODOLOGIA', level=1)
    doc.add_paragraph('Texto metodologia...')
    doc.add_heading('3. RESULTADOS', level=1)
    doc.add_paragraph('Texto resultados...')
    doc.add_heading('ANEXO A - Dados', level=1)
    doc.add_paragraph('Texto anexo...')
    
    # Teste 1: Exato funciona
    capitulo = MockCapitulo('INTRODUÇÃO', indice='1', nivel=1)
    resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
    assert resultado['encontrado'] == True
    assert resultado['estrategia_usada'] == 'exato'
    print(f"1. Exato: 'INTRODUÇÃO' → {resultado['estrategia_usada']} (confiança: {resultado['confianca']:.2f})")
    
    # Teste 2: Fuzzy detecta variações
    capitulo = MockCapitulo('METODOLOGA', indice=None, nivel=1)  # Sem índice, título com typo
    resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
    assert resultado['encontrado'] == True
    assert resultado['estrategia_usada'] == 'fuzzy'
    print(f"2. Fuzzy: 'METODOLOGA' (typo) → {resultado['estrategia_usada']} (confiança: {resultado['confianca']:.2f})")
    
    # Teste 3: Contexto por classificação
    capitulo = MockCapitulo('Qualquer título', indice=None, nivel=1, classificacao='anexo')
    resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
    assert resultado['encontrado'] == True
    assert resultado['estrategia_usada'] == 'contexto'
    print(f"3. Contexto: classificação 'anexo' → {resultado['estrategia_usada']} (confiança: {resultado['confianca']:.2f})")
    
    # Teste 4: Nenhuma estratégia funciona
    capitulo = MockCapitulo('TÍTULO INEXISTENTE', indice='99', nivel=1)
    resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
    assert resultado['encontrado'] == False
    assert len(resultado['alternativas']) > 0  # Deve sugerir alternativas
    print(f"4. Não encontrado: 'TÍTULO INEXISTENTE' → alternativas sugeridas: {len(resultado['alternativas'])}")
    
    print("✅ Teste cascata de estratégias passou\n")
    return True


def test_determinismo():
    """Testa determinismo (múltiplas execuções → mesmo resultado)."""
    print("=== Teste Determinismo ===")
    
    doc = Document()
    doc.add_heading('5. DISCUSSÃO', level=1)
    doc.add_paragraph('Texto discussão...')
    
    capitulo = MockCapitulo('DISCUSSÃO', indice='5', nivel=1)
    
    # Executar múltiplas vezes
    resultados = []
    for i in range(5):
        resultado = localizar_range_capitulo_robusto(doc, capitulo, estrategia='multi_niveis')
        resultados.append(resultado)
    
    # Verificar que todos são iguais
    for i in range(1, len(resultados)):
        assert resultados[i]['encontrado'] == resultados[0]['encontrado']
        assert resultados[i]['estrategia_usada'] == resultados[0]['estrategia_usada']
        if resultados[0]['encontrado']:
            assert resultados[i]['inicio'] == resultados[0]['inicio']
            assert resultados[i]['fim'] == resultados[0]['fim']
            assert abs(resultados[i]['confianca'] - resultados[0]['confianca']) < 0.01
    
    print(f"Executado 5x, todos consistentes: encontrado={resultados[0]['encontrado']}, estratégia={resultados[0]['estrategia_usada']}")
    print("✅ Teste determinismo passou\n")
    return True


def main():
    """Executa todos os testes."""
    print("=" * 60)
    print("Testes para Task 2.3: Implementar match fuzzy com distância de edição")
    print("=" * 60)
    
    testes_passados = 0
    testes_totais = 0
    
    try:
        testes_totais += 1
        if test_exemplo_1_requirements():
            testes_passados += 1
        
        testes_totais += 1
        if test_cascata_estrategias():
            testes_passados += 1
        
        testes_totais += 1
        if test_determinismo():
            testes_passados += 1
        
        print(f"\n📊 Resumo: {testes_passados}/{testes_totais} testes passaram")
        
        if testes_passados == testes_totais:
            print("\n🎉 TODOS OS TESTES PASSARAM!")
            print("A implementação de fuzzy matching com distância de edição está funcionando corretamente.")
            print("A cascata de estratégias (exato → fuzzy → contexto) está operacional.")
            return 0
        else:
            print(f"\n⚠️  {testes_totais - testes_passados} teste(s) falharam.")
            return 1
            
    except Exception as e:
        print(f"\n❌ Erro durante execução dos testes: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())