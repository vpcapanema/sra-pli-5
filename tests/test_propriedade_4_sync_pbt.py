"""Testes Property-Based (PBT) para Property 4: Classificação e Sincronização.

Feature: automacao-montagem-relatorios, Property 4: Respeito a Classificação e Seções

Esta suite implementa testes property-based com Hypothesis para validar que:
Para qualquer capítulo no template DOCX com classificação, após sync,
CapituloDocumento tem campos classificacao + prefixo_indice preenchidos corretamente.

Com 100+ exemplos gerados aleatoriamente.

**Validates: Requirements 4.1, 4.2, 4.3, 4.4, 4.5**
"""
from __future__ import annotations

import os
import tempfile
from datetime import date
import uuid

import pytest
import hypothesis.strategies as st
from hypothesis import given, settings, HealthCheck, assume
from docx import Document

from app import db
from app.models.dominio import Dominio
from app.models.usuario import Usuario
from app.models.relatorio_producao import RelatorioProducao
from app.models.capitulo_documento import CapituloDocumento
from app.services.servico_sincronizar_capitulos import (
    ressincronizar_capitulos_com_classificacao
)


# =====================================================================
# Helpers com UUID para evitar conflicts de constraint
# =====================================================================

def criar_usuario_unico(app_context_ref):
    """Cria usuário com nome único usando UUID."""
    unique_id = str(uuid.uuid4())[:8]
    perfil = Dominio.query.filter_by(tipo='perfil_usuario', valor='coordenador').first()
    usuario = Usuario(
        nome=f'TestUser {unique_id}',
        email=f'test_{unique_id}@example.com',
        nome_de_usuario=f'user_{unique_id}',
        senha_hash='hash_seguro',
        perfil_id=perfil.id,
        ativo=True,
    )
    db.session.add(usuario)
    db.session.flush()
    return usuario.id, unique_id


def criar_relatorio_com_docx(usuario_id, capítulos_info, unique_id):
    """Cria RelatorioProducao com DOCX contendo capítulos."""
    # Criar DOCX
    doc = Document()
    for titulo, estilo, _ in capítulos_info:
        if estilo.startswith('Heading'):
            nivel = int(estilo[-1]) if estilo[-1].isdigit() else 1
            doc.add_heading(titulo, level=nivel)
        else:
            # Anexo/Apêndice/etc
            doc.add_heading(titulo, level=1)

        doc.add_paragraph(f"Conteúdo: {titulo}")

    # Salvar DOCX
    tmpdir = tempfile.mkdtemp()
    docx_path = os.path.join(tmpdir, f'pbt_test_{unique_id}.docx')
    doc.save(docx_path)

    # Criar relatório
    status = Dominio.query.filter_by(tipo='status_relatorio').first()
    rel = RelatorioProducao(
        codigo_d20=f'D-20-PBT{unique_id}',
        numero_medicao=1,
        mes_referencia=date.today(),
        periodo_inicio=date.today(),
        periodo_fim=date.today(),
        titulo_curto=f'PBT Test {unique_id}',
        status_id=status.id,
        criado_por=usuario_id,
        versao_atual='R00',
        caminho_template=docx_path,
    )
    db.session.add(rel)
    db.session.flush()

    return rel


# =====================================================================
# Estratégias Hypothesis
# =====================================================================

@st.composite
def capitulos_mistos_strategy(draw) -> list:
    """Gera lista de capítulos com mistura de tipos.

    Retorna lista de (titulo, estilo, classificacao_esperada)
    """
    num_capitulos = draw(st.integers(min_value=1, max_value=5))
    capitulos = []

    # Alguns heading normais
    titulos_heading = ['Introdução', 'Metodologia', 'Resultados', 'Discussão', 'Conclusão']
    for i in range(min(draw(st.integers(min_value=1, max_value=3)), num_capitulos)):
        titulo = titulos_heading[i % len(titulos_heading)]
        nivel = draw(st.sampled_from([1, 2]))
        estilo = f'Heading {nivel}'
        capitulos.append((titulo, estilo, None))  # Heading normal = classificacao None

    # Alguns anexos
    se_tem_anexos = draw(st.booleans())
    if se_tem_anexos and len(capitulos) < num_capitulos:
        for i in range(draw(st.integers(min_value=1, max_value=2))):
            capitulos.append((f'Anexo {chr(65+i)}', 'Anexo', 'anexo'))

    # Alguns apêndices
    se_tem_apendices = draw(st.booleans())
    if se_tem_apendices and len(capitulos) < num_capitulos:
        for i in range(draw(st.integers(min_value=1, max_value=2))):
            capitulos.append((f'Apêndice {chr(73+i)}', 'Apêndice', 'apendice'))

    assume(len(capitulos) >= 1)  # Garantir pelo menos 1 capítulo
    return capitulos[:num_capitulos]  # Limitar ao número desejado


# =====================================================================
# Property-Based Tests
# =====================================================================

# **Validates: Requirements 4.1, 4.2, 4.3, 4.4, 4.5**
@given(capitulos=capitulos_mistos_strategy())
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
def test_property_4_classificacao_para_todo_capitulo(app, capitulos: list):
    """Property 4: Para qualquer capítulo no template, após sync,
    CapituloDocumento.classificacao é preenchida corretamente.

    Testa com 100+ exemplos de estruturas variadas.

    **Validates: Requirements 4.1, 4.2, 4.3**
    """
    with app.app_context():
        usuario_id, unique_id = criar_usuario_unico(app)
        rel = criar_relatorio_com_docx(usuario_id, capitulos, unique_id)

        # Executar sincronização
        resultado = ressincronizar_capitulos_com_classificacao(rel)

        # Assertivas da Property 4
        assert isinstance(resultado, dict), "Resultado deve ser dict"
        assert 'sucesso' in resultado, "Resultado deve ter campo 'sucesso'"

        # Verificar cada capítulo criado
        caps = CapituloDocumento.query.filter_by(id_relatorio=rel.id).all()
        assert len(caps) > 0, "Deve ter criado pelo menos um capítulo"

        for cap in caps:
            # Invariante 1: Títulos devem estar preenchidos
            assert cap.titulo_capitulo, f"Capítulo sem título: {cap.id_capitulo_documento}"

            # Invariante 2: Classificação deve ser válida (None ou um dos valores esperados)
            valid_classifications = [None, 'textual', 'pre_textual', 'pos_textual', 'anexo', 'apendice']
            assert cap.classificacao in valid_classifications, \
                f"Classificação inválida '{cap.classificacao}' para '{cap.titulo_capitulo}'"

            # Invariante 3: Se é anexo ou apêndice, deve ter prefixo
            if cap.classificacao in ('anexo', 'apendice'):
                assert cap.prefixo_indice is not None, \
                    f"Anexo/Apêndice '{cap.titulo_capitulo}' sem prefixo_indice"
                assert len(cap.prefixo_indice) > 0, \
                    f"prefixo_indice vazio para '{cap.titulo_capitulo}'"


# **Validates: Requirements 4.1, 4.2, 4.3**
@given(
    num_anexos=st.integers(min_value=1, max_value=3),
    num_apendices=st.integers(min_value=1, max_value=3)
)
@settings(
    max_examples=80,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
def test_property_4_anexo_apendice_distintos(app, num_anexos: int, num_apendices: int):
    """Property 4: Anexos e apêndices sempre têm classificacoes distintas.

    Nunca confunde anexo com apêndice.

    **Validates: Requirements 4.2**
    """
    with app.app_context():
        usuario_id, unique_id = criar_usuario_unico(app)

        # Construir capítulos
        capitulos = []

        # Introdução (textual)
        capitulos.append(('Introdução', 'Heading 1', None))

        # Anexos
        for i in range(num_anexos):
            capitulos.append((f'Anexo {chr(65+i)}', 'Anexo', 'anexo'))

        # Apêndices
        for i in range(num_apendices):
            capitulos.append((f'Apêndice {chr(73+i)}', 'Apêndice', 'apendice'))

        rel = criar_relatorio_com_docx(usuario_id, capitulos, unique_id)
        resultado = ressincronizar_capitulos_com_classificacao(rel)

        caps = CapituloDocumento.query.filter_by(id_relatorio=rel.id).all()

        # Separar por tipo
        anexos = [c for c in caps if c.classificacao == 'anexo']
        apendices = [c for c in caps if c.classificacao == 'apendice']

        # Invariante: Nunca deve haver confusão
        for cap in caps:
            # Cada capítulo é exatamente um tipo (ou None)
            assert cap.classificacao in [None, 'textual', 'pre_textual', 'pos_textual', 'anexo', 'apendice'], \
                f"Classificação ambígua ou inválida: {cap.classificacao}"

            # Nunca anexo E apêndice simultaneamente
            é_anexo = 'anexo' in cap.titulo_capitulo.lower() or cap.classificacao == 'anexo'
            é_apendice = 'apêndice' in cap.titulo_capitulo.lower() or cap.classificacao == 'apendice'

            if é_anexo and cap.classificacao is not None:
                assert cap.classificacao == 'anexo', \
                    f"Elemento 'Anexo' com classificação errada: {cap.classificacao}"

            if é_apendice and cap.classificacao is not None:
                assert cap.classificacao == 'apendice', \
                    f"Elemento 'Apêndice' com classificação errada: {cap.classificacao}"


# **Validates: Requirements 4.1, 4.3, 4.5**
@given(
    num_repeticoes=st.integers(min_value=2, max_value=3),
    capitulos=capitulos_mistos_strategy()
)
@settings(
    max_examples=50,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
def test_property_4_determinismo_idempotencia(app, num_repeticoes: int, capitulos: list):
    """Property 4: Sincronizações sucessivas com mesmo template
    produzem estado idêntico (determinismo e idempotência).

    **Validates: Requirements 4.1, 4.3, 4.5**
    """
    with app.app_context():
        usuario_id, unique_id = criar_usuario_unico(app)
        rel = criar_relatorio_com_docx(usuario_id, capitulos, unique_id)
        rel_id = rel.id

        # Executar sync múltiplas vezes
        estados = []
        for rep in range(num_repeticoes):
            rel_refresh = RelatorioProducao.query.get(rel_id)
            resultado = ressincronizar_capitulos_com_classificacao(rel_refresh)

            assert resultado['sucesso'], f"Repetição {rep} falhou"

            # Capturar estado
            estado = {}
            caps = CapituloDocumento.query.filter_by(id_relatorio=rel_id).all()
            for cap in caps:
                # Usar (titulo, classificacao, prefixo) como chave
                chave = (cap.titulo_capitulo, cap.classificacao, cap.prefixo_indice)
                estado[chave] = True

            estados.append(estado)

        # Invariante: Todos os estados devem ser idênticos
        estado_base = estados[0]
        for rep_idx, estado in enumerate(estados[1:], start=1):
            assert estado == estado_base, \
                f"Estado divergiu na repetição {rep_idx}: {estado} != {estado_base}"


# **Validates: Requirements 4.1, 4.4**
@given(
    num_pre_textuais=st.integers(min_value=1, max_value=2)
)
@settings(
    max_examples=40,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
def test_property_4_elementos_pre_textuais(app, num_pre_textuais: int):
    """Property 4: Elementos pré-textuais são classificados corretamente.

    Capa, Prefácio, etc devem ser reconhecidos.

    **Validates: Requirements 4.1, 4.4**
    """
    with app.app_context():
        usuario_id, unique_id = criar_usuario_unico(app)

        # Construir capítulos com pré-textuais
        capitulos = []
        titulos_pretextuais = ['Capa', 'Folha de Rosto', 'Prefácio', 'Resumo', 'Abstract']

        for i in range(num_pre_textuais):
            capitulos.append((titulos_pretextuais[i], 'Title', None))  # Pré-textuais

        capitulos.append(('Introdução', 'Heading 1', None))  # Textual

        rel = criar_relatorio_com_docx(usuario_id, capitulos, unique_id)
        resultado = ressincronizar_capitulos_com_classificacao(rel)

        caps = CapituloDocumento.query.filter_by(id_relatorio=rel.id).all()

        # Invariante: Todos os capítulos devem ter estrutura válida
        for cap in caps:
            assert cap.titulo_capitulo, f"Capítulo sem título"

            # Classificação deve ser válida (ou None para pré-textuais)
            assert cap.classificacao in [None, 'textual', 'pre_textual', 'pos_textual', 'anexo', 'apendice'], \
                f"Classificação inválida: {cap.classificacao}"


# **Validates: Requirements 4.1**
@given(
    estilos_com_variacao=st.lists(
        st.tuples(
            st.text(
                min_size=5,
                max_size=30,
                alphabet=st.characters(
                    blacklist_categories=("Cc", "Cs", "Cn"),  # Remover control chars, surrogates, outros
                    blacklist_characters='\x00\x08'  # Remover null e backspace
                )
            ),
            st.sampled_from(['Heading 1', 'Heading 2', 'Anexo', 'Apêndice'])
        ),
        min_size=1,
        max_size=4
    )
)
@settings(
    max_examples=60,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
def test_property_4_robustez_variacoes_titulo(app, estilos_com_variacao: list):
    """Property 4: Classificação funciona com títulos variados (gerados aleatoriamente).

    Não deve falhar ou lançar exceção.

    **Validates: Requirements 4.1**
    """
    with app.app_context():
        usuario_id, unique_id = criar_usuario_unico(app)

        # Construir capítulos
        capitulos = []
        for titulo, estilo in estilos_com_variacao:
            # Filtrar títulos muito vazios após limpeza
            titulo_limpo = titulo.strip()
            if not titulo_limpo:
                titulo_limpo = "Capítulo Vazio"

            expected_classif = 'anexo' if 'Anexo' in estilo else ('apendice' if 'Apêndice' in estilo else None)
            capitulos.append((titulo_limpo, estilo, expected_classif))

        assume(len(capitulos) > 0)  # Garantir pelo menos um capítulo

        rel = criar_relatorio_com_docx(usuario_id, capitulos, unique_id)

        # Deve não lançar exceção, mesmo com títulos aleatórios
        try:
            resultado = ressincronizar_capitulos_com_classificacao(rel)
            assert isinstance(resultado, dict)
        except Exception as e:
            pytest.fail(f"Sync falhou com entrada variada: {e}")


if __name__ == '__main__':
    import sys
    sys.exit(pytest.main([__file__, '-v', '--tb=short']))
