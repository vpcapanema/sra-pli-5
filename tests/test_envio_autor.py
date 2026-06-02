"""Testes do fluxo de envio do autor: upload, prévia e confirmação."""
# pylint: disable=redefined-outer-name
import io
import os
import tempfile

import pytest
from docx import Document

from app import create_app, db
from app.config import Config
from app.models.dominio import Dominio
from app.models.usuario import Usuario
from app.models.relatorio_producao import RelatorioProducao
from app.models.capitulo_documento import CapituloDocumento
from app.services.servico_envio_autor import ServicoEnvioAutor
from app.services.servico_extracao_canonica import ServicoExtracaoCanonica


class TestConfig(Config):
    """Configuração isolada para testes deste módulo."""

    TESTING = True
    SQLALCHEMY_DATABASE_URI = 'sqlite:///:memory:'
    WTF_CSRF_ENABLED = False


@pytest.fixture
def app():
    app = create_app(TestConfig)
    with app.app_context():
        db.create_all()
        # Seed mínimo
        dominio_autor = Dominio()
        dominio_autor.tipo = 'perfil_usuario'
        dominio_autor.valor = 'autor'
        dominio_autor.descricao = 'Autor'
        dominio_coordenador = Dominio()
        dominio_coordenador.tipo = 'perfil_usuario'
        dominio_coordenador.valor = 'coordenador'
        dominio_coordenador.descricao = 'Coord'
        status_relatorio = Dominio()
        status_relatorio.tipo = 'status_relatorio'
        status_relatorio.valor = 'em_producao'
        status_relatorio.descricao = 'EP'
        db.session.add_all([dominio_autor, dominio_coordenador, status_relatorio])
        db.session.commit()
        yield app
        db.session.remove()
        db.drop_all()


@pytest.fixture
def dados(app):
    perfil_autor = Dominio.query.filter_by(
        tipo='perfil_usuario', valor='autor'
    ).first()
    status = Dominio.query.filter_by(tipo='status_relatorio').first()
    assert perfil_autor is not None
    assert status is not None
    u = Usuario()
    u.nome = 'Autor T'
    u.email = 'a@t.com'
    u.nome_de_usuario = 'autort'
    u.senha_hash = 'x'
    u.perfil_id = perfil_autor.id
    u.ativo = True
    db.session.add(u)
    db.session.flush()

    hoje = __import__('datetime').date.today()
    rp = RelatorioProducao()
    rp.codigo_d20 = 'D-20'
    rp.numero_medicao = 1
    rp.mes_referencia = hoje
    rp.periodo_inicio = hoje
    rp.periodo_fim = hoje
    rp.titulo_curto = 'Teste'
    rp.status_id = status.id
    rp.criado_por = u.id
    rp.versao_atual = 'R00'
    db.session.add(rp)
    db.session.flush()

    c1 = CapituloDocumento()
    c1.id_relatorio = rp.id
    c1.titulo_capitulo = 'Introdução'
    c1.ordem_capitulo = 1
    c1.nivel_capitulo = 1
    c1.indice_capitulo = '1'
    c1.tipo_elemento = 'textual'
    c1.id_usuario_responsavel = u.id
    c2 = CapituloDocumento()
    c2.id_relatorio = rp.id
    c2.titulo_capitulo = 'Metodologia'
    c2.ordem_capitulo = 2
    c2.nivel_capitulo = 1
    c2.indice_capitulo = '2'
    c2.tipo_elemento = 'textual'
    c2.id_usuario_responsavel = u.id
    db.session.add_all([c1, c2])
    db.session.commit()
    return {
        'usuario_id': u.id,
        'relatorio_id': rp.id,
        'cap_ids': [c1.id_capitulo_documento, c2.id_capitulo_documento],
    }


def _criar_docx_simples(headings, paragrafos_por_heading=2):
    """Cria um DOCX em memória com cabeçalhos e parágrafos."""
    d = Document()
    for h in headings:
        d.add_heading(h, level=1)
        for k in range(paragrafos_por_heading):
            d.add_paragraph(f'Parágrafo {k+1} de {h}')
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


class _FakeUpload:
    """Imita werkzeug FileStorage para testes."""
    def __init__(self, buf, filename):
        self._buf = buf
        self.filename = filename

    def save(self, path):
        with open(path, 'wb') as f:
            f.write(self._buf.getvalue())


def test_upload_e_previa_geram_segmento_por_capitulo(app, dados, tmp_path):
    buf = _criar_docx_simples(['Introdução', 'Metodologia'])
    fake = _FakeUpload(buf, 'envio.docx')
    envio = ServicoEnvioAutor.processar_upload(
        id_relatorio=dados['relatorio_id'],
        id_usuario=dados['usuario_id'],
        arquivo_storage=fake,
        base_dir=str(tmp_path),
        id_capitulo_destino=dados['cap_ids'][0],
    )
    previas = list(envio.previsualizacoes)
    tipos = {p.tipo_previsualizacao for p in previas}
    assert 'parcial' in tipos
    cap_ids = {
        p.caminho_saida
        for p in previas
        if p.tipo_previsualizacao == 'parcial' and p.caminho_saida
    }
    assert len(cap_ids) == 2


@pytest.mark.skip(reason=(
    'Obsoleto pos-Fase 1: o servico_envio_autor.confirmar nao persiste '
    'mais cap.conteudo_docx; faz merge in-place no DOCX em producao. '
    'Substituido pelo smoke test em test_merge_docx.py.'
))
def test_confirmar_importa_conteudo_nos_capitulos(app, dados, tmp_path):
    assert app
    assert dados
    assert tmp_path


@pytest.mark.skip(reason=(
    'Obsoleto pos-Fase 1: cap.conteudo_docx foi removido; o cancelamento '
    'apenas muda status_envio para rejeitado.'
))
def test_confirmar_rejeitar_nao_altera_capitulos(app, dados, tmp_path):
    assert app
    assert dados
    assert tmp_path


def test_extracao_deduplica_titulos_repetidos(app):
    d = Document()
    d.add_paragraph('SUMÁRIO')
    d.add_paragraph('Introdução')
    d.add_paragraph('Metodologia')
    d.add_heading('Introdução', level=1)
    d.add_paragraph('texto')
    d.add_heading('Metodologia', level=1)
    d.add_paragraph('texto')

    buf = io.BytesIO()
    d.save(buf)
    with tempfile.NamedTemporaryFile(
        suffix='.docx', delete=False
    ) as tmp:
        tmp.write(buf.getvalue())
        caminho = tmp.name
    try:
        doc = Document(caminho)
        arvore = ServicoExtracaoCanonica.extrair_capitulos(doc)
        titulos = [n['titulo'].strip().lower() for n in arvore]
        assert titulos.count('introdução') == 1
        assert titulos.count('metodologia') == 1
    finally:
        os.unlink(caminho)


def test_endpoint_baixar_envio_docx(app, dados, tmp_path):
    """Endpoint /api/envios/<id>/docx devolve o DOCX original."""
    buf = _criar_docx_simples(['Introdução'])
    fake = _FakeUpload(buf, 'envio.docx')
    envio = ServicoEnvioAutor.processar_upload(
        id_relatorio=dados['relatorio_id'],
        id_usuario=dados['usuario_id'],
        arquivo_storage=fake,
        base_dir=str(tmp_path),
        id_capitulo_destino=dados['cap_ids'][0],
    )
    client = app.test_client()
    with client.session_transaction() as sess:
        sess['_user_id'] = str(dados['usuario_id'])
        sess['perfil_ativo'] = 'autor'
        sess['csrf_token'] = 'tk'

    resp = client.get(
        f'/api/envios/{envio.id_envio_conteudo}/docx'
    )
    assert resp.status_code == 200
    assert resp.data[:2] == b'PK'  # DOCX zip


def test_endpoint_baixar_segmento_docx(app, dados, tmp_path):
    """Endpoint /api/envios/<id>/capitulos/<id_cap>/docx devolve
    segmento DOCX classificado por heading."""
    buf = _criar_docx_simples(['Introdução', 'Metodologia'])
    fake = _FakeUpload(buf, 'envio.docx')
    envio = ServicoEnvioAutor.processar_upload(
        id_relatorio=dados['relatorio_id'],
        id_usuario=dados['usuario_id'],
        arquivo_storage=fake,
        base_dir=str(tmp_path),
        id_capitulo_destino=dados['cap_ids'][0],
    )
    client = app.test_client()
    with client.session_transaction() as sess:
        sess['_user_id'] = str(dados['usuario_id'])
        sess['perfil_ativo'] = 'autor'
        sess['csrf_token'] = 'tk'

    resp = client.get(
        f'/api/envios/{envio.id_envio_conteudo}'
        f'/capitulos/{dados["cap_ids"][0]}/docx'
    )
    assert resp.status_code == 200
    assert resp.data[:2] == b'PK'
    # O segmento abre e contém os parágrafos da Introdução
    out_doc = Document(io.BytesIO(resp.data))
    textos = ' '.join(p.text for p in out_doc.paragraphs)
    assert 'Parágrafo' in textos


@pytest.mark.skip(reason=(
    'Endpoint salvar_envio_segmento_docx foi depreciado pos-Fase 1 '
    '(retorna 410 Gone). Substituido pelo fluxo upload + merge.'
))
def test_endpoint_salvar_segmento_docx_html(app, dados, tmp_path):
    assert app
    assert dados
    assert tmp_path


@pytest.mark.skip(reason=(
    '_gerar_docx_versao foi removido pos-Fase 1; o DOCX em producao '
    '(caminho_template) e a fonte unica e e servido direto.'
))
def test_gerar_docx_versao_fallback(app, dados):
    assert app
    assert dados
