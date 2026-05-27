"""Testes para Tarefa 2.1: Função auxiliar de extração de range respeitando seções.

**Feature**: automacao-montagem-relatorios
**Task**: 2.1 Criar função auxiliar de extração de range respeitando seções
**Property**: 2 (Determinismo e Idempotência de Localização)

Testa a função _calcular_range_respeitando_secao que:
1. Encontra próximo heading de nível ≤ nivel_inicio
2. Respeita quebras de seção (sectPr)
3. Retorna dict com inicio, fim, secao_inicio, secao_fim

**Validates: Requirements 2.1, 2.4**
"""
from __future__ import annotations

import hypothesis.strategies as st
from hypothesis import given, settings
from docx import Document
from lxml import etree

from app.services.servico_merge_docx import (
    _calcular_range_respeitando_secao,
    _eh_paragrafo_heading,
    W_NS,
)


@st.composite
def documento_com_secoes_fixture(draw) -> Document:
    """Gera um documento DOCX com múltiplas seções e headings.
    
    Cria um documento com:
    - Múltiplos headings de diferentes níveis
    - Seções (sectPr)
    - Conteúdo variado entre seções
    """
    doc = Document()
    
    num_secoes = draw(st.integers(min_value=1, max_value=3))
    
    for secao_num in range(num_secoes):
        # Adicionar heading principal
        titulo = f"Seção {secao_num + 1}"
        doc.add_heading(titulo, level=1)
        
        # Adicionar conteúdo
        num_paragrafos = draw(st.integers(min_value=1, max_value=3))
        for i in range(num_paragrafos):
            doc.add_paragraph(f"Conteúdo do parágrafo {i+1} da seção {secao_num + 1}")
        
        # Adicionar subheadings
        num_subheadings = draw(st.integers(min_value=0, max_value=2))
        for j in range(num_subheadings):
            doc.add_heading(f"Subseção {secao_num + 1}.{j + 1}", level=2)
            doc.add_paragraph(f"Conteúdo da subseção {secao_num + 1}.{j + 1}")
        
        # Quebra de página (cria nova seção implícita)
        if secao_num < num_secoes - 1:
            doc.add_page_break()
    
    return doc


def test_calcular_range_primeiro_heading():
    """Testa _calcular_range_respeitando_secao com primeiro heading."""
    doc = Document()
    doc.add_heading("Capítulo 1", level=1)
    doc.add_paragraph("Conteúdo 1")
    doc.add_heading("Capítulo 2", level=1)
    doc.add_paragraph("Conteúdo 2")
    
    # Encontrar índice do primeiro heading
    body = doc.element.body
    indice_heading_1 = None
    for i, child in enumerate(body):
        if child.tag == f'{{{W_NS}}}p':
            nivel = _eh_paragrafo_heading(child)
            if nivel == 1:
                indice_heading_1 = i
                break
    
    assert indice_heading_1 is not None, "Heading 1 não encontrado"
    
    # Calcular range
    resultado = _calcular_range_respeitando_secao(doc, indice_heading_1, 1)
    
    # Validações
    assert isinstance(resultado, dict), "Resultado deve ser dict"
    assert 'inicio' in resultado, "'inicio' deve estar no resultado"
    assert 'fim' in resultado, "'fim' deve estar no resultado"
    assert 'secao_inicio' in resultado, "'secao_inicio' deve estar no resultado"
    assert 'secao_fim' in resultado, "'secao_fim' deve estar no resultado"
    assert 'encontrou_limite_secao' in resultado, "'encontrou_limite_secao' deve estar no resultado"
    
    # O resultado deve ter começo e fim
    assert resultado['inicio'] == indice_heading_1, "Início deve ser igual ao índice do heading"
    assert resultado['fim'] >= resultado['inicio'], "Fim deve ser >= início"
    
    print(f"✅ Resultado: {resultado}")


def test_calcular_range_determinismo():
    """Testa determinismo de _calcular_range_respeitando_secao."""
    doc = Document()
    doc.add_heading("Capítulo 1", level=1)
    doc.add_paragraph("Conteúdo 1")
    doc.add_paragraph("Mais conteúdo")
    doc.add_heading("Capítulo 2", level=1)
    doc.add_paragraph("Conteúdo 2")
    
    # Encontrar primeiro heading
    body = doc.element.body
    indice_heading_1 = None
    for i, child in enumerate(body):
        if child.tag == f'{{{W_NS}}}p':
            nivel = _eh_paragrafo_heading(child)
            if nivel == 1:
                indice_heading_1 = i
                break
    
    assert indice_heading_1 is not None
    
    # Executar múltiplas vezes
    resultados = []
    for _ in range(5):
        resultado = _calcular_range_respeitando_secao(doc, indice_heading_1, 1)
        resultados.append(resultado)
    
    # Todos os resultados devem ser idênticos
    primeiro = resultados[0]
    for i, resultado in enumerate(resultados[1:], 1):
        assert resultado == primeiro, f"Execução {i+1} divergiu do resultado 1: {resultado} vs {primeiro}"
    
    print(f"✅ Determinismo validado: todos os {len(resultados)} resultados idênticos")


def test_calcular_range_com_multiplos_niveis():
    """Testa _calcular_range_respeitando_secao com múltiplos níveis de heading."""
    doc = Document()
    doc.add_heading("Capítulo 1", level=1)
    doc.add_paragraph("Conteúdo 1")
    doc.add_heading("Seção 1.1", level=2)
    doc.add_paragraph("Conteúdo 1.1")
    doc.add_heading("Subseção 1.1.1", level=3)
    doc.add_paragraph("Conteúdo 1.1.1")
    doc.add_heading("Capítulo 2", level=1)
    doc.add_paragraph("Conteúdo 2")
    
    # Encontrar heading nível 2
    body = doc.element.body
    indice_heading_2 = None
    for i, child in enumerate(body):
        if child.tag == f'{{{W_NS}}}p':
            nivel = _eh_paragrafo_heading(child)
            if nivel == 2:
                indice_heading_2 = i
                break
    
    assert indice_heading_2 is not None, "Heading nível 2 não encontrado"
    
    # Calcular range com nivel_inicio = 2
    # Deve ir até o próximo heading de nível <= 2 (que é o heading de nível 1)
    resultado = _calcular_range_respeitando_secao(doc, indice_heading_2, 2)
    
    # O fim deve ser ANTES do próximo heading de nível 1
    assert resultado['fim'] < len(body), "Fim deve ser antes do final do documento"
    
    # Verificar que fim aponta para posição antes do heading de nível 1
    fim_elemento = body[resultado['fim']]
    proximo_elemento = body[resultado['fim'] + 1] if resultado['fim'] + 1 < len(body) else None
    
    if proximo_elemento is not None and proximo_elemento.tag == f'{{{W_NS}}}p':
        nivel_proximo = _eh_paragrafo_heading(proximo_elemento)
        if nivel_proximo is not None:
            # O próximo heading deve ter nível menor ou igual que 2
            # já que estamos com nivel_inicio=2
            assert nivel_proximo <= 2, f"Próximo heading deve ter nível <= 2, mas é {nivel_proximo}"
    
    print(f"✅ Múltiplos níveis: resultado={resultado}")


@given(
    doc=documento_com_secoes_fixture(),
    nivel_teste=st.integers(min_value=1, max_value=3)
)
@settings(max_examples=30, deadline=5000)
def test_property_2_range_respeitando_secao_determinismo(doc: Document, nivel_teste: int):
    """Property-based test para determinismo de _calcular_range_respeitando_secao.
    
    Para qualquer documento e nível de heading, múltiplas execuções devem
    retornar exatamente o mesmo resultado.
    
    **Validates: Property 2: Determinismo e Idempotência de Localização**
    """
    # Encontrar um heading do nível desejado
    body = doc.element.body
    indice_heading = None
    nivel_encontrado = None
    
    for i, child in enumerate(body):
        if child.tag == f'{{{W_NS}}}p':
            nivel = _eh_paragrafo_heading(child)
            if nivel is not None and nivel <= nivel_teste:
                indice_heading = i
                nivel_encontrado = nivel
                break
    
    # Se não encontrou heading do nível desejado, pular este exemplo
    if indice_heading is None:
        return
    
    # Executar a função múltiplas vezes
    resultados = []
    num_execucoes = 5
    
    for _ in range(num_execucoes):
        resultado = _calcular_range_respeitando_secao(doc, indice_heading, nivel_teste)
        resultados.append(resultado)
    
    # Verificar que todos os resultados são idênticos
    primeiro = resultados[0]
    for resultado in resultados[1:]:
        assert resultado == primeiro, f"Divergência detectada: {resultado} vs {primeiro}"
    
    # Validar estrutura do resultado
    assert primeiro['inicio'] == indice_heading, "'inicio' deve ser igual ao índice do heading"
    assert primeiro['fim'] >= primeiro['inicio'], "'fim' deve ser >= 'inicio'"
    assert primeiro['secao_inicio'] >= 0, "'secao_inicio' deve ser >= 0"
    assert primeiro['secao_fim'] >= primeiro['secao_inicio'], "'secao_fim' deve ser >= 'secao_inicio'"
    assert isinstance(primeiro['encontrou_limite_secao'], bool), "'encontrou_limite_secao' deve ser bool"


@given(
    doc=documento_com_secoes_fixture()
)
@settings(max_examples=20, deadline=5000)
def test_property_2_range_coerencia_estrutura(doc: Document):
    """Testa que o resultado sempre tem estrutura coerente.
    
    Para qualquer documento, _calcular_range_respeitando_secao deve retornar
    um dict com campos específicos e valores válidos.
    
    **Validates: Property 3: Coerência de Estrutura Retornada**
    """
    # Encontrar primeiro heading
    body = doc.element.body
    indice_heading = None
    
    for i, child in enumerate(body):
        if child.tag == f'{{{W_NS}}}p':
            nivel = _eh_paragrafo_heading(child)
            if nivel is not None:
                indice_heading = i
                break
    
    # Se não encontrou heading, pular
    if indice_heading is None:
        return
    
    # Calcular range
    resultado = _calcular_range_respeitando_secao(doc, indice_heading, 1)
    
    # Validar estrutura
    campos_obrigatorios = [
        'inicio',
        'fim',
        'secao_inicio',
        'secao_fim',
        'encontrou_limite_secao'
    ]
    
    for campo in campos_obrigatorios:
        assert campo in resultado, f"Campo '{campo}' faltando no resultado"
    
    # Validar tipos
    assert isinstance(resultado['inicio'], int), "'inicio' deve ser int"
    assert isinstance(resultado['fim'], int), "'fim' deve ser int"
    assert isinstance(resultado['secao_inicio'], int), "'secao_inicio' deve ser int"
    assert isinstance(resultado['secao_fim'], int), "'secao_fim' deve ser int"
    assert isinstance(resultado['encontrou_limite_secao'], bool), "'encontrou_limite_secao' deve ser bool"
    
    # Validar invariantes
    assert resultado['inicio'] >= 0, "'inicio' deve ser >= 0"
    assert resultado['fim'] >= resultado['inicio'], "'fim' deve ser >= 'inicio'"
    assert resultado['secao_inicio'] >= 0, "'secao_inicio' deve ser >= 0"
    assert resultado['secao_fim'] >= resultado['secao_inicio'], "'secao_fim' deve ser >= 'secao_inicio'"


if __name__ == '__main__':
    import pytest
    import sys
    
    print("\n" + "="*70)
    print("Testes para Tarefa 2.1: Função auxiliar de extração de range")
    print("="*70 + "\n")
    
    # Executar testes unitários primeiro
    print("Executando testes unitários...\n")
    
    try:
        test_calcular_range_primeiro_heading()
        test_calcular_range_determinismo()
        test_calcular_range_com_multiplos_niveis()
        print("\n✅ Todos os testes unitários passaram")
    except Exception as e:
        print(f"\n❌ Erro em teste unitário: {e}")
        import traceback
        traceback.print_exc()
    
    print("\n" + "="*70)
    print("Executando testes property-based...\n")
    
    # Executar testes property-based com pytest
    sys.exit(pytest.main([__file__, '-v', '--tb=short']))
