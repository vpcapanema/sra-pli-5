"""Debug test for _match_contexto."""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)) + '/..')

from docx import Document
from app.services.servico_merge_docx import _match_contexto
import re

class MockCapitulo:
    def __init__(self, titulo, indice=None, nivel=None, classificacao=None, tipo_elemento=None):
        self.titulo_capitulo = titulo
        self.indice_capitulo = indice
        self.nivel_capitulo = nivel
        self.classificacao = classificacao
        self.tipo_elemento = tipo_elemento

# Test regex extraction
print("Testando extração de número do índice:")
indice = '2'
match = re.match(r'^(\d+)', indice)
if match:
    print(f"Índice '{indice}' → número: {match.group(1)}")
else:
    print(f"Índice '{indice}' → sem match")

# Test with actual document
doc = Document()
doc.add_heading('1. INTRODUÇÃO', level=1)
doc.add_paragraph('Texto...')
doc.add_heading('2. METODOLOGIA', level=1)
doc.add_paragraph('Texto...')

capitulo = MockCapitulo('Qualquer título', indice='2', nivel=1)
resultado = _match_contexto(doc, capitulo)

print(f"\nResultado: encontrado={resultado['encontrado']}")
print(f"Diagnóstico: {resultado['diagnostico']}")
print(f"Confiança: {resultado['confianca']}")

# Check what's in the document
print("\nHeadings no documento:")
body = doc.element.body
for i, child in enumerate(body):
    if child.tag.endswith('}p'):
        from app.services.servico_merge_docx import _texto_paragrafo, _eh_paragrafo_heading
        nivel = _eh_paragrafo_heading(child)
        if nivel is not None:
            texto = _texto_paragrafo(child)
            print(f"  [{i}] Nível {nivel}: '{texto}'")