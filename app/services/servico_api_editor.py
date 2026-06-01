"""Servicos usados pela API do editor colaborativo."""
from __future__ import annotations

import json
import os
from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document

from app import db
from app.models.biblioteca_formatacao import BibliotecaFormatacaoCanonica
from app.models.capitulo_documento import CapituloDocumento
from app.models.dominio import Dominio
from app.models.envio_conteudo import EnvioConteudo
from app.models.previsualizacao_conteudo import PrevisualizacaoConteudo
from app.models.relatorio_producao import RelatorioProducao
from app.models.usuario import Usuario
from app.services.servico_finalizar_relatorio import finalizar
from app.services.servico_sanitizar_docx import sanitizar_docx, sanitizar_docx_bytes


MAX_UPLOAD_BYTES = 50 * 1024 * 1024


def obter_versao_api(id_vt):
    """Retorna versao de trabalho serializada com capitulos raiz."""
    vt = RelatorioProducao.query.get_or_404(id_vt)
    caps = _listar_capitulos_raiz(id_vt)
    return {
        'id': vt.id,
        'titulo': vt.titulo_curto,
        'status': vt.status.codigo if vt.status else None,
        'id_biblioteca': None,
        'capitulos': [serializar_capitulo(c) for c in caps],
    }


def listar_capitulos_api(id_vt):
    """Lista capitulos raiz serializados."""
    return [serializar_capitulo(c) for c in _listar_capitulos_raiz(id_vt)]


def criar_capitulo_api(id_vt, dados, id_usuario):
    """Cria capitulo pela API."""
    RelatorioProducao.query.get_or_404(id_vt)
    titulo = dados['titulo']
    cap = CapituloDocumento(
        id_relatorio=id_vt,
        titulo_capitulo=titulo,
        nome_capitulo=dados.get('nome') or titulo,
        ordem_capitulo=dados.get('ordem', 0),
        nivel_capitulo=dados.get('nivel', 1),
        id_capitulo_pai=dados.get('id_pai'),
        id_usuario_responsavel=dados.get('id_responsavel'),
        status_capitulo='em_edicao',
        criado_por=id_usuario,
    )
    db.session.add(cap)
    db.session.commit()
    return serializar_capitulo(cap)


def atualizar_capitulo_api(id_cap, dados, id_usuario):
    """Atualiza metadados de capitulo pela API."""
    cap = CapituloDocumento.query.get_or_404(id_cap)
    if 'titulo' in dados:
        cap.titulo_capitulo = dados['titulo']
        if not cap.nome_capitulo:
            cap.nome_capitulo = dados['titulo']
    if 'ordem' in dados:
        cap.ordem_capitulo = dados['ordem']
    if 'nivel' in dados:
        cap.nivel_capitulo = dados['nivel']
    if 'id_responsavel' in dados:
        cap.id_usuario_responsavel = dados['id_responsavel']
    cap.atualizado_por = id_usuario
    db.session.commit()
    return serializar_capitulo(cap)


def excluir_capitulo_api(id_cap):
    """Marca capitulo como inativo."""
    cap = CapituloDocumento.query.get_or_404(id_cap)
    cap.ativo = False
    db.session.commit()


def usuario_pode_acessar_envio(envio, id_usuario, perfil_ativo):
    """Verifica acesso ao envio do autor."""
    return envio.id_usuario == id_usuario or perfil_ativo in ('coordenador', 'admin')


def obter_estrutura_envio(id_envio):
    """Retorna estrutura processada de um envio."""
    envio = EnvioConteudo.query.get_or_404(id_envio)
    try:
        if envio.sugestoes_json:
            return json.loads(envio.sugestoes_json)
        return {'capitulos': [], 'legendas': []}
    except (ValueError, TypeError) as erro:
        raise ValueError(f'Erro ao carregar estrutura: {erro}') from erro


def atualizar_renomeacao_envio(id_envio, id_capitulo, aprovado):
    """Atualiza aprovacao de renomeacao pendente."""
    envio = EnvioConteudo.query.get_or_404(id_envio)
    try:
        estrutura = json.loads(envio.sugestoes_json or '{}')
    except (ValueError, TypeError):
        estrutura = {}
    pendentes = estrutura.get('renomeacoes_pendentes') or []
    for renomeacao in pendentes:
        if renomeacao.get('id_capitulo_documento') == id_capitulo:
            renomeacao['aprovado'] = aprovado
            estrutura['renomeacoes_pendentes'] = pendentes
            envio.sugestoes_json = json.dumps(estrutura)
            db.session.commit()
            return {
                'ok': True,
                'id_capitulo_documento': id_capitulo,
                'aprovado': aprovado,
            }
    return None


def listar_segmentos_envio(id_envio):
    """Lista segmentos de previa de um envio."""
    previas = PrevisualizacaoConteudo.query.filter_by(
        id_envio_conteudo=id_envio
    ).all()
    segmentos = []
    for previa in previas:
        cap = _capitulo_de_previa(previa)
        if cap:
            segmentos.append({
                'id_previsualizacao': previa.id_previsualizacao,
                'id_capitulo': cap.id_capitulo_documento,
                'titulo_capitulo': cap.titulo_capitulo,
                'indice_capitulo': cap.indice_capitulo,
                'resultado_html': previa.resultado_html,
                'tipo_previsualizacao': previa.tipo_previsualizacao,
            })
    return segmentos


def gerar_segmento_docx(id_envio, id_capitulo):
    """Gera bytes DOCX de um segmento de envio."""
    from app.services.servico_envio_autor import gerar_docx_segmento

    envio = EnvioConteudo.query.get_or_404(id_envio)
    cap = CapituloDocumento.query.get_or_404(id_capitulo)
    if cap.id_relatorio != envio.id_relatorio:
        raise ValueError('Capítulo não pertence ao envio')
    return gerar_docx_segmento(envio, cap)


def converter_html_para_docx(dados):
    """Converte HTML simples do editor para bytes DOCX."""
    soup = BeautifulSoup(dados.decode('utf-8'), 'html.parser')
    doc = Document()
    for el in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
        txt = el.get_text(strip=True)
        if not txt:
            continue
        if el.name.startswith('h'):
            nivel = int(el.name[1])
            doc.add_heading(txt, level=min(nivel, 9))
        else:
            doc.add_paragraph(txt)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def extrair_conteudo_capitulo(id_cap):
    """Extrai e sanitiza DOCX de capitulo a partir do documento em producao."""
    from app.services.servico_merge_docx import extrair_capitulo_como_docx

    cap = CapituloDocumento.query.get_or_404(id_cap)
    rel = RelatorioProducao.query.get(cap.id_relatorio)
    if not rel or not rel.caminho_template:
        return None, 'Relatorio sem DOCX em producao', 404
    if not os.path.exists(rel.caminho_template):
        return None, 'Arquivo DOCX em producao nao encontrado', 404
    conteudo = extrair_capitulo_como_docx(rel.caminho_template, cap)
    if not conteudo:
        return None, 'Capitulo nao localizado no DOCX em producao', 404
    return sanitizar_docx_bytes(conteudo) or conteudo, '', 200


def finalizar_relatorio_api(id_relatorio, id_usuario):
    """Finaliza relatorio de producao e retorna dados serializados."""
    rf = finalizar(id_relatorio=id_relatorio, id_usuario=id_usuario)
    return {
        'mensagem': 'Relatorio finalizado com sucesso',
        'id_finalizado': rf.id,
        'versao': rf.versao,
        'checksum': rf.checksum_docx,
        'nome_arquivo': rf.nome_arquivo,
    }


def finalizar_capitulo_api(id_cap, id_usuario):
    """Autor marca capitulo como finalizado."""
    cap = CapituloDocumento.query.get_or_404(id_cap)
    if cap.status_capitulo not in ('em_edicao', 'reprovado'):
        return None, 'Status inválido para finalizar', 400
    if cap.id_usuario_responsavel and cap.id_usuario_responsavel != id_usuario:
        return None, 'Sem permissão', 403
    cap.status_capitulo = 'finalizado'
    db.session.commit()
    return serializar_capitulo(cap), '', 200


def aprovar_capitulo_api(id_cap):
    """Coordenador aprova capitulo."""
    cap = CapituloDocumento.query.get_or_404(id_cap)
    if cap.status_capitulo != 'finalizado':
        return None, 'Capítulo não está finalizado', 400
    cap.status_capitulo = 'aprovado'
    cap.observacao_coordenador = None
    if cap.id_usuario_responsavel:
        notificar(
            cap.id_usuario_responsavel,
            f'Capítulo "{cap.titulo_capitulo}" aprovado pelo coordenador.',
        )
    db.session.commit()
    return serializar_capitulo(cap), '', 200


def reprovar_capitulo_api(id_cap, observacao):
    """Coordenador reprova capitulo."""
    cap = CapituloDocumento.query.get_or_404(id_cap)
    if cap.status_capitulo != 'finalizado':
        return None, 'Capítulo não está finalizado', 400
    cap.status_capitulo = 'reprovado'
    cap.observacao_coordenador = observacao or ''
    if cap.id_usuario_responsavel:
        notificar(
            cap.id_usuario_responsavel,
            f'Capítulo "{cap.titulo_capitulo}" reprovado. '
            f'Observação: {cap.observacao_coordenador}',
        )
    db.session.commit()
    return serializar_capitulo(cap), '', 200


def servir_docx_producao(vt):
    """Retorna bytes sanitizados ou caminho do DOCX em producao."""
    if not vt.caminho_template or not os.path.exists(vt.caminho_template):
        return None, 'DOCX em producao indisponivel'
    return sanitizar_docx(vt.caminho_template) or vt.caminho_template, ''


def vincular_biblioteca_api(id_vt, id_bib):
    """Valida vinculo de biblioteca em versao de trabalho."""
    RelatorioProducao.query.get_or_404(id_vt)
    if not id_bib:
        return False, 'id_biblioteca obrigatório'
    BibliotecaFormatacaoCanonica.query.get_or_404(id_bib)
    db.session.commit()
    return True, ''


def listar_autores_api():
    """Lista autores ativos serializados."""
    perfil_autor = Dominio.query.filter_by(
        tipo="perfil_usuario", valor="autor"
    ).first()
    if not perfil_autor:
        return []
    autores = (
        Usuario.query.filter(
            Usuario.perfil_id == perfil_autor.id_dominio,
            Usuario.ativo == True,  # noqa: E712
        )
        .order_by(Usuario.nome)
        .all()
    )
    return [
        {
            'id': autor.id,
            'nome': autor.nome,
            'email': autor.email,
        }
        for autor in autores
    ]


def notificar(id_usuario, mensagem):
    """Cria notificacao simples para usuario."""
    from app.models.notificacao import Notificacao

    notif = Notificacao(
        id_usuario_destino=id_usuario,
        tipo_notificacao='workflow',
        mensagem=mensagem,
    )
    db.session.add(notif)


def serializar_capitulo(cap):
    """Serializa capitulo com filhos."""
    filhos = CapituloDocumento.query.filter_by(
        id_capitulo_pai=cap.id_capitulo_documento,
        ativo=True,
    ).order_by(CapituloDocumento.ordem_capitulo).all()
    return {
        'id': cap.id_capitulo_documento,
        'titulo': cap.titulo_capitulo,
        'ordem': cap.ordem_capitulo,
        'nivel': cap.nivel_capitulo,
        'status': cap.status_capitulo,
        'id_responsavel': cap.id_usuario_responsavel,
        'responsavel_nome': cap.responsavel.nome if cap.responsavel else None,
        'tem_conteudo': True,
        'observacao_coordenador': cap.observacao_coordenador,
        'filhos': [serializar_capitulo(filho) for filho in filhos],
    }


def _listar_capitulos_raiz(id_vt):
    return CapituloDocumento.query.filter_by(
        id_relatorio=id_vt,
        id_capitulo_pai=None,
        ativo=True,
    ).order_by(CapituloDocumento.ordem_capitulo).all()


def _capitulo_de_previa(previa):
    if not previa.caminho_saida:
        return None
    try:
        return CapituloDocumento.query.get(int(previa.caminho_saida))
    except (ValueError, TypeError):
        return None
