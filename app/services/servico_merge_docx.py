"""Serviço de mescla in-place de conteúdo do autor no DOCX em produção.

Fluxo principal (Fase 1 do plano):

1. `localizar_range_capitulo`: dado o DOCX em produção e um
   `CapituloDocumento`, identifica o intervalo [inicio, fim] (em
   índices relativos ao corpo do XML, NÃO a paragraphs[]) que
   corresponde ao capítulo (heading + conteúdo até antes do próximo
   heading de nível <= ao do capítulo).
2. `substituir_capitulo`: substitui o conteúdo do capítulo no DOCX
   de produção pelo conteúdo do DOCX do autor, preservando heading,
   estilos, seções e imagens (rIds remapeados).

Decisão arquitetural: a tabela `capitulos_documento` NÃO armazena
mais o conteúdo (`conteudo_docx` será removido em migração futura).
A fonte única de verdade do conteúdo é o `.docx` em produção em
`storage/relatorios_producao/<arquivo>.docx`.
"""
from __future__ import annotations

import difflib
import re
import unicodedata
from io import BytesIO
from typing import Optional

import lxml.etree as etree

from docx import Document
from docxcompose.composer import Composer

from app.models.capitulo_documento import CapituloDocumento
from app.services._ooxml_helpers import texto_paragrafo as _texto_paragrafo_base
from app.services.servico_nivelador_erros import ServicoNiveladorErros


W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _normalizar(texto: str) -> str:
    """Lowercase + sem acentos + colapsa espaços + tira numeração inicial."""
    if not texto:
        return ''
    s = unicodedata.normalize('NFD', texto)
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    s = re.sub(r'\s+', ' ', s).strip().lower()
    s = re.sub(r'^\s*(?:\d+(?:\.\d+)*|[ivx]+|[a-z])[\.\)]?\s+', '', s)
    return s


def _heading_nivel_de_estilo(style_name: str) -> Optional[int]:
    """Retorna nível 1..9 se o estilo for um heading, senão None.

    Aceita variações de styleId/styleName comuns:
    - Inglês: 'Heading 1', 'Heading1', 'heading 2', ...
    - Português: 'Título 1', 'Titulo 1', 'Ttulo1' (styleId Word
      sem acentos), 'Título1', ...
    - Title (estilo de título do documento) → nível 0.
    """
    if not style_name:
        return None
    # Normaliza: lowercase + sem acentos + sem espaços
    s = unicodedata.normalize('NFD', style_name)
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    s = s.strip().lower().replace(' ', '')

    for prefixo in ('heading', 'titulo', 'ttulo'):
        if s.startswith(prefixo):
            sufixo = s[len(prefixo):].strip()
            if not sufixo:
                return 1
            try:
                return int(sufixo)
            except ValueError:
                return None
    if s in ('title',):
        return 0
    return None


def _eh_paragrafo_heading(p_element) -> Optional[int]:
    """Inspeciona o elemento <w:p> e retorna o nível do heading
    (1..9) se for um heading, senão None.

    Olha o `<w:pStyle w:val="Heading N"/>` dentro de `<w:pPr>`.
    """
    pPr = p_element.find(f'{{{W_NS}}}pPr')
    if pPr is None:
        return None
    pStyle = pPr.find(f'{{{W_NS}}}pStyle')
    if pStyle is None:
        return None
    val = pStyle.get(f'{{{W_NS}}}val', '')
    return _heading_nivel_de_estilo(val)


def _texto_paragrafo(p_element) -> str:
    """Extrai texto plano de um `<w:p>` (concatena todos os `<w:t>`)."""
    return _texto_paragrafo_base(p_element, strip=True)


def _calcular_range_respeitando_secao(doc, indice_inicio: int, nivel_inicio: int) -> dict:
    """Calcula range respeitando limites de seção.

    Encontra próximo heading de nível ≤ nivel_inicio, respeitando quebras
    de seção (sectPr). Retorna dict com informações do range.

    Args:
        doc: Documento python-docx
        indice_inicio: Índice do elemento de início (heading do capítulo)
        nivel_inicio: Nível do heading do capítulo

    Returns:
        Dict com:
            - inicio: índice do heading do capítulo
            - fim: índice do último elemento antes do próximo heading ou seção
            - secao_inicio: índice da seção onde o capítulo começa
            - secao_fim: índice da seção onde o capítulo termina
            - encontrou_limite_secao: True se encontrou sectPr antes de heading
    """
    body = doc.element.body
    total_elementos = len(body)

    # Encontrar a seção atual (procurar sectPr anterior mais próximo)
    secao_atual = 0
    for i in range(indice_inicio, -1, -1):
        if body[i].tag == f'{{{W_NS}}}sectPr':
            secao_atual = i
            break

    # Procurar fim: próximo heading de nível ≤ nivel_inicio OU próximo sectPr
    fim = None
    encontrou_limite_secao = False
    secao_fim = secao_atual

    for j in range(indice_inicio + 1, total_elementos):
        child = body[j]

        # Verificar se é sectPr (fim da seção atual)
        if child.tag == f'{{{W_NS}}}sectPr':
            fim = j - 1
            encontrou_limite_secao = True
            # A seção termina neste sectPr
            secao_fim = j
            break

        # Verificar se é heading de nível ≤ nivel_inicio
        if child.tag == f'{{{W_NS}}}p':
            nivel_j = _eh_paragrafo_heading(child)
            if nivel_j is not None and nivel_j <= nivel_inicio:
                fim = j - 1
                break

    # Se não encontrou limite, vai até o final (excluindo sectPr final)
    if fim is None:
        # Encontrar último elemento que não é sectPr
        for k in range(total_elementos - 1, indice_inicio, -1):
            if body[k].tag != f'{{{W_NS}}}sectPr':
                fim = k
                break
        if fim is None:
            fim = indice_inicio

    return {
        'inicio': indice_inicio,
        'fim': fim,
        'secao_inicio': secao_atual,
        'secao_fim': secao_fim,
        'encontrou_limite_secao': encontrou_limite_secao
    }


def _localizar_range_capitulo_interno(doc, capitulo) -> Optional[tuple[int, int]]:
    """Implementação interna de localizar_range_capitulo."""
    body = doc.element.body
    titulo_alvo = _normalizar(capitulo.titulo_capitulo)
    indice_alvo = (capitulo.indice_capitulo or '').strip()
    nivel_alvo = capitulo.nivel_capitulo or 1

    inicio = None
    for i, child in enumerate(body):
        if child.tag != f'{{{W_NS}}}p':
            continue
        nivel = _eh_paragrafo_heading(child)
        if nivel is None:
            continue
        texto = _texto_paragrafo(child)
        texto_norm = _normalizar(texto)

        casou_titulo = bool(titulo_alvo) and texto_norm == titulo_alvo
        casou_indice = (
            bool(indice_alvo)
            and texto.lstrip().startswith(indice_alvo)
        )
        if casou_titulo or casou_indice:
            inicio = i
            nivel_alvo_real = nivel  # confia no nível encontrado
            break
    # nivel_alvo (vindo do parâmetro) é apenas dica; o nível real
    # usado para delimitar o fim é o do heading encontrado.
    _ = nivel_alvo  # silenciar warning de variável não usada

    if inicio is None:
        return None

    # Fim: último elemento antes do próximo heading com nível <=
    fim = None
    for j in range(inicio + 1, len(body)):
        child = body[j]
        if child.tag == f'{{{W_NS}}}p':
            nivel_j = _eh_paragrafo_heading(child)
            if nivel_j is not None and nivel_j <= nivel_alvo_real:
                fim = j - 1
                break
        if child.tag == f'{{{W_NS}}}sectPr':
            fim = j - 1
            break

    if fim is None:
        # Foi até o fim sem encontrar próximo heading; fim é o último
        # elemento que NÃO é sectPr.
        ultimos = list(range(inicio + 1, len(body)))
        while ultimos and body[ultimos[-1]].tag == f'{{{W_NS}}}sectPr':
            ultimos.pop()
        fim = ultimos[-1] if ultimos else inicio

    return (inicio, fim)


def localizar_range_capitulo(doc, capitulo, relatorio_id=None, capitulo_id=None) -> Optional[tuple[int, int]]:
    """Localiza o range de elementos do corpo (`body`) que pertencem
    ao capítulo informado.

    Retorna `(inicio, fim)` como índices inclusivos em
    `doc.element.body[*]`, onde `inicio` é o parágrafo do heading
    e `fim` é o último elemento (parágrafo ou tabela) antes do
    próximo heading de nível <= `capitulo.nivel_capitulo`, ou o
    último elemento do corpo (excluindo `<w:sectPr>`).

    Retorna `None` se o capítulo não for localizado.

    Estratégia de casamento:
    - Normaliza `capitulo.titulo_capitulo` e compara com texto
      normalizado de cada parágrafo que tem estilo `Heading N`.
    - Quando o `indice_capitulo` está presente, tenta também
      casar pelo prefixo `"<indice> <titulo>"` no texto do
      parágrafo (caso o DOCX traga numeração explícita).
    - Pega o primeiro match (DOCX bem formatado não duplica
      headings de capítulo no mesmo nível).
    """
    resultado = ServicoNiveladorErros.executar_com_tratamento(
        _localizar_range_capitulo_interno,
        doc, capitulo,
        relatorio_id=relatorio_id,
        capitulo_id=capitulo_id,
        etapa='localizacao_capitulo'
    )

    # Se o resultado é um dict de erro, retorna None
    if isinstance(resultado, dict) and not resultado.get('sucesso', True):
        return None

    return resultado


def _listar_subheadings_no_range_interno(
    doc, inicio: int, fim: int, nivel_pai: int
) -> list[dict]:
    """Implementação interna de listar_subheadings_no_range."""
    body = doc.element.body
    encontrados = []
    # Saltar o próprio heading do capítulo pai (em body[inicio])
    for k in range(inicio + 1, min(fim + 1, len(body))):
        child = body[k]
        if child.tag != f'{{{W_NS}}}p':
            continue
        nivel = _eh_paragrafo_heading(child)
        if nivel is None or nivel <= nivel_pai:
            continue
        texto = _texto_paragrafo(child)
        if not texto:
            continue
        # Tentar extrair prefixo numérico do título (ex.: "5.4.6.1
        # Sistema" -> indice "5.4.6.1", título "Sistema").
        indice = None
        m = re.match(
            r'^\s*(\d+(?:\.\d+)*)\s*[\.\)\-–—:]?\s+(.+?)\s*$',
            texto,
        )
        if m:
            indice = m.group(1)
            titulo_limpo = m.group(2).strip()
        else:
            titulo_limpo = texto
        encontrados.append({
            'titulo': titulo_limpo,
            'indice': indice,
            'nivel': nivel,
            'indice_no_body': k,
        })
    return encontrados


def listar_subheadings_no_range(
    doc, inicio: int, fim: int, nivel_pai: int,
    relatorio_id: Optional[int] = None,
    capitulo_id: Optional[int] = None,
) -> list[dict]:
    """Varre o range [inicio..fim] do `body` (incluindo o heading do
    capítulo pai) e devolve uma lista de subheadings encontrados,
    em ordem de aparecimento.

    Cada item: {'titulo': str, 'nivel': int, 'indice_no_body': int}
    onde `nivel` é o do heading no DOCX (>= nivel_pai + 1) e
    `indice_no_body` é a posição do parágrafo dentro de `body`.

    Inclui subheadings de TODOS os níveis abaixo de `nivel_pai`
    (não só nivel_pai+1) — a hierarquia entre eles é resolvida na
    montagem da árvore (sincronizar_subcapitulos).
    """
    resultado = ServicoNiveladorErros.executar_com_tratamento(
        _listar_subheadings_no_range_interno,
        doc, inicio, fim, nivel_pai,
        relatorio_id=relatorio_id,
        capitulo_id=capitulo_id,
        etapa='listagem_subheadings'
    )

    # Se o resultado é um dict de erro, retorna lista vazia
    if isinstance(resultado, dict) and not resultado.get('sucesso', True):
        return []

    return resultado


def _sincronizar_subcapitulos_interno(
    db_session,
    capitulo_pai,
    caminho_docx: str,
) -> dict:
    """Implementação interna de sincronizar_subcapitulos."""
    doc = Document(caminho_docx)
    rng = localizar_range_capitulo(doc, capitulo_pai)
    if rng is None:
        return {'inseridos': 0, 'atualizados': 0, 'desativados': 0,
                'erro': 'capitulo_pai não localizado no DOCX'}
    inicio, fim = rng

    nivel_pai = capitulo_pai.nivel_capitulo or 1
    detectados = listar_subheadings_no_range(doc, inicio, fim, nivel_pai)

    # Filhos atuais do capítulo pai no banco
    filhos_db = (
        CapituloDocumento.query
        .filter_by(
            id_relatorio=capitulo_pai.id_relatorio,
            id_capitulo_pai=capitulo_pai.id_capitulo_documento,
        )
        .all()
    )
    filhos_por_titulo = {
        _normalizar(f.titulo_capitulo): f for f in filhos_db
    }
    titulos_detectados = set()

    inseridos = atualizados = 0
    indice_pai = (capitulo_pai.indice_capitulo or '').strip()
    for ordem, item in enumerate(detectados, start=1):
        chave = _normalizar(item['titulo'])
        titulos_detectados.add(chave)
        if chave in filhos_por_titulo:
            f = filhos_por_titulo[chave]
            f.ordem_capitulo = ordem
            f.ativo = True
            atualizados += 1
        else:
            indice_filho = item.get('indice')
            if not indice_filho and indice_pai:
                indice_filho = f'{indice_pai}.{ordem}'
            elif not indice_filho:
                indice_filho = str(ordem)
            from app.utils.auditoria import usuario_atual_id  # noqa: C0415
            novo = CapituloDocumento(
                id_relatorio=capitulo_pai.id_relatorio,
                id_capitulo_pai=capitulo_pai.id_capitulo_documento,
                titulo_capitulo=item['titulo'],
                # `nome_capitulo` espelha o titulo na criacao automatica.
                nome_capitulo=item['titulo'],
                ordem_capitulo=ordem,
                nivel_capitulo=nivel_pai + 1,
                tipo_elemento=capitulo_pai.tipo_elemento or 'textual',
                indice_capitulo=indice_filho,
                status_capitulo='em_edicao',
                ativo=True,
                # Subcapitulo detectado a partir do upload do autor —
                # preenche com o usuario logado (autor) se houver
                # request context; senao fica como sistema.
                criado_por=usuario_atual_id(),
            )
            db_session.add(novo)
            inseridos += 1

    # Desativar filhos que sumiram do DOCX
    desativados = 0
    for chave, f in filhos_por_titulo.items():
        if chave not in titulos_detectados and f.ativo:
            f.ativo = False
            desativados += 1

    db_session.flush()
    return {
        'inseridos': inseridos,
        'atualizados': atualizados,
        'desativados': desativados,
    }


def sincronizar_subcapitulos(
    db_session,
    capitulo_pai,
    caminho_docx: str,
    relatorio_id: Optional[int] = None,
    capitulo_id: Optional[int] = None,
) -> dict:
    """Após um merge no DOCX em produção, sincroniza os subcapítulos
    no banco a partir dos subheadings detectados dentro do range do
    `capitulo_pai` no DOCX.

    Política:
    - Subheadings novos (não presentes no banco) → INSERT como filhos
      diretos de `capitulo_pai` (nivel = nivel_pai + 1, idependente do
      nivel real no heading — uma vez que o autor pode ter usado um
      Heading 4 onde devia ser Heading 3).
    - Subheadings existentes (mesmo título normalizado, mesmo pai) →
      atualizam ordem_capitulo para refletir nova ordem de aparecimento.
      Mantêm `id_usuario_responsavel`, `status_capitulo`, `id`, etc.
    - Subheadings que sumiram do DOCX → marcados ativo=False (não
      excluídos: preserva histórico e referências).

    Retorna dict com contadores: {'inseridos': N, 'atualizados': N,
    'desativados': N}.
    """
    resultado = ServicoNiveladorErros.executar_com_tratamento(
        _sincronizar_subcapitulos_interno,
        db_session, capitulo_pai, caminho_docx,
        relatorio_id=relatorio_id,
        capitulo_id=capitulo_id,
        etapa='sincronizacao_subcapitulos'
    )

    # Se o resultado é um dict de erro, retorna dict de erro
    if isinstance(resultado, dict) and not resultado.get('sucesso', True):
        return resultado

    return resultado


def _filhos_corpo_doc(doc) -> list:
    """Retorna lista dos elementos filhos diretos do `<w:body>`,
    excluindo `<w:sectPr>` final (configuração da seção)."""
    body = doc.element.body
    return [c for c in body if c.tag != f'{{{W_NS}}}sectPr']


def _extrair_capitulo_como_docx_interno(
    caminho_master: str,
    capitulo,
    *,
    incluir_heading: bool = True,
) -> Optional[bytes]:
    """Implementação interna de extrair_capitulo_como_docx."""
    doc = Document(caminho_master)
    rng = localizar_range_capitulo(doc, capitulo)
    if rng is None:
        return None
    inicio, fim = rng

    body = doc.element.body
    primeiro_manter = inicio if incluir_heading else inicio + 1
    ultimo_manter = fim  # inclusivo

    # Coletar referências a remover: tudo fora do range, exceto
    # sectPr (configuração da seção). Capturar ANTES para evitar
    # shift de índices durante o loop de remoção.
    a_remover = []
    for k, child in enumerate(list(body)):
        if child.tag == f'{{{W_NS}}}sectPr':
            continue
        if k < primeiro_manter or k > ultimo_manter:
            a_remover.append(child)
    for el in a_remover:
        body.remove(el)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def extrair_capitulo_como_docx(
    caminho_master: str,
    capitulo,
    *,
    incluir_heading: bool = True,
    relatorio_id: Optional[int] = None,
    capitulo_id: Optional[int] = None,
) -> Optional[bytes]:
    """Extrai o range do `capitulo` do DOCX em `caminho_master` e
    retorna um DOCX autônomo (em bytes) contendo apenas aquele
    capítulo.

    Útil para o endpoint `/api/capitulos/<id>/conteudo` (usado pelo
    eigenpal modal e pelos editores legados) — substitui o antigo
    `cap.conteudo_docx` agora que o DOCX em produção é a fonte
    única.

    Estratégia minimalista: clona o master, identifica o range
    `[inicio..fim]`, remove TODO o resto do body (preservando
    `<w:sectPr>` para que o DOCX final continue válido) e salva
    em memória.

    Retorna `None` se o capítulo não for localizado.
    """
    resultado = ServicoNiveladorErros.executar_com_tratamento(
        _extrair_capitulo_como_docx_interno,
        caminho_master, capitulo,
        incluir_heading=incluir_heading,
        relatorio_id=relatorio_id,
        capitulo_id=capitulo_id,
        etapa='extracao_capitulo_docx'
    )

    # Se o resultado é um dict de erro, retorna None
    if isinstance(resultado, dict) and not resultado.get('sucesso', True):
        return None

    return resultado


def _substituir_texto_heading(p_element, novo_texto: str) -> None:
    """Substitui o texto visível de um `<w:p>` heading preservando
    pPr (estilo, numeração) e os bookmarks anchor de cross-refs.

    Estratégia:
    - Mantém o primeiro `<w:r>` (cópia das propriedades de run) e
      substitui o texto do primeiro `<w:t>` por `novo_texto`.
    - Remove os demais `<w:r>` (que carregavam o restante do título
      antigo) sem mexer em `<w:bookmarkStart/>` `<w:bookmarkEnd/>`
      `<w:pPr>` (filhos diretos do `<w:p>` que NÃO são `<w:r>`).
    - Se o heading não tinha nenhum `<w:r>`, cria um novo.
    """
    runs = p_element.findall(f'{{{W_NS}}}r')
    if runs:
        primeiro = runs[0]
        # Localiza o primeiro <w:t> dentro do primeiro run
        t_primeiro = primeiro.find(f'{{{W_NS}}}t')
        if t_primeiro is None:
            t_primeiro = etree.SubElement(primeiro, f'{{{W_NS}}}t')
        t_primeiro.text = novo_texto
        # Limpa demais <w:t> dentro do primeiro run para não duplicar
        for t in primeiro.findall(f'{{{W_NS}}}t')[1:]:
            primeiro.remove(t)
        # Remove os runs subsequentes (carregavam o resto do título)
        for r in runs[1:]:
            p_element.remove(r)
    else:
        novo_run = etree.SubElement(p_element, f'{{{W_NS}}}r')
        novo_t = etree.SubElement(novo_run, f'{{{W_NS}}}t')
        novo_t.text = novo_texto


def _atualizar_titulo_capitulo_interno(
    caminho_master: str,
    capitulo,
    novo_titulo: str,
) -> bool:
    """Implementação interna de atualizar_titulo_capitulo."""
    if not novo_titulo or not novo_titulo.strip():
        return False
    master = Document(caminho_master)
    rng = localizar_range_capitulo(master, capitulo)
    if rng is None:
        return False
    inicio, _ = rng
    body = master.element.body
    p_heading = body[inicio]
    if p_heading.tag != f'{{{W_NS}}}p':
        return False
    _substituir_texto_heading(p_heading, novo_titulo.strip())
    master.save(caminho_master)
    return True


def atualizar_titulo_capitulo(
    caminho_master: str,
    capitulo,
    novo_titulo: str,
    relatorio_id: Optional[int] = None,
    capitulo_id: Optional[int] = None,
) -> bool:
    """Atualiza apenas o texto do parágrafo de heading do capítulo
    no DOCX em produção, preservando estilo, numeração automática,
    bookmarks de cross-reference e demais marcadores.

    Retorna True se o heading foi localizado e atualizado.
    """
    resultado = ServicoNiveladorErros.executar_com_tratamento(
        _atualizar_titulo_capitulo_interno,
        caminho_master, capitulo, novo_titulo,
        relatorio_id=relatorio_id,
        capitulo_id=capitulo_id,
        etapa='atualizacao_titulo_capitulo'
    )

    # Se o resultado é um dict de erro, retorna False
    if isinstance(resultado, dict) and not resultado.get('sucesso', True):
        return False

    return resultado


def _substituir_capitulo_interno(
    caminho_master: str,
    capitulo,
    caminho_autor: str,
    *,
    preservar_heading: bool = True,
) -> dict | bool:
    """Implementação interna de substituir_capitulo.

    Returns:
        dict ou bool: Retorna dict de erro com 'sugestoes' se capítulo não encontrado,
                      ou True se bem-sucedido, ou False se erro na substituição.
    """
    master = Document(caminho_master)
    autor = Document(caminho_autor)

    # Usar localização robusta com cascata de estratégias
    resultado_localizacao = localizar_range_capitulo_robusto(master, capitulo)
    if not resultado_localizacao['encontrado']:
        # Retornar dict de erro com alternativas sugeridas
        return {
            'sucesso': False,
            'erro': f"Capítulo '{capitulo.titulo_capitulo}' não localizado no documento",
            'sugestoes': resultado_localizacao.get('alternativas', []),
            'diagnostico': resultado_localizacao.get('diagnostico', 'Capítulo não encontrado')
        }

    inicio = resultado_localizacao['inicio']
    fim = resultado_localizacao['fim']

    body = master.element.body
    # NOTA sobre lxml: `id(elemento)` NÃO é estável entre iterações
    # (cada acesso retorna um proxy Python novo). Por isso usamos
    # contagem de filhos (estável) para identificar os elementos
    # apendados pelo Composer.
    n_pre_append = len(body)

    composer = Composer(master)
    composer.append(autor)

    n_pos_append = len(body)
    # Os apendados ficam no FINAL do body, normalmente antes do
    # `<w:sectPr>` (que continua sendo o último elemento da seção
    # original) — docxcompose insere `body.insert(len(body)-1, …)`.
    # Por segurança, filtramos elementos sectPr da janela apendada.
    inicio_apendados = n_pre_append
    fim_apendados = n_pos_append  # exclusivo
    novos = []
    for k in range(inicio_apendados, fim_apendados):
        el = body[k]
        if el.tag == f'{{{W_NS}}}sectPr':
            continue
        novos.append(el)

    # Remover o range do capítulo. inicio/fim foram capturados ANTES
    # do append; como o append acrescenta SOMENTE no fim do body,
    # esses índices continuam válidos para os elementos originais.
    if preservar_heading:
        primeiro_remover = inicio + 1
    else:
        primeiro_remover = inicio
    fim_remover = fim + 1  # exclusivo

    # Coletar referências (lxml.Element) aos elementos a remover
    # ANTES de remover qualquer um — assim a iteração não shifta
    # índices durante o loop.
    a_remover = [body[k] for k in range(primeiro_remover, fim_remover)]
    for elem in a_remover:
        body.remove(elem)

    # Posição final de inserção (após remoção):
    pos_insercao = (inicio + 1) if preservar_heading else inicio

    # Mover os apendados (que estão no fim do body) para a posição
    # correta. Removemos do fim e re-inserimos preservando ordem.
    for elem in novos:
        body.remove(elem)
    for offset, elem in enumerate(novos):
        body.insert(pos_insercao + offset, elem)

    master.save(caminho_master)
    return True


def substituir_capitulo(
    caminho_master: str,
    capitulo,
    caminho_autor: str,
    *,
    preservar_heading: bool = True,
    relatorio_id: Optional[int] = None,
    capitulo_id: Optional[int] = None,
) -> dict | bool:
    """Substitui o conteúdo do `capitulo` no DOCX em
    `caminho_master` pelo conteúdo do DOCX do autor em
    `caminho_autor`. Salva o resultado em `caminho_master`.

    Quando `preservar_heading=True` (padrão), o parágrafo de heading
    do capítulo no master é MANTIDO; só os elementos seguintes (até
    antes do próximo heading <= nível) são substituídos.

    Estratégia:
    1. `Composer(master).append(autor_doc)` → docxcompose copia
       todos os elementos do autor para o final do master,
       remapeando rIds de imagens, estilos numerados, etc. Isso
       é o que garante que figuras/numerações cheguem íntegras.
    2. Identificamos os elementos APENDADOS (novos) localizando
       a posição que era o último elemento antes do append.
    3. Removemos o range do capítulo (preservando heading conforme
       parâmetro).
    4. Movemos os elementos apendados para a posição correta.
    5. Salvamos.

    Retorna:
        - True em sucesso
        - dict com 'sucesso': False, 'sugestoes': [...] se capítulo não localizado
        - dict com 'sucesso': False se outro erro
    """
    resultado = ServicoNiveladorErros.executar_com_tratamento(
        _substituir_capitulo_interno,
        caminho_master, capitulo, caminho_autor,
        preservar_heading=preservar_heading,
        relatorio_id=relatorio_id,
        capitulo_id=capitulo_id,
        etapa='substituicao_capitulo'
    )

    # Se o resultado é um dict de erro, retorna o dict com sugestões
    if isinstance(resultado, dict) and not resultado.get('sucesso', True):
        return resultado

    # Se é True (sucesso)
    if resultado is True:
        return True

    # Se é False ou outro tipo, retorna False
    return False


def remover_capitulo_do_docx(caminho_docx: str, capitulo) -> bool:
    """Remove o range de um capítulo do DOCX de produção."""
    doc = Document(caminho_docx)
    rng = localizar_range_capitulo(doc, capitulo)
    if rng is None:
        return False
    inicio, fim = rng
    body = doc.element.body
    for idx in range(fim, inicio - 1, -1):
        body.remove(body[idx])
    doc.save(caminho_docx)
    return True


def _match_fuzzy(
    doc,
    capitulo,
    max_distancia_edicao: int = 2
) -> dict:
    """Match por fuzzy (distância de edição).

    Procura heading com distância de Levenshtein ≤ max_distancia_edicao
    em relação ao título normalizado do capítulo.

    Args:
        doc: Documento python-docx
        capitulo: Objeto CapituloDocumento com título a buscar
        max_distancia_edicao: Distância máxima de edição permitida (default: 2)

    Returns:
        Dict com:
            - encontrado: bool
            - indice: int ou None (índice do heading encontrado)
            - confianca: float (0.5-0.9 baseado em ratio)
            - titulo_encontrado: str ou None
            - diagnostico: str
            - alternativas: list de dicts com melhores matches
    """
    # Normalizar título do capítulo
    titulo_alvo = _normalizar(capitulo.titulo_capitulo)
    if not titulo_alvo:
        return {
            'encontrado': False,
            'indice': None,
            'confianca': 0.0,
            'titulo_encontrado': None,
            'diagnostico': 'Título do capítulo vazio após normalização',
            'alternativas': []
        }

    # Coletar todos os headings do documento
    body = doc.element.body
    headings = []

    for i, child in enumerate(body):
        if child.tag != f'{{{W_NS}}}p':
            continue
        nivel = _eh_paragrafo_heading(child)
        if nivel is None:
            continue
        texto = _texto_paragrafo(child)
        if not texto:
            continue
        texto_norm = _normalizar(texto)
        headings.append({
            'indice': i,
            'texto_original': texto,
            'texto_normalizado': texto_norm,
            'nivel': nivel
        })

    if not headings:
        return {
            'encontrado': False,
            'indice': None,
            'confianca': 0.0,
            'titulo_encontrado': None,
            'diagnostico': 'Nenhum heading encontrado no documento',
            'alternativas': []
        }

    # Calcular similaridade para cada heading
    matches = []
    for heading in headings:
        texto_norm = heading['texto_normalizado']

        # Usar SequenceMatcher para calcular similaridade
        ratio = difflib.SequenceMatcher(None, titulo_alvo, texto_norm).ratio()

        # Calcular distância de edição aproximada
        # Para strings curtas, podemos estimar: distância ≈ (1 - ratio) * max(len(a), len(b))
        max_len = max(len(titulo_alvo), len(texto_norm))
        distancia_estimada = int((1 - ratio) * max_len)

        if distancia_estimada <= max_distancia_edicao:
            # Mapear ratio para confiança 0.5-0.9
            # ratio 0.8-1.0 → confiança 0.7-0.9
            # ratio 0.6-0.8 → confiança 0.5-0.7
            if ratio >= 0.8:
                confianca = 0.7 + (ratio - 0.8) * 1.0  # 0.7-0.9
            elif ratio >= 0.6:
                confianca = 0.5 + (ratio - 0.6) * 1.0  # 0.5-0.7
            else:
                confianca = 0.5

            # Limitar confiança entre 0.5 e 0.9
            confianca = max(0.5, min(0.9, confianca))

            matches.append({
                'indice': heading['indice'],
                'texto_original': heading['texto_original'],
                'texto_normalizado': texto_norm,
                'nivel': heading['nivel'],
                'ratio': ratio,
                'distancia_estimada': distancia_estimada,
                'confianca': confianca
            })

    # Ordenar por confiança (maior primeiro)
    matches.sort(key=lambda x: x['confianca'], reverse=True)

    if matches:
        melhor = matches[0]
        # Preparar alternativas (melhores 3 matches)
        alternativas = []
        for i, match in enumerate(matches[:3]):
            alternativas.append({
                'posicao': i + 1,
                'titulo': match['texto_original'],
                'titulo_normalizado': match['texto_normalizado'],
                'confianca': match['confianca'],
                'ratio': match['ratio'],
                'distancia_estimada': match['distancia_estimada'],
                'nivel': match['nivel']
            })

        return {
            'encontrado': True,
            'indice': melhor['indice'],
            'confianca': melhor['confianca'],
            'titulo_encontrado': melhor['texto_original'],
            'diagnostico': f'Match fuzzy encontrado (distância={melhor["distancia_estimada"]}, ratio={melhor["ratio"]:.2f})',
            'alternativas': alternativas
        }
    else:
        # Nenhum match dentro da distância máxima
        # Retornar os 3 melhores matches mesmo fora do limite para sugestões
        todas_matches = []
        for heading in headings:
            texto_norm = heading['texto_normalizado']
            ratio = difflib.SequenceMatcher(None, titulo_alvo, texto_norm).ratio()
            max_len = max(len(titulo_alvo), len(texto_norm))
            distancia_estimada = int((1 - ratio) * max_len)

            todas_matches.append({
                'indice': heading['indice'],
                'texto_original': heading['texto_original'],
                'texto_normalizado': texto_norm,
                'nivel': heading['nivel'],
                'ratio': ratio,
                'distancia_estimada': distancia_estimada,
                'confianca': ratio  # Usar ratio como confiança para ordenação
            })

        todas_matches.sort(key=lambda x: x['confianca'], reverse=True)
        alternativas = []
        for i, match in enumerate(todas_matches[:3]):
            alternativas.append({
                'posicao': i + 1,
                'titulo': match['texto_original'],
                'titulo_normalizado': match['texto_normalizado'],
                'confianca': match['ratio'],  # Ratio como confiança
                'ratio': match['ratio'],
                'distancia_estimada': match['distancia_estimada'],
                'nivel': match['nivel']
            })

        return {
            'encontrado': False,
            'indice': None,
            'confianca': 0.0,
            'titulo_encontrado': None,
            'diagnostico': f'Nenhum match fuzzy dentro da distância máxima ({max_distancia_edicao})',
            'alternativas': alternativas
        }


def _match_exato(
    doc,
    capitulo,
    headings_cache: dict = None
) -> dict:
    """Match por casamento exato de estilo + título + nível.

    Args:
        doc: Documento python-docx
        capitulo: Objeto CapituloDocumento com título a buscar
        headings_cache: Cache opcional de headings para performance

    Returns:
        Dict com:
            - encontrado: bool
            - indice: int ou None
            - confianca: float (0.95 se exato)
            - titulo_encontrado: str ou None
            - diagnostico: str
            - alternativas: list vazia (não aplicável para match exato)
    """
    # Normalizar título do capítulo
    titulo_alvo = _normalizar(capitulo.titulo_capitulo)
    indice_alvo = (capitulo.indice_capitulo or '').strip()
    nivel_alvo = capitulo.nivel_capitulo or 1

    if not titulo_alvo and not indice_alvo:
        return {
            'encontrado': False,
            'indice': None,
            'confianca': 0.0,
            'titulo_encontrado': None,
            'diagnostico': 'Título e índice do capítulo vazios',
            'alternativas': []
        }

    body = doc.element.body

    # Usar cache se fornecido, senão buscar headings
    if headings_cache is not None and 'headings' in headings_cache:
        headings = headings_cache['headings']
    else:
        headings = []
        for i, child in enumerate(body):
            if child.tag != f'{{{W_NS}}}p':
                continue
            nivel = _eh_paragrafo_heading(child)
            if nivel is None:
                continue
            texto = _texto_paragrafo(child)
            if not texto:
                continue
            texto_norm = _normalizar(texto)
            headings.append({
                'indice': i,
                'texto_original': texto,
                'texto_normalizado': texto_norm,
                'nivel': nivel
            })

        # Atualizar cache se fornecido
        if headings_cache is not None:
            headings_cache['headings'] = headings

    # Buscar match exato
    for heading in headings:
        texto_norm = heading['texto_normalizado']

        # Verificar casamento por título normalizado
        casou_titulo = bool(titulo_alvo) and texto_norm == titulo_alvo

        # Verificar casamento por índice (prefixo do texto original)
        casou_indice = False
        if indice_alvo and heading['texto_original']:
            # Verificar se texto original começa com índice + separador
            texto_original = heading['texto_original'].lstrip()
            if texto_original.startswith(indice_alvo):
                # Verificar se após o índice tem separador (espaço, ponto, etc.)
                if len(texto_original) > len(indice_alvo):
                    char_apos = texto_original[len(indice_alvo)]
                    if char_apos in ' .)':
                        casou_indice = True

        if casou_titulo or casou_indice:
            # Verificar se nível corresponde (opcional, mas aumenta confiança)
            nivel_corresponde = heading['nivel'] == nivel_alvo
            confianca = 0.95 if nivel_corresponde else 0.9

            diagnostico_parts = []
            if casou_titulo:
                diagnostico_parts.append('título normalizado')
            if casou_indice:
                diagnostico_parts.append('índice')
            if nivel_corresponde:
                diagnostico_parts.append('nível correspondente')

            diagnostico = f'Match exato por {", ".join(diagnostico_parts)}'

            return {
                'encontrado': True,
                'indice': heading['indice'],
                'confianca': confianca,
                'titulo_encontrado': heading['texto_original'],
                'diagnostico': diagnostico,
                'alternativas': []
            }

    # Nenhum match exato encontrado
    return {
        'encontrado': False,
        'indice': None,
        'confianca': 0.0,
        'titulo_encontrado': None,
        'diagnostico': 'Nenhum match exato encontrado',
        'alternativas': []
    }


def _match_contexto(
    doc,
    capitulo,
    indice_esperado: int = None
) -> dict:
    """Match por contexto: índice + tipo + classificação.

    Usa número do capítulo ou classificação para inferir posição
    quando título não é encontrado.

    Args:
        doc: Documento python-docx
        capitulo: Objeto CapituloDocumento
        indice_esperado: Índice esperado do capítulo (opcional)

    Returns:
        Dict com:
            - encontrado: bool
            - indice: int ou None
            - confianca: float (0.6-0.8)
            - titulo_encontrado: str ou None
            - diagnostico: str
            - alternativas: list vazia
    """
    # Extrair número do capítulo para match por contexto
    # Prioridade: 1. parâmetro indice_esperado, 2. propriedade numero_capitulo_esperado do modelo
    numero_capitulo = None

    # 1. Usar parâmetro se fornecido (sobrescreve qualquer valor do modelo)
    if indice_esperado is not None:
        numero_capitulo = indice_esperado
    # 2. Usar propriedade numero_capitulo_esperado do modelo
    else:
        numero_capitulo = getattr(capitulo, 'numero_capitulo_esperado', None)

    # Obter classificação e tipo do capítulo
    classificacao = getattr(capitulo, 'classificacao', None)
    tipo_elemento = getattr(capitulo, 'tipo_elemento', None)

    if not numero_capitulo and not classificacao and not tipo_elemento:
        return {
            'encontrado': False,
            'indice': None,
            'confianca': 0.0,
            'titulo_encontrado': None,
            'diagnostico': 'Sem número de capítulo, classificação ou tipo para match por contexto',
            'alternativas': []
        }

    # Coletar todos os headings do documento
    body = doc.element.body
    headings = []

    for i, child in enumerate(body):
        if child.tag != f'{{{W_NS}}}p':
            continue
        nivel = _eh_paragrafo_heading(child)
        if nivel is None:
            continue
        texto = _texto_paragrafo(child)
        if not texto:
            continue
        texto_norm = _normalizar(texto)
        headings.append({
            'indice': i,
            'texto_original': texto,
            'texto_normalizado': texto_norm,
            'nivel': nivel
        })

    if not headings:
        return {
            'encontrado': False,
            'indice': None,
            'confianca': 0.0,
            'titulo_encontrado': None,
            'diagnostico': 'Nenhum heading encontrado no documento',
            'alternativas': []
        }

    # Tentar match por número de capítulo
    if numero_capitulo:
        # Procurar heading que comece com o número do capítulo
        for heading in headings:
            texto_original = heading['texto_original']
            # Verificar se texto começa com número do capítulo + separador
            match = re.match(rf'^\s*{numero_capitulo}[\.\)\s]', texto_original)
            if match:
                return {
                    'encontrado': True,
                    'indice': heading['indice'],
                    'confianca': 0.7,
                    'titulo_encontrado': heading['texto_original'],
                    'diagnostico': f'Match por contexto: número do capítulo ({numero_capitulo})',
                    'alternativas': []
                }

    # Tentar match por classificação (ex: "ANEXO", "APÊNDICE")
    if classificacao:
        classificacao_lower = classificacao.lower()

        # Mapear classificações para padrões de busca
        padroes_busca = {
            'anexo': r'ANEXO\s*[A-Z]?',
            'apendice': r'AP[EÊ]NDICE\s*[A-Z]?',
            'pre_textual': r'(SUM[ÁA]RIO|RESUMO|ABSTRACT|LISTA)',
            'pos_textual': r'(REFER[EÊ]NCIAS|BIBLIOGRAFIA)',
            'textual': r'^\s*\d+[\.\)\s]'  # Capítulos textuais começam com número
        }

        if classificacao_lower in padroes_busca:
            padrao = padroes_busca[classificacao_lower]
            for heading in headings:
                texto_original = heading['texto_original'].upper()
                if re.search(padrao, texto_original, re.IGNORECASE):
                    return {
                        'encontrado': True,
                        'indice': heading['indice'],
                        'confianca': 0.6,
                        'titulo_encontrado': heading['texto_original'],
                        'diagnostico': f'Match por contexto: classificação ({classificacao})',
                        'alternativas': []
                    }

    # Tentar match por tipo_elemento (fallback quando não há classificação)
    if tipo_elemento and not classificacao:
        tipo_lower = tipo_elemento.lower()

        # Mapear tipos para padrões de busca
        padroes_tipo = {
            'pre_textual': r'(SUM[ÁA]RIO|RESUMO|ABSTRACT|LISTA)',
            'pos_textual': r'(REFER[EÊ]NCIAS|BIBLIOGRAFIA)',
            'textual': r'^\s*\d+[\.\)\s]'  # Capítulos textuais começam com número
        }

        if tipo_lower in padroes_tipo:
            padrao = padroes_tipo[tipo_lower]
            for heading in headings:
                texto_original = heading['texto_original'].upper()
                if re.search(padrao, texto_original, re.IGNORECASE):
                    return {
                        'encontrado': True,
                        'indice': heading['indice'],
                        'confianca': 0.65,  # Ligeiramente maior que classificação por ser mais genérico
                        'titulo_encontrado': heading['texto_original'],
                        'diagnostico': f'Match por contexto: tipo ({tipo_elemento})',
                        'alternativas': []
                    }

    # Nenhum match por contexto encontrado
    return {
        'encontrado': False,
        'indice': None,
        'confianca': 0.0,
        'titulo_encontrado': None,
        'diagnostico': 'Nenhum match por contexto encontrado',
        'alternativas': []
    }


def localizar_range_capitulo_robusto(
    doc,
    capitulo,
    estrategia: str = 'multi_niveis'
) -> dict:
    """Localiza range com múltiplas estratégias e contexto.

    Implementa cascata de estratégias: exato → fuzzy → contexto.

    Args:
        doc: Documento python-docx
        capitulo: Objeto CapituloDocumento
        estrategia: Estratégia a usar ('multi_niveis', 'exato', 'fuzzy', 'contexto')

    Returns:
        Dict com:
            - encontrado: bool
            - inicio: int ou None (índice do heading)
            - fim: int ou None (índice do último elemento)
            - secao_inicio: int ou None
            - secao_fim: int ou None
            - titulo_encontrado: str ou None
            - confianca: float (0.0-1.0)
            - estrategia_usada: str
            - diagnostico: str
            - alternativas: list de dicts com melhores matches
    """
    # Cache de headings para performance (usado por múltiplas estratégias)
    headings_cache = {}

    # Resultado padrão (não encontrado)
    resultado_padrao = {
        'encontrado': False,
        'inicio': None,
        'fim': None,
        'secao_inicio': None,
        'secao_fim': None,
        'titulo_encontrado': None,
        'confianca': 0.0,
        'estrategia_usada': 'nenhuma',
        'diagnostico': 'Capítulo não localizado',
        'alternativas': []
    }

    # Definir ordem das estratégias baseado no parâmetro
    if estrategia == 'exato':
        estrategias = ['exato']
    elif estrategia == 'fuzzy':
        estrategias = ['fuzzy']
    elif estrategia == 'contexto':
        estrategias = ['contexto']
    else:  # 'multi_niveis' (padrão)
        estrategias = ['exato', 'fuzzy', 'contexto']

    # Executar estratégias em cascata
    for estrategia_nome in estrategias:
        if estrategia_nome == 'exato':
            match_result = _match_exato(doc, capitulo, headings_cache)
        elif estrategia_nome == 'fuzzy':
            match_result = _match_fuzzy(doc, capitulo)
        elif estrategia_nome == 'contexto':
            match_result = _match_contexto(doc, capitulo)
        else:
            continue

        if match_result['encontrado']:
            # Encontrou! Calcular range completo
            indice_inicio = match_result['indice']

            # Determinar nível do heading encontrado
            nivel_inicio = None
            body = doc.element.body
            if indice_inicio is not None and indice_inicio < len(body):
                child = body[indice_inicio]
                if child.tag == f'{{{W_NS}}}p':
                    nivel_inicio = _eh_paragrafo_heading(child)

            # Usar nível do capítulo como fallback
            if nivel_inicio is None:
                nivel_inicio = capitulo.nivel_capitulo or 1

            # Calcular range respeitando seção
            range_result = _calcular_range_respeitando_secao(
                doc, indice_inicio, nivel_inicio
            )

            # Construir resultado completo
            return {
                'encontrado': True,
                'inicio': indice_inicio,
                'fim': range_result['fim'],
                'secao_inicio': range_result['secao_inicio'],
                'secao_fim': range_result['secao_fim'],
                'titulo_encontrado': match_result['titulo_encontrado'],
                'confianca': match_result['confianca'],
                'estrategia_usada': estrategia_nome,
                'diagnostico': (
                    f'{match_result["diagnostico"]}. '
                    f'{"Respeitou limite de seção." if range_result.get("encontrou_limite_secao", False) else ""}'
                ),
                'alternativas': match_result.get('alternativas', [])
            }

        # Se não encontrou nesta estratégia, adicionar alternativas ao resultado
        if match_result.get('alternativas'):
            resultado_padrao['alternativas'].extend(match_result['alternativas'])

    # Nenhuma estratégia encontrou o capítulo
    # Consolidar alternativas de todas as estratégias
    if resultado_padrao['alternativas']:
        # Remover duplicados por título normalizado
        alternativas_unicas = {}
        for alt in resultado_padrao['alternativas']:
            chave = alt.get('titulo_normalizado', alt.get('titulo', ''))
            if (
                chave not in alternativas_unicas
                or alt.get('confianca', 0) > alternativas_unicas[chave].get('confianca', 0)
            ):
                alternativas_unicas[chave] = alt

        resultado_padrao['alternativas'] = list(alternativas_unicas.values())
        # Ordenar por confiança
        resultado_padrao['alternativas'].sort(key=lambda x: x.get('confianca', 0), reverse=True)
        # Manter apenas as 3 melhores
        resultado_padrao['alternativas'] = resultado_padrao['alternativas'][:3]

        if resultado_padrao['alternativas']:
            titulos = [
                alt.get("titulo", "sem título")
                for alt in resultado_padrao["alternativas"][:2]
            ]
            resultado_padrao['diagnostico'] = (
                'Capítulo não localizado. Alternativas sugeridas: '
                f'{", ".join(titulos)}'
            )

    return resultado_padrao
