"""
Serviço de extração canônica.
Recebe um DOCX modelo (relatório finalizado) e extrai todos os parâmetros
de formatação, gerando um JSON canônico pronto para consumo.
"""

import json
import os
import re

from docx import Document


# Regex para detectar prefixo de numeração hierárquica no início do
# título de um capítulo no DOCX. Aceita formas como:
#   "1 Introdução", "1. Introdução", "5.4.6.1 Sistema",
#   "5.4.6.1. Sistema", "12.1 — Declaração", "5.4.1.1) Tema".
# Captura o índice (grupo 1) e o título limpo (grupo 2).
_INDICE_PREFIXO_RE = re.compile(
    r'^\s*(\d+(?:\.\d+)*)\s*[\.\)\-–—:]?\s+(.+?)\s*$'
)


def extrair_indice_e_titulo(texto):
    """Detecta prefixo numérico hierárquico no título de um heading.

    Retorna (indice, titulo_limpo) — `indice` é uma string como
    "5.4.6.1" ou None se não houver prefixo numérico.
    """
    if not texto:
        return None, texto
    m = _INDICE_PREFIXO_RE.match(texto)
    if not m:
        return None, texto.strip()
    return m.group(1), m.group(2).strip()


NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}


class ServicoExtracaoCanonica:
    """Extrai parâmetros canônicos de um DOCX modelo."""

    # Nomes dos arquivos canônicos produzidos
    ARQUIVO_FORMATACAO = 'canonico_formatacao.json'
    ARQUIVO_MACRO = 'canonico_estrutura_macro.json'
    ARQUIVO_CAPITULOS = 'canonico_capitulos.json'

    @classmethod
    def extrair(cls, caminho_docx, diretorio_saida):
        """
        Extrai parâmetros canônicos de um DOCX modelo e salva em 3 JSONs:
        - canonico_formatacao.json: estilos, seções, margens, numeração
        - canonico_macro.json: estrutura macro (capa, pré, textual, pós)
        - canonico_capitulos.json: árvore hierárquica de capítulos
        Retorna dict com as três chaves.
        """
        doc = Document(caminho_docx)

        secoes = cls._extrair_secoes(doc)

        formatacao = {
            'secoes': secoes,
            'estilos_paragrafo': cls._extrair_estilos_paragrafo(doc),
            'estilos_caractere': cls._extrair_estilos_caractere(doc),
            'estilos_tabela': cls._extrair_estilos_tabela(doc),
            'numeracao': cls._extrair_numeracao(doc),
            'legendas': cls._extrair_legendas(doc),
            'propriedades_documento': cls._extrair_props_documento(doc),
        }

        macro = cls._extrair_macro(doc, secoes)

        # Enriquecer macro com inventario detalhado da CAPA: shapes
        # (text boxes flutuantes), imagens flutuantes (anchors com
        # pic:pic), tabela de Folha de Rosto, tabela de Controle de
        # Versoes, SDT do Sumario nativo. Esses elementos nao sao
        # parágrafos comuns e por isso NAO seriam capturados pelo
        # `_extrair_macro` original (que so olha estilos de paragrafo).
        try:
            from app.services.servico_capa import (
                extrair_estrutura_capa,
            )
            capa_detalhada = extrair_estrutura_capa(doc)
            # Anexa ao primeiro bloco de tipo 'capa' (se existir) ou
            # cria um bloco dedicado.
            anexado = False
            for bloco in macro:
                if bloco.get('tipo') == 'capa':
                    bloco['detalhe'] = capa_detalhada
                    anexado = True
                    break
            if not anexado:
                macro.insert(0, {
                    'tipo': 'capa',
                    'inicio_paragrafo': 0,
                    'fim_paragrafo': capa_detalhada.get(
                        'indice_fim_capa', 0,
                    ),
                    'titulos': [],
                    'secoes_indices': [],
                    'detalhe': capa_detalhada,
                })
        except Exception as exc:  # pragma: no cover - extracao opcional
            # A capa detalhada e enriquecimento, nao bloqueia extracao
            for bloco in macro:
                if bloco.get('tipo') == 'capa':
                    bloco['detalhe_erro'] = str(exc)
                    break

        capitulos = cls._extrair_capitulos(doc)

        os.makedirs(diretorio_saida, exist_ok=True)

        caminho_fmt = os.path.join(
            diretorio_saida, cls.ARQUIVO_FORMATACAO
        )
        with open(caminho_fmt, 'w', encoding='utf-8') as f:
            json.dump(formatacao, f, ensure_ascii=False, indent=2)

        caminho_macro = os.path.join(
            diretorio_saida, cls.ARQUIVO_MACRO
        )
        with open(caminho_macro, 'w', encoding='utf-8') as f:
            json.dump(macro, f, ensure_ascii=False, indent=2)

        caminho_cap = os.path.join(
            diretorio_saida, cls.ARQUIVO_CAPITULOS
        )
        with open(caminho_cap, 'w', encoding='utf-8') as f:
            json.dump(capitulos, f, ensure_ascii=False, indent=2)

        return {
            'formatacao': formatacao,
            'macro': macro,
            'capitulos': capitulos,
        }

    @classmethod
    def _extrair_secoes(cls, doc):
        """
        Extrai propriedades de seção com intervalo de parágrafos.

        No DOCX, sectPr inline (dentro de pPr) marca o FIM de uma
        seção; a sectPr do body é a última seção. Retorna lista de
        dicts, cada um com inicio_paragrafo e fim_paragrafo.
        """
        secoes = []
        w_ns = NSMAP['w']
        paragraphs = doc.element.body.findall('w:p', NSMAP)
        inicio_sec = 0

        for idx, p_el in enumerate(paragraphs):
            ppr = p_el.find('w:pPr', NSMAP)
            sect_pr = ppr.find('w:sectPr', NSMAP) if ppr is not None else None
            if sect_pr is not None:
                secao = cls._parse_sect_pr(sect_pr, w_ns)
                secao['inicio_paragrafo'] = inicio_sec
                secao['fim_paragrafo'] = idx
                secoes.append(secao)
                inicio_sec = idx + 1

        # Última seção (body-level sectPr)
        body_sect = doc.element.body.find('w:sectPr', NSMAP)
        if body_sect is not None:
            secao = cls._parse_sect_pr(body_sect, w_ns)
            secao['inicio_paragrafo'] = inicio_sec
            secao['fim_paragrafo'] = len(paragraphs) - 1
            secoes.append(secao)

        return secoes

    @classmethod
    def _parse_sect_pr(cls, sect_pr, w_ns):
        """Extrai propriedades de um elemento sectPr."""
        secao = {}

        pg_sz = sect_pr.find('w:pgSz', NSMAP)
        if pg_sz is not None:
            w_val = int(pg_sz.get(f'{{{w_ns}}}w', '0'))
            secao['largura_pagina_mm'] = cls._emu_para_mm(w_val)
            h_val = int(pg_sz.get(f'{{{w_ns}}}h', '0'))
            secao['altura_pagina_mm'] = cls._emu_para_mm(h_val)
            orient = pg_sz.get(f'{{{w_ns}}}orient')
            secao['orientacao'] = orient if orient else 'retrato'

        pg_mar = sect_pr.find('w:pgMar', NSMAP)
        if pg_mar is not None:
            for lado in ['top', 'right', 'bottom', 'left',
                         'header', 'footer', 'gutter']:
                val = pg_mar.get(f'{{{w_ns}}}{lado}')
                if val:
                    secao[f'margem_{lado}_mm'] = (
                        cls._twips_para_mm(int(val))
                    )

        cols = sect_pr.find('w:cols', NSMAP)
        if cols is not None:
            secao['colunas'] = int(
                cols.get(f'{{{w_ns}}}num', '1')
            )
            sp = int(cols.get(f'{{{w_ns}}}space', '0'))
            secao['espaco_colunas_mm'] = cls._twips_para_mm(sp)

        tipo = sect_pr.find('w:type', NSMAP)
        if tipo is not None:
            secao['tipo_quebra'] = tipo.get(
                f'{{{w_ns}}}val', ''
            )

        # Formato de numeração de páginas — sinal forte para distinguir
        # pré-textual (lowerRoman/upperRoman) de textual (decimal)
        pg_num = sect_pr.find('w:pgNumType', NSMAP)
        if pg_num is not None:
            fmt = pg_num.get(f'{{{w_ns}}}fmt')
            if fmt:
                secao['formato_numero_pagina'] = fmt
            inicio = pg_num.get(f'{{{w_ns}}}start')
            if inicio:
                secao['inicio_numero_pagina'] = inicio

        return secao

    @staticmethod
    def _extrair_estilos_paragrafo(doc):
        """Extrai estilos de parágrafo definidos no documento."""
        estilos = []
        for style in doc.styles:
            if style.type != 1:  # WD_STYLE_TYPE.PARAGRAPH = 1
                continue
            est = {
                'nome': style.name,
                'style_id': style.style_id,
                'base_style': (
                    style.base_style.name
                    if style.base_style else None
                ),
                'builtin': style.builtin,
            }

            pf = style.paragraph_format
            if pf:
                est['alinhamento'] = (
                    str(pf.alignment)
                    if pf.alignment else None
                )
                est['espacamento_antes_pt'] = (
                    pf.space_before.pt if pf.space_before else None
                )
                est['espacamento_depois_pt'] = (
                    pf.space_after.pt if pf.space_after else None
                )
                est['entre_linhas'] = (
                    pf.line_spacing if pf.line_spacing else None
                )
                est['recuo_esquerda_cm'] = (
                    pf.left_indent.cm if pf.left_indent else None
                )
                est['recuo_direita_cm'] = (
                    pf.right_indent.cm if pf.right_indent else None
                )
                est['recuo_primeira_linha_cm'] = (
                    pf.first_line_indent.cm
                    if pf.first_line_indent else None
                )

            font = style.font
            if font:
                est['fonte_nome'] = font.name
                est['fonte_tamanho_pt'] = (
                    font.size.pt if font.size else None
                )
                est['negrito'] = font.bold
                est['italico'] = font.italic
                est['sublinhado'] = (
                    bool(font.underline)
                    if font.underline is not None else None
                )
                est['cor_rgb'] = (
                    str(font.color.rgb)
                    if font.color and font.color.rgb
                    else None
                )

            estilos.append(est)
        return estilos

    @staticmethod
    def _extrair_estilos_caractere(doc):
        """Extrai estilos de caractere definidos no documento."""
        estilos = []
        for style in doc.styles:
            if style.type != 2:  # WD_STYLE_TYPE.CHARACTER = 2
                continue
            est = {
                'nome': style.name,
                'style_id': style.style_id,
                'base_style': (
                    style.base_style.name
                    if style.base_style else None
                ),
            }
            font = style.font
            if font:
                est['fonte_nome'] = font.name
                est['fonte_tamanho_pt'] = (
                    font.size.pt if font.size else None
                )
                est['negrito'] = font.bold
                est['italico'] = font.italic
                est['cor_rgb'] = (
                    str(font.color.rgb)
                    if font.color and font.color.rgb
                    else None
                )
            estilos.append(est)
        return estilos

    @staticmethod
    def _extrair_estilos_tabela(doc):
        """Extrai estilos de tabela definidos no documento."""
        estilos = []
        for style in doc.styles:
            if style.type != 3:  # WD_STYLE_TYPE.TABLE = 3
                continue
            estilos.append({
                'nome': style.name,
                'style_id': style.style_id,
                'base_style': (
                    style.base_style.name
                    if style.base_style else None
                ),
            })
        return estilos

    @staticmethod
    def _extrair_numeracao(doc):
        """Extrai definições de numeração (listas, headings)."""
        numeracoes = []
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return numeracoes

        # pylint: disable=protected-access
        numbering_xml = numbering_part._element  # noqa: SLF001
        for abstract_num in numbering_xml.findall('.//w:abstractNum', NSMAP):
            abstract_id = abstract_num.get('{%s}abstractNumId' % NSMAP['w'])
            niveis = []
            for lvl in abstract_num.findall('w:lvl', NSMAP):
                nivel_info = {
                    'nivel': int(lvl.get('{%s}ilvl' % NSMAP['w'], '0')),
                }
                num_fmt = lvl.find('w:numFmt', NSMAP)
                if num_fmt is not None:
                    nivel_info['formato'] = num_fmt.get(
                        '{%s}val' % NSMAP['w'], ''
                    )

                lvl_text = lvl.find('w:lvlText', NSMAP)
                if lvl_text is not None:
                    nivel_info['texto_nivel'] = (
                        lvl_text.get(
                            '{%s}val' % NSMAP['w'], ''
                        )
                    )

                start = lvl.find('w:start', NSMAP)
                if start is not None:
                    nivel_info['inicio'] = int(
                        start.get(
                            '{%s}val' % NSMAP['w'], '1'
                        )
                    )

                p_style = lvl.find('w:pStyle', NSMAP)
                if p_style is not None:
                    nivel_info['estilo_vinculado'] = (
                        p_style.get(
                            '{%s}val' % NSMAP['w'], ''
                        )
                    )

                niveis.append(nivel_info)

            numeracoes.append({
                'abstract_num_id': abstract_id,
                'niveis': niveis,
            })
        return numeracoes

    @staticmethod
    def _extrair_legendas(doc):
        """
        Detecta padrões de legendas de figuras, tabelas e fontes
        no documento. Identifica:
        - estilos usados para legenda de figuras e tabelas
        - posição da legenda (acima / abaixo do objeto)
        - existência e estilo da indicação de fonte
        """
        ns = {
            'w': ('http://schemas.openxmlformats.org'
                  '/wordprocessingml/2006/main'),
        }
        paragrafos = doc.paragraphs
        total = len(paragrafos)

        fig_legendas = []
        tab_legendas = []
        fonte_legendas = []

        for i, p in enumerate(paragrafos):
            txt = p.text.strip()
            sn = p.style.name or ''
            txt_lower = txt.lower()

            # Legenda de figura (texto começa com "Figura")
            if txt_lower.startswith('figura') and len(txt) > 8:
                prev_el = getattr(  # type: ignore[misc]
                    paragrafos[i - 1], '_element', None
                ) if i > 0 else None
                has_img_prev = (
                    prev_el is not None
                    and prev_el.findall('.//w:drawing', ns)
                )
                next_el = getattr(  # type: ignore[misc]
                    paragrafos[i + 1], '_element', None
                ) if i < total - 1 else None
                has_img_next = (
                    next_el is not None
                    and next_el.findall('.//w:drawing', ns)
                )
                pos = 'acima'
                if has_img_prev:
                    pos = 'abaixo'
                elif has_img_next:
                    pos = 'acima'
                fig_legendas.append({
                    'indice_paragrafo': i,
                    'estilo': sn,
                    'posicao': pos,
                    'texto_exemplo': txt[:80],
                })

            # Legenda de tabela
            if txt_lower.startswith('tabela') and len(txt) > 8:
                tab_legendas.append({
                    'indice_paragrafo': i,
                    'estilo': sn,
                    'posicao': 'acima',
                    'texto_exemplo': txt[:80],
                })

            # Indicação de fonte
            if txt_lower.startswith('fonte:'):
                fonte_legendas.append({
                    'indice_paragrafo': i,
                    'estilo': sn,
                    'texto_exemplo': txt[:80],
                })

        # Consolidar padrão predominante
        def padrao(lst):
            """Extrai estilo e posição mais frequentes."""
            if not lst:
                return None
            estilos = {}
            posicoes = {}
            for item in lst:
                e = item.get('estilo', '')
                estilos[e] = estilos.get(e, 0) + 1
                p = item.get('posicao', '')
                if p:
                    posicoes[p] = posicoes.get(p, 0) + 1
            est_top = max(estilos, key=lambda k: estilos[k])
            pos_top = (max(posicoes, key=lambda k: posicoes[k])
                       if posicoes else None)
            return {
                'estilo_predominante': est_top,
                'posicao_predominante': pos_top,
                'total_ocorrencias': len(lst),
                'exemplos': [
                    x['texto_exemplo'] for x in lst[:3]
                ],
            }

        return {
            'figura': padrao(fig_legendas),
            'tabela': padrao(tab_legendas),
            'fonte': padrao(fonte_legendas),
        }

    @staticmethod
    def _extrair_props_documento(doc):
        """Extrai propriedades gerais do documento."""
        props = {}
        core = doc.core_properties
        if core:
            props['titulo'] = core.title or ''
            props['assunto'] = core.subject or ''
            props['autor'] = core.author or ''
            props['categoria'] = core.category or ''

        # Contar elementos para referência
        props['total_paragrafos'] = len(doc.paragraphs)
        props['total_tabelas'] = len(doc.tables)
        props['total_secoes'] = len(doc.sections)

        return props

    @classmethod
    def _extrair_macro(cls, doc, secoes=None):
        """
        Extrai a estrutura macro do documento:
        Capa, Pré-textuais, Textuais, Pós-textuais.

        Cada bloco inclui `secoes_indices` — lista de índices das
        seções (w:sectPr) que cobrem aquele intervalo de parágrafos,
        permitindo ao visualizador saber qual tamanho de página e
        margens aplicar a cada bloco.
        """
        pos_textuais_kw = [
            'referências', 'referencias', 'bibliography',
            'apêndice', 'apendice', 'appendix',
            'anexo', 'annex',
        ]
        pre_textuais_kw = [
            'sumário', 'sumario', 'table of contents',
            'lista de figuras', 'lista de tabelas',
            'lista de abreviaturas', 'lista de siglas',
            'resumo', 'abstract',
            'folha de rosto', 'folha de aprovação',
            'dedicatória', 'agradecimentos', 'epígrafe',
        ]

        if secoes is None:
            secoes = cls._extrair_secoes(doc)

        # Sinal por seções DOCX: quando uma seção transiciona de numeração
        # romana (lowerRoman/upperRoman) para decimal, é forte indicador de
        # que a parte textual começa no primeiro parágrafo dessa seção.
        indice_textual_por_secao = None
        formato_anterior = None
        for sec in secoes:
            fmt = (sec.get('formato_numero_pagina') or '').lower()
            if (formato_anterior in ('lowerroman', 'upperroman')
                    and fmt in ('decimal', '')):
                indice_textual_por_secao = sec.get('inicio_paragrafo')
                break
            if fmt:
                formato_anterior = fmt

        primeiro_heading1_encontrado = False
        indice_inicio_textual = None
        indice_inicio_pos = None

        # Títulos canônicos pré-textuais (match exato) usados para descartar
        # Heading 1 que na verdade rotulam Sumário, Resumo, Listas, etc.
        pre_textuais_titulos_exatos = cls._PRE_TEXTUAIS_TITULOS

        def _eh_titulo_pre_textual(texto_lower, style_lower):
            """Verifica se um parágrafo (ainda que com estilo Heading) deve
            ser tratado como pré-textual e portanto ignorado para fins
            de detecção do início do conteúdo textual.

            Tolerante a prefixos numericos automaticos do Word
            (ex.: '1 SUMÁRIO' → compara com 'sumário').
            """
            if 'toc' in style_lower:
                return True
            # Remove eventual prefixo numerico antes de comparar.
            _idx, titulo_limpo = extrair_indice_e_titulo(texto_lower)
            alvo = (titulo_limpo or texto_lower).strip().lower()
            if alvo in pre_textuais_titulos_exatos:
                return True
            for kw in pre_textuais_kw:
                if kw == alvo or alvo.startswith(kw + ' '):
                    return True
                # Compara tambem com texto bruto (caso indice nao seja
                # detectado mas o titulo comece com o keyword).
                if kw == texto_lower or texto_lower.startswith(kw + ' '):
                    return True
            return False

        for i, para in enumerate(doc.paragraphs):
            style_name = (para.style.name or '').lower()
            texto = para.text.strip().lower()

            if (style_name.startswith('heading 1') or
                    style_name == 'heading 1'):
                # Pula Heading 1 que rotula elemento pré-textual
                # (Sumário, Resumo, Lista de Figuras, etc.)
                if _eh_titulo_pre_textual(texto, style_name):
                    continue
                if not primeiro_heading1_encontrado:
                    primeiro_heading1_encontrado = True
                    indice_inicio_textual = i

            if primeiro_heading1_encontrado and indice_inicio_pos is None:
                if style_name.startswith('heading'):
                    for kw in pos_textuais_kw:
                        if kw in texto:
                            indice_inicio_pos = i
                            break

        # Refinar com sinal de seções DOCX:
        # - Se nenhum Heading 1 textual foi encontrado, usar a seção com
        #   numeração decimal como início textual.
        # - Se o Heading 1 detectado vem ANTES da transição decimal, ainda
        #   pode ser ruído; preferir o índice da seção, que é mais confiável.
        if indice_textual_por_secao is not None:
            if indice_inicio_textual is None:
                indice_inicio_textual = indice_textual_por_secao
            elif indice_inicio_textual < indice_textual_por_secao:
                indice_inicio_textual = indice_textual_por_secao

        macro = []
        for i, para in enumerate(doc.paragraphs):
            style_name = (para.style.name or '').lower()
            texto = para.text.strip()

            if indice_inicio_pos and i >= indice_inicio_pos:
                tipo = 'pos_textual'
            elif (indice_inicio_textual
                  and i >= indice_inicio_textual):
                tipo = 'textual'
            elif (indice_inicio_textual
                  and i < indice_inicio_textual):
                texto_lower = texto.lower()
                is_pre = any(
                    kw in texto_lower for kw in pre_textuais_kw
                )
                is_toc_style = 'toc' in style_name
                if is_pre or is_toc_style:
                    tipo = 'pre_textual'
                elif i < 3:
                    tipo = 'capa'
                else:
                    tipo = 'pre_textual'
            else:
                tipo = 'capa'

            if not macro or macro[-1]['tipo'] != tipo:
                macro.append({
                    'tipo': tipo,
                    'inicio_paragrafo': i,
                    'fim_paragrafo': i,
                    'titulos': [],
                })
            else:
                macro[-1]['fim_paragrafo'] = i

            if style_name.startswith('heading') and texto:
                macro[-1]['titulos'].append(texto)
            elif tipo == 'pre_textual' and texto:
                texto_lower = texto.lower()
                if any(kw in texto_lower for kw in pre_textuais_kw):
                    macro[-1]['titulos'].append(texto)

        # Consolidar: um registro por tipo
        resultado = {}
        for bloco in macro:
            t = bloco['tipo']
            if t not in resultado:
                resultado[t] = {
                    'tipo': t,
                    'inicio_paragrafo': bloco['inicio_paragrafo'],
                    'fim_paragrafo': bloco['fim_paragrafo'],
                    'titulos': list(bloco['titulos']),
                }
            else:
                resultado[t]['fim_paragrafo'] = bloco['fim_paragrafo']
                resultado[t]['titulos'].extend(bloco['titulos'])

        # Vincular seções aos blocos macro
        for bloco in resultado.values():
            bi = bloco['inicio_paragrafo']
            bf = bloco['fim_paragrafo']
            indices = []
            for si, sec in enumerate(secoes):
                if sec['fim_paragrafo'] >= bi and \
                   sec['inicio_paragrafo'] <= bf:
                    indices.append(si)
            bloco['secoes_indices'] = indices

        ordem = ['capa', 'pre_textual', 'textual', 'pos_textual']
        return [resultado[t] for t in ordem if t in resultado]

    # Pré-textuais AUTO-GERADOS: NUNCA viram capítulos no banco.
    # São produzidos a partir do conteúdo dos demais capítulos
    # (Sumário lista os capítulos; Lista de Figuras lista as figuras
    # dos capítulos; etc.). A camada de geração final do DOCX é que
    # monta esses elementos. Tratá-los como capítulo no banco
    # geraria recursão de conteúdo (capítulo "Sumário" listaria a si
    # mesmo) e duplicação visual.
    _PRE_TEXTUAIS_AUTO_GERADOS = {
        'sumário', 'sumario', 'table of contents',
        'lista de figuras', 'lista de tabelas',
        'lista de abreviaturas', 'lista de siglas',
        'lista de quadros', 'lista de gráficos',
        'lista de símbolos',
        'folha de rosto', 'folha de aprovação',
        'capa',
    }
    # Pré-textuais AUTORAIS: o autor escreve conteúdo livre dentro
    # deles. Esses sim viram capítulos (tipo='pre_textual').
    _PRE_TEXTUAIS_AUTORAIS = {
        'resumo', 'abstract',
        'dedicatória', 'dedicatoria',
        'agradecimentos',
        'epígrafe', 'epigrafe',
    }
    # União para verificação rápida (compatibilidade interna).
    _PRE_TEXTUAIS_TITULOS = (
        _PRE_TEXTUAIS_AUTO_GERADOS | _PRE_TEXTUAIS_AUTORAIS
    )
    _POS_TEXTUAIS_PREFIXOS = (
        'referências', 'referencias', 'bibliography',
        'apêndice', 'apendice', 'appendix',
        'anexo', 'annex', 'glossário', 'glossario', 'índice', 'indice',
    )

    @classmethod
    def _extrair_capitulos(cls, doc):
        """
        Extrai a árvore hierárquica de capítulos.

        Regras (para evitar duplicação e ruído):
        - Só considera parágrafos com estilo Heading N (1..9) ou
          parágrafos cujo texto coincide exatamente com um título
          canônico pré-textual (Sumário, Resumo, etc.) ou começa
          com prefixo pós-textual (Apêndice, Anexo, ...).
        - Deduplica por (título_normalizado, nível) dentro do mesmo
          pai — entradas repetidas no DOCX (TOC + Heading real) não
          geram capítulos duplicados.
        - Itens pré-textuais ficam no nível 1 com tipo 'pre_textual';
          pós-textuais no nível 1 com tipo 'pos_textual'; demais
          headings ficam como 'textual'.
        """
        # Estrutura macro para tipagem
        macro = cls._extrair_macro(doc)
        tipo_por_indice = {}
        for bloco in macro:
            for i in range(bloco['inicio_paragrafo'],
                           bloco['fim_paragrafo'] + 1):
                tipo_por_indice[i] = bloco['tipo']

        headings_flat = []
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name or ''
            texto = para.text.strip()
            if not texto or len(texto) < 2:
                continue
            tipo_elemento = tipo_por_indice.get(i, 'textual')
            texto_lower = texto.lower()
            style_lower = style_name.lower()

            # Estilos do Sumário gerados pelo Word (TOC 1..9) NUNCA são
            # capítulos — são entradas do índice do próprio sumário.
            if style_lower.startswith('toc'):
                continue

            nivel = None
            incluir = False

            if style_name.startswith('Heading'):
                try:
                    nivel = int(style_name.replace('Heading ', '').strip())
                    incluir = True
                except ValueError:
                    nivel = None

            # Pre-compute o titulo SEM prefixo numerico (ex.: "1 SUMÁRIO"
            # -> "sumário") para casar com os sets de keywords mesmo
            # quando o Word ja injetou numeracao automatica no heading.
            _idx_pref, titulo_limpo = extrair_indice_e_titulo(texto)
            alvo_match = (
                (titulo_limpo or texto).strip().lower()
            )

            # No bloco pré-textual:
            #   - títulos AUTO-GERADOS (Sumário, Listas de Figuras/Tabelas,
            #     Capa, Folha de Rosto) NUNCA viram capítulo. Eles são
            #     produzidos pela camada de renderização final a partir
            #     do conteúdo dos próprios capítulos.
            #   - títulos AUTORAIS (Resumo, Dedicatória, Agradecimentos,
            #     Epígrafe) viram capítulo de tipo 'pre_textual'.
            #   - qualquer outro Heading no bloco pré-textual é ignorado
            #     (parágrafos espúrios marcados como heading não devem
            #     gerar capítulos fantasmas).
            if tipo_elemento == 'pre_textual':
                if (alvo_match in cls._PRE_TEXTUAIS_AUTO_GERADOS
                        or texto_lower in cls._PRE_TEXTUAIS_AUTO_GERADOS):
                    incluir = False
                elif (alvo_match in cls._PRE_TEXTUAIS_AUTORAIS
                        or texto_lower in cls._PRE_TEXTUAIS_AUTORAIS):
                    nivel = 1
                    incluir = True
                else:
                    incluir = False

            # Pós-textuais: apenas se prefixo conhecido (além de Heading).
            # Tambem tolera prefixo numerico antes do keyword
            # (ex.: "16 APÊNDICE 01" -> casa com 'apêndice').
            if (not incluir and tipo_elemento == 'pos_textual'
                    and any(
                        alvo_match.startswith(p)
                        or texto_lower.startswith(p)
                        for p in cls._POS_TEXTUAIS_PREFIXOS
                    )):
                nivel = 1
                incluir = True

            if not incluir or nivel is None:
                continue

            # Tentar extrair prefixo numérico do título no DOCX
            # ("5.4.6.1 Sistema SIGMA-PLI" -> indice="5.4.6.1",
            #  titulo="Sistema SIGMA-PLI"). Se houver, o nível
            # vira o número de níveis do índice (mais confiável que
            # o estilo Heading, que pode estar inconsistente no DOCX).
            indice_extraido, titulo_limpo = extrair_indice_e_titulo(texto)
            if indice_extraido:
                titulo_final = titulo_limpo
                nivel_pelo_indice = indice_extraido.count('.') + 1
                nivel_final = nivel_pelo_indice
            else:
                titulo_final = texto
                nivel_final = nivel

            headings_flat.append({
                'titulo': titulo_final,
                'indice': indice_extraido,
                'nivel': nivel_final,
                'estilo': style_name,
                'tipo_elemento': tipo_elemento,
                'indice_paragrafo': i,
                'filhos': [],
            })

        # Montar árvore hierárquica com deduplicação por (titulo, nível, pai)
        raiz = []
        pilha = []  # lista de (nivel, nó)

        def _existe_irmao(lista, titulo_norm, nivel):
            for irmao in lista:
                if (irmao['nivel'] == nivel
                        and irmao['titulo'].strip().lower()
                        == titulo_norm):
                    return True
            return False

        for item in headings_flat:
            nv = item['nivel']
            titulo_norm = item['titulo'].strip().lower()
            while pilha and pilha[-1][0] >= nv:
                pilha.pop()

            destino = pilha[-1][1]['filhos'] if pilha else raiz
            if _existe_irmao(destino, titulo_norm, nv):
                # Já registrado neste nível e pai — pular duplicata
                continue

            destino.append(item)
            pilha.append((nv, item))

        return raiz

    @staticmethod
    def _twips_para_mm(twips):
        """Converte twips para milímetros."""
        return round(twips * 25.4 / 1440, 2)

    @staticmethod
    def _emu_para_mm(emu_val):
        """Converte EMU (no contexto pgSz em twips) para mm."""
        return round(emu_val * 25.4 / 1440, 2)

    @staticmethod
    def validar_estrutura_canonica(caminho_docx, perfil=None):
        """Valida se o DOCX tem elementos canônicos esperados.

        Args:
            caminho_docx: Caminho para o arquivo DOCX.
            perfil: Perfil de formatação (opcional).

        Returns:
            Dict com chave 'problemas' (lista de strings) ou mensagem de erro.
        """
        del perfil
        if not os.path.exists(caminho_docx):
            return {'problemas': [f'Arquivo não encontrado: {caminho_docx}']}

        try:
            doc = Document(caminho_docx)
            problemas = []

            if len(doc.paragraphs) == 0:
                problemas.append('Documento sem parágrafos')

            if len(doc.sections) == 0:
                problemas.append('Documento sem seções')

            return {'problemas': problemas}
        except Exception as exc:
            return {'problemas': [f'Erro ao ler DOCX: {exc}']}
