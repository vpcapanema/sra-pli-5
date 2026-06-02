"""Serviço para extrair seções e quebras de página de documentos DOCX."""

import re
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from app.models.secao_docx import SecaoDOCX
from app.models.quebra_pagina import QuebraPagina
from app.models.capitulo_documento import CapituloDocumento


@dataclass
class ElementoDOCX:
    """Representa um elemento extraído do DOCX."""
    tipo: str  # 'secao', 'quebra_pagina', 'paragrafo', 'tabela', 'figura'
    posicao: int  # Posição no documento
    conteudo: str
    estilo: Optional[str] = None
    propriedades: Optional[Dict] = None


class ServicoExtracaoSecoes:
    """Serviço para extrair seções, quebras e elementos do DOCX."""
    
    @staticmethod
    def extrair_secoes_do_docx(docx_path: str, id_relatorio: int) -> List[SecaoDOCX]:
        """Extrai todas as seções (w:sectPr) de um documento DOCX.
        
        Nota: python-docx tem limitações na extração de seções.
        Em produção, pode ser necessário usar lxml diretamente.
        """
        try:
            from docx import Document
        except ImportError:
            # Fallback para quando python-docx não estiver disponível
            return ServicoExtracaoSecoes._extrair_secoes_fallback(docx_path, id_relatorio)
        
        try:
            doc = Document(docx_path)
            secoes = []
            
            # python-docx expõe seções através de document.sections
            for i, section in enumerate(doc.sections):
                secao = SecaoDOCX(
                    id_relatorio=id_relatorio,
                    ordem_secao=i,
                    tipo_secao=ServicoExtracaoSecoes._determinar_tipo_secao(section),
                    reiniciar_numero_pagina=section.start_type != 'continuous',
                    orientacao='landscape' if section.page_width > section.page_height else 'portrait',
                    colunas=ServicoExtracaoSecoes._contar_colunas(section),
                    margem_superior=section.top_margin.twips,
                    margem_inferior=section.bottom_margin.twips,
                    margem_esquerda=section.left_margin.twips,
                    margem_direita=section.right_margin.twips,
                    propriedades_raw=ServicoExtracaoSecoes._extrair_propriedades_raw(section)
                )
                secoes.append(secao)
            
            return secoes
            
        except Exception as e:
            # Fallback em caso de erro
            print(f"Erro ao extrair seções: {e}")
            return ServicoExtracaoSecoes._extrair_secoes_fallback(docx_path, id_relatorio)
    
    @staticmethod
    def _determinar_tipo_secao(section) -> str:
        """Determina o tipo de seção baseado nas propriedades."""
        # python-docx tem limitações aqui
        # Em produção, usar lxml para ler w:type diretamente
        if hasattr(section, 'start_type'):
            return section.start_type  # 'continuous', 'nextPage', etc.
        return 'nextPage'  # Default
    
    @staticmethod
    def _contar_colunas(section) -> int:
        """Conta o número de colunas na seção."""
        # python-docx não expõe colunas diretamente
        # Em produção, usar lxml
        return 1  # Default
    
    @staticmethod
    def _extrair_propriedades_raw(section) -> str:
        """Extrai propriedades OOXML brutas da seção."""
        # python-docx não expõe XML bruto facilmente
        # Em produção, usar lxml
        return "{}"
    
    @staticmethod
    def _extrair_secoes_fallback(docx_path: str, id_relatorio: int) -> List[SecaoDOCX]:
        """Fallback para extração de seções quando python-docx falha."""
        # Heurística simples: assumir que cada Heading 1 inicia nova seção
        # Em produção, implementar com lxml
        secoes = [
            SecaoDOCX(
                id_relatorio=id_relatorio,
                ordem_secao=0,
                tipo_secao='nextPage',  # Primeira seção começa em nova página
                reiniciar_numero_pagina=True,
                numero_pagina_inicial=1,
                estilo_numero_pagina='decimal'
            )
        ]
        
        # Adicionar seção para anexos/apêndices se existirem
        secoes.append(SecaoDOCX(
            id_relatorio=id_relatorio,
            ordem_secao=1,
            tipo_secao='continuous',
            reiniciar_numero_pagina=False,
            estilo_numero_pagina='decimal'
        ))
        
        return secoes
    
    @staticmethod
    def extrair_quebras_pagina_do_docx(docx_path: str, secoes: List[SecaoDOCX]) -> List[QuebraPagina]:
        """Extrai quebras de página do documento DOCX."""
        try:
            from docx import Document
        except ImportError:
            return []
        
        try:
            doc = Document(docx_path)
            quebras = []
            secao_atual = 0
            posicao_na_secao = 0
            
            for i, paragraph in enumerate(doc.paragraphs):
                # Verificar se parágrafo contém quebra de página
                if hasattr(paragraph, 'runs'):
                    for run in paragraph.runs:
                        # Verificar por texto que indica quebra
                        if run.text and '\\page' in run.text:
                            quebra = QuebraPagina(
                                id_secao=secoes[secao_atual].id_secao if secao_atual < len(secoes) else 1,
                                posicao_na_secao=posicao_na_secao,
                                tipo_posicao='paragrafo',
                                tipo_quebra='page',
                                forcar_nova_pagina=True,
                                contexto_anterior=paragraph.text[:100] if paragraph.text else None
                            )
                            quebras.append(quebra)
                
                posicao_na_secao += 1
                
                # Heurística: verificar se estilo indica nova seção
                if paragraph.style and paragraph.style.name in ['Heading 1', 'Título 1']:
                    # Possível início de nova seção
                    secao_atual = min(secao_atual + 1, len(secoes) - 1)
                    posicao_na_secao = 0
            
            return quebras
            
        except Exception as e:
            print(f"Erro ao extrair quebras de página: {e}")
            return []
    
    @staticmethod
    def mapear_capitulos_para_secoes(
        capitulos: List[CapituloDocumento], 
        secoes: List[SecaoDOCX]
    ) -> List[CapituloDocumento]:
        """Mapeia capítulos para as seções onde começam e terminam.
        
        Heurística: capítulos começam na seção atual e terminam na seção
        anterior ao próximo capítulo.
        """
        if not capitulos or not secoes:
            return capitulos
        
        # Ordenar capítulos por ordem_capitulo
        capitulos.sort(key=lambda x: x.ordem_capitulo)
        
        # Para cada capítulo, determinar seção de início
        for i, cap in enumerate(capitulos):
            # Heurística simples: capítulos textuais começam na seção 1
            # anexos/apêndices começam na seção 2
            if cap.tipo_elemento == 'textual' and not cap.classificacao:
                cap.id_secao_inicio = secoes[0].id_secao if secoes else 1
            elif cap.classificacao in ('anexo', 'apendice'):
                cap.id_secao_inicio = secoes[1].id_secao if len(secoes) > 1 else secoes[0].id_secao
            
            # Determinar seção de fim (se houver próximo capítulo)
            if i < len(capitulos) - 1:
                proximo_cap = capitulos[i + 1]
                # Capítulo termina na seção anterior ao próximo capítulo
                cap.id_secao_fim = proximo_cap.id_secao_inicio
            else:
                # Último capítulo: termina na última seção
                cap.id_secao_fim = secoes[-1].id_secao if secoes else cap.id_secao_inicio
        
        return capitulos
    
    @staticmethod
    def analisar_estrutura_documento(docx_path: str, id_relatorio: int) -> Dict:
        """Análise completa da estrutura do documento DOCX.
        
        Retorna:
            Dict com seções, quebras, capítulos e análise estrutural
        """
        # Extrair seções
        secoes = ServicoExtracaoSecoes.extrair_secoes_do_docx(docx_path, id_relatorio)
        
        # Extrair quebras de página
        quebras = ServicoExtracaoSecoes.extrair_quebras_pagina_do_docx(docx_path, secoes)
        
        # Análise estrutural
        analise = {
            'total_secoes': len(secoes),
            'total_quebras': len(quebras),
            'secoes_com_numero_diferente': sum(1 for s in secoes if s.tem_numero_pagina_diferente),
            'secoes_com_orientacao_diferente': sum(1 for s in secoes if s.orientacao != 'portrait'),
            'quebras_visiveis': sum(1 for q in quebras if q.e_quebra_visivel),
            'estrutura_secoes': [
                {
                    'ordem': s.ordem_secao,
                    'tipo': s.tipo_secao,
                    'orientacao': s.orientacao,
                    'colunas': s.colunas,
                    'reinicia_numero': s.reiniciar_numero_pagina
                }
                for s in secoes
            ]
        }
        
        return {
            'secoes': secoes,
            'quebras': quebras,
            'analise': analise
        }
    
    @staticmethod
    def validar_mapeamento_secoes_capitulos(
        secoes: List[SecaoDOCX], 
        capitulos: List[CapituloDocumento]
    ) -> List[str]:
        """Valida o mapeamento entre seções e capítulos."""
        erros = []
        
        # Verificar se todos os capítulos têm seção de início
        for cap in capitulos:
            if not cap.id_secao_inicio:
                erros.append(f"Capítulo '{cap.titulo_capitulo}' não tem seção de início")
            
            # Verificar se seção de início existe
            secao_inicio = next((s for s in secoes if s.id_secao == cap.id_secao_inicio), None)
            if not secao_inicio:
                erros.append(f"Capítulo '{cap.titulo_capitulo}' referencia seção inexistente: {cap.id_secao_inicio}")
            
            # Verificar se capítulos textuais não começam em seção de anexos
            if (cap.tipo_elemento == 'textual' and 
                secao_inicio and 
                secao_inicio.ordem_secao > 0 and  # Não é a primeira seção
                any(c.classificacao in ('anexo', 'apendice') for c in capitulos if c.id_secao_inicio == secao_inicio.id_secao)):
                erros.append(f"Capítulo textual '{cap.titulo_capitulo}' começa em seção de anexos/apêndices")
        
        return erros