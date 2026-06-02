"""Serviço de envio, extração, classificação e confirmação de conteúdo do autor.

Pipeline:
1. Receber upload do DOCX do autor (rota chama `processar_upload`).
2. Extrair elementos (parágrafos, headings, tabelas, imagens inline).
3. Classificar elementos e tentar posicioná-los nos capítulos
   (`CapituloDocumento`) da versão de trabalho, casando por:
   - Heading 1..N: bate pelo título normalizado contra capítulos.
   - Parágrafos/tabelas entre dois headings: ficam no último heading casado.
   - Conteúdo solto (sem heading): vai para o capítulo destino indicado
     (ou o primeiro capítulo, como fallback).
4. Extrair sugestões do DOCX upado (títulos, figuras, tabelas com legendas).
5. Gerar uma `PrevisualizacaoConteudo` por capítulo destino, com HTML
   básico para o autor revisar antes da confirmação.
6. Confirmação:
   - 'importar' → persiste o DOCX por capítulo em `conteudo_docx` e marca
     o envio como 'importado'.
   - 'rejeitar' → descarta o envio (status 'rejeitado') sem alterar
     capítulos.
"""

import json
import os
import re
import tempfile
import unicodedata
from io import BytesIO

from docx import Document

from app import db
from app.models.envio_conteudo import EnvioConteudo
from app.models.previsualizacao_conteudo import PrevisualizacaoConteudo
from app.models.capitulo_documento import CapituloDocumento

TIPO_PREVIA_DOCX_SUGERIDO = "docx_sugerido"
VERSAO_DOCX_SUGERIDO = "analise_upload_v2"


def _normalizar(texto):
    """Lowercase + sem acentos + colapsa espaços."""
    if not texto:
        return ""
    s = unicodedata.normalize("NFD", texto)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    s = re.sub(r"\s+", " ", s).strip().lower()
    # Remover numeração inicial: "1. Introdução" → "introdução"
    s = re.sub(r"^\s*(?:\d+(?:\.\d+)*|[ivx]+|[a-z])[\.\)]\s*", "", s)
    return s


def _nome_estilo_paragrafo(paragrafo):
    """Retorna nome do estilo do parágrafo com fallback seguro."""
    estilo = getattr(paragrafo, "style", None)
    return getattr(estilo, "name", "") or ""


def _elemento_xml(objeto):
    """Retorna o elemento XML interno de objetos python-docx, se existir."""
    return getattr(objeto, "_element", None)


def _xml_findall(objeto, caminho, namespaces=None):
    """Executa findall no XML interno quando ele existe."""
    elemento = _elemento_xml(objeto)
    if elemento is None:
        return []
    return elemento.findall(caminho, namespaces)


def _xml_xpath(objeto, caminho):
    """Executa xpath no XML interno quando ele existe."""
    elemento = _elemento_xml(objeto)
    if elemento is None:
        return []
    return elemento.xpath(caminho)


def _heading_nivel(estilo):
    """Retorna nível do heading (1..9) ou None."""
    if not estilo:
        return None
    s = estilo.strip().lower()
    if s.startswith("heading"):
        try:
            return int(s.replace("heading", "").strip() or "1")
        except ValueError:
            return None
    if s in ("title", "titulo", "título"):
        return 0
    return None


def gerar_docx_segmento(envio, capitulo):
    """Gera em memória um DOCX correspondente apenas ao segmento do
    envio que foi classificado para o capítulo informado.

    Estratégia: percorre os parágrafos do DOCX original, identifica
    o capítulo ativo por casamento de Heading com `titulo_capitulo`
    e acumula os parágrafos entre o heading casado e o próximo
    heading casado.

    Retorna bytes (.docx) ou None se nada foi atribuído.
    """
    if not envio.caminho_arquivo or not os.path.exists(envio.caminho_arquivo):
        return None

    capitulos = (
        CapituloDocumento.query.filter_by(
            id_relatorio=envio.id_relatorio,
            ativo=True,
        )
        .order_by(CapituloDocumento.ordem_capitulo)
        .all()
    )
    mapa = {}
    for cap in capitulos:
        chave = _normalizar(cap.titulo_capitulo)
        if chave:
            mapa.setdefault(chave, cap)

    alvo_norm = _normalizar(capitulo.titulo_capitulo)

    doc_origem = Document(envio.caminho_arquivo)
    novo = Document()

    coletando = False
    qtd = 0
    for para in doc_origem.paragraphs:
        estilo = _nome_estilo_paragrafo(para)
        texto = para.text.strip()
        nivel = _heading_nivel(estilo)
        if nivel is not None and texto:
            norm = _normalizar(texto)
            if norm in mapa:
                coletando = norm == alvo_norm
                continue
        if not coletando:
            continue
        # Copia o parágrafo (preserva runs básicos)
        novo_para = novo.add_paragraph()
        for run in para.runs:
            r = novo_para.add_run(run.text)
            if run.bold:
                r.bold = True
            if run.italic:
                r.italic = True
            if run.underline:
                r.underline = True
        qtd += 1

    if qtd == 0:
        # Se nada foi coletado pelo casamento de heading, devolve
        # placeholder mínimo (1 parágrafo) para o editor abrir
        novo.add_paragraph(f'(Sem conteúdo classificado para "{capitulo.titulo_capitulo}")')

    buf = BytesIO()
    novo.save(buf)
    return buf.getvalue()


class ServicoEnvioAutor:
    """Orquestra upload, extração, classificação e confirmação de envios."""

    @staticmethod
    def diretorio_uploads(base_dir, id_relatorio):
        """Diretório onde os uploads são salvos: storage/uploads/{id}/."""
        return os.path.join(base_dir, "storage", "uploads", str(id_relatorio))

    @staticmethod
    def diretorio_versoes_sugeridas(base_dir, id_relatorio):
        """Diretório das versões sugeridas: storage/VERSAO_SUGERIDA/{id}/."""
        return os.path.join(base_dir, "storage", "VERSAO_SUGERIDA", str(id_relatorio))

    @classmethod
    def _descartar_envios_anteriores(
        cls, *, id_relatorio, id_capitulo_destino, id_envio_atual=None
    ):
        """Descarta TODOS os envios anteriores do mesmo
        (relatório, capítulo) — qualquer status. Garante a regra
        de unicidade total (1 envio por capítulo) suportada pelo
        UNIQUE INDEX `ux_envios_por_capitulo` no Postgres.

        Quando `id_envio_atual` for informado, esse envio é
        preservado; caso contrário, todos são descartados (uso
        antes de criar um novo).
        """
        q = EnvioConteudo.query.filter(
            EnvioConteudo.id_relatorio == id_relatorio,
            EnvioConteudo.id_capitulo_destino == id_capitulo_destino,
        )
        if id_envio_atual is not None:
            q = q.filter(EnvioConteudo.id_envio_conteudo != id_envio_atual)
        for ev in q.all():
            cls._descartar_envio(ev)

    @classmethod
    def _descartar_envio(cls, envio):
        """Remove envio do banco e seu arquivo de storage. Cascata
        para PrevisualizacaoConteudo (não há cascade no model)."""
        try:
            if envio.caminho_arquivo and os.path.exists(envio.caminho_arquivo):
                os.remove(envio.caminho_arquivo)
        except OSError:
            pass
        for prev in list(envio.previsualizacoes or []):
            if prev.caminho_saida and os.path.exists(prev.caminho_saida):
                try:
                    os.remove(prev.caminho_saida)
                except OSError:
                    pass
            db.session.delete(prev)
        db.session.delete(envio)

    @staticmethod
    def caminho_docx_sugerido(envio):
        """Retorna o caminho da versão DOCX sugerida/editável do envio."""
        for prev in envio.previsualizacoes or []:
            if prev.tipo_previsualizacao == TIPO_PREVIA_DOCX_SUGERIDO:
                return prev.caminho_saida
        return None

    @classmethod
    def garantir_docx_sugerido(cls, envio):
        """Cria a versão sugerida processada a partir da análise do upload."""
        caminho = cls.caminho_docx_sugerido(envio)
        if caminho and os.path.exists(caminho) and cls._docx_sugerido_esta_processado(envio):
            return caminho
        if not envio.caminho_arquivo or not os.path.exists(envio.caminho_arquivo):
            return None

        if not caminho:
            caminho = cls._caminho_novo_docx_sugerido(envio)
        cls._gerar_docx_sugerido_processado(envio, caminho)
        if cls.caminho_docx_sugerido(envio) is None:
            db.session.add(
                PrevisualizacaoConteudo(
                    id_envio_conteudo=envio.id_envio_conteudo,
                    tipo_previsualizacao=TIPO_PREVIA_DOCX_SUGERIDO,
                    caminho_saida=caminho,
                )
            )
        db.session.flush()
        return caminho

    @classmethod
    def _caminho_novo_docx_sugerido(cls, envio):
        """Monta caminho com vínculo explícito ao upload correspondente."""
        base_dir = os.path.dirname(
            os.path.dirname(
                os.path.dirname(os.path.dirname(os.path.abspath(envio.caminho_arquivo)))
            )
        )
        dir_sugerido = cls.diretorio_versoes_sugeridas(base_dir, envio.id_relatorio)
        os.makedirs(dir_sugerido, exist_ok=True)
        nome_upload = os.path.splitext(os.path.basename(envio.caminho_arquivo))[0]
        nome_sugerido = f"envio_{envio.id_envio_conteudo}_" f"upload_{nome_upload}_sugerido.docx"
        return os.path.join(dir_sugerido, nome_sugerido)

    @staticmethod
    def _docx_sugerido_esta_processado(envio):
        try:
            estrutura = json.loads(envio.sugestoes_json or "{}")
        except (TypeError, ValueError):
            return False
        return estrutura.get("docx_sugerido", {}).get("gerado_por") == (VERSAO_DOCX_SUGERIDO)

    @classmethod
    def _gerar_docx_sugerido_processado(cls, envio, caminho_saida):
        """Gera DOCX sugerido por serviço canônico especializado."""
        from app.models.relatorio_producao import RelatorioProducao  # noqa: C0415
        from app.services.servico_perfil_formatacao import PerfilFormatacao  # noqa: C0415
        from app.services.servico_versao_sugerida_canonica import (  # noqa: C0415
            ServicoVersaoSugeridaCanonica,
        )

        rel = RelatorioProducao.query.get(envio.id_relatorio)
        perfil = PerfilFormatacao.de_relatorio(rel) if rel else PerfilFormatacao()
        metricas_aplicadas, metricas = ServicoVersaoSugeridaCanonica.gerar(
            envio=envio,
            relatorio=rel,
            perfil=perfil,
            caminho_saida=caminho_saida,
        )
        cls._registrar_diagnostico_sugestao(envio, perfil, metricas_aplicadas, metricas)

    @staticmethod
    def _registrar_diagnostico_sugestao(envio, perfil, metricas_aplicadas, metricas):
        try:
            estrutura = json.loads(envio.sugestoes_json or "{}")
        except (TypeError, ValueError):
            estrutura = {}
        estrutura["docx_sugerido"] = {
            "gerado_por": VERSAO_DOCX_SUGERIDO,
            "base": "upload_original",
            "processamentos": metricas_aplicadas,
            "metricas_obrigatorias_aplicadas": True,
            "perfil_formatacao": getattr(perfil, "origem", "default"),
            "avisos_perfil": getattr(perfil, "avisos", []),
            "biblioteca_canonica": {
                "diretorio": (metricas or {}).get("diretorio"),
                "docx_base": (metricas or {}).get("docx_base"),
                "tem_formatacao": bool((metricas or {}).get("formatacao")),
                "tem_capitulos": bool((metricas or {}).get("capitulos")),
                "tem_macro": bool((metricas or {}).get("macro")),
            },
        }
        envio.sugestoes_json = json.dumps(estrutura, ensure_ascii=False)

    @classmethod
    def salvar_docx_sugerido(cls, envio, conteudo):
        """Persiste alterações do autor na versão sugerida já gerada."""
        caminho = cls.caminho_docx_sugerido(envio)
        if not caminho or not os.path.exists(caminho):
            return False, ("DOCX sugerido ainda não foi gerado pelo sistema.")
        with open(caminho, "wb") as arquivo:
            arquivo.write(conteudo)
        db.session.commit()
        return True, "DOCX sugerido salvo."

    @classmethod
    def rejeitar_para_novo_upload(cls, envio):
        """Remove upload, prévia sugerida e registros para permitir novo envio."""
        id_relatorio = envio.id_relatorio
        id_capitulo = envio.id_capitulo_destino
        cls._descartar_envio(envio)
        db.session.commit()
        return id_relatorio, id_capitulo

    @classmethod
    def processar_upload(
        cls, *, id_relatorio, id_usuario, arquivo_storage, base_dir, id_capitulo_destino=None
    ):
        """Persiste o arquivo e gera registros de envio + prévias.

        O `id_capitulo_destino` é OBRIGATÓRIO no novo fluxo: o
        autor sempre acessa o upload via uma URL específica de
        capítulo (`/capitulo/<id>/upload`), e todo o conteúdo do
        DOCX upado será mesclado naquele capítulo (preservando o
        heading e sobrescrevendo o conteúdo antigo) na confirmação.

        Retorna o `EnvioConteudo` criado, já com prévias associadas.
        """
        if not id_capitulo_destino:
            raise ValueError(
                "id_capitulo_destino é obrigatório no fluxo atual de "
                "envio do autor — o destino do conteúdo é fixado pela "
                "URL de upload."
            )

        # Regra: só 1 envio por (relatório, capítulo), qualquer status.
        # Descarta tudo que ja existe para esse par antes de criar o
        # novo registro. Garante consistencia com o UNIQUE INDEX
        # `ux_envios_por_capitulo` no Postgres.
        cls._descartar_envios_anteriores(
            id_relatorio=id_relatorio,
            id_capitulo_destino=id_capitulo_destino,
        )

        dir_destino = cls.diretorio_uploads(base_dir, id_relatorio)
        os.makedirs(dir_destino, exist_ok=True)

        from werkzeug.utils import secure_filename

        nome = secure_filename(arquivo_storage.filename or "envio.docx")
        # Evitar colisão preservando histórico
        timestamp = __import__("datetime").datetime.now().strftime("%Y%m%d%H%M%S")
        nome_final = f"{timestamp}_{nome}"
        caminho_final = os.path.join(dir_destino, nome_final)
        arquivo_storage.save(caminho_final)

        envio = EnvioConteudo(
            id_relatorio=id_relatorio,
            id_usuario=id_usuario,
            nome_arquivo=nome,
            caminho_arquivo=caminho_final,
            status_envio="em_previa",
            id_capitulo_destino=id_capitulo_destino,
        )
        db.session.add(envio)
        db.session.flush()

        # Extração + classificação + prévia
        cls._gerar_previas(envio, id_capitulo_destino)
        cls.garantir_docx_sugerido(envio)

        db.session.commit()
        return envio

    @classmethod
    def _gerar_previas(cls, envio, id_capitulo_destino):
        """Lê o DOCX e gera PrevisualizacaoConteudo por capítulo destino.

        Extrai também a estrutura completa do DOCX:
        - Árvore hierárquica de capítulos e subcapítulos
        - Figuras com legendas organizadas por capítulo
        - Tabelas com legendas organizadas por capítulo
        """
        capitulos = (
            CapituloDocumento.query.filter_by(
                id_relatorio=envio.id_relatorio,
                ativo=True,
            )
            .order_by(CapituloDocumento.ordem_capitulo)
            .all()
        )

        # Mapa de capítulos por título normalizado
        mapa = {}
        for cap in capitulos:
            chave = _normalizar(cap.titulo_capitulo)
            if chave:
                mapa.setdefault(chave, cap)

        try:
            doc = Document(envio.caminho_arquivo)
        except (OSError, ValueError) as e:
            prev = PrevisualizacaoConteudo(
                id_envio_conteudo=envio.id_envio_conteudo,
                tipo_previsualizacao="erro",
                resultado_html=(f'<div class="ew__erro">Erro ao ler DOCX: {e}</div>'),
            )
            db.session.add(prev)
            return

        # Extrair estrutura completa do DOCX usando ServicoExtracaoCanonica
        estrutura = cls._extrair_estrutura_completa(doc)

        # Detectar renomeações de capítulos: o autor pode ter mudado
        # o título de capítulos cujos índices já existem no banco.
        # - nível 1 e 2: vão para fila de aprovação do coordenador
        # - nível >= 3: aplicam automaticamente NO MOMENTO DA IMPORTAÇÃO
        # Aqui apenas registramos as duas listas; a aplicação real
        # acontece em `confirmar(importar)` onde temos `caminho_master`.
        renomeacoes = cls._detectar_renomeacoes(estrutura.get("arvore_estrutural", []), capitulos)
        pendentes = renomeacoes.get("pendentes", [])
        automaticas = renomeacoes.get("automaticas", [])
        if automaticas:
            estrutura["renomeacoes_automaticas_pendentes"] = automaticas
        if pendentes:
            estrutura["renomeacoes_pendentes"] = pendentes
            cls._notificar_coordenadores_renomeacao(envio, pendentes)

        # Armazenar estrutura no envio para uso na prévia
        envio.sugestoes_json = json.dumps(estrutura)

        # Particionar conteúdo: header (antes do primeiro heading casado)
        # + listas de "segmentos" por capítulo destino.
        cap_atual = None
        if id_capitulo_destino:
            cap_atual = CapituloDocumento.query.get(id_capitulo_destino)

        segmentos_por_cap = {}
        # Se há destino fixo, todo conteúdo vai para ele.
        forcar_destino = cap_atual is not None and not mapa

        # Iterar parágrafos
        for para in doc.paragraphs:
            estilo = _nome_estilo_paragrafo(para)
            texto = para.text.strip()
            nivel = _heading_nivel(estilo)

            if nivel is not None and texto:
                norm = _normalizar(texto)
                if norm in mapa:
                    cap_atual = mapa[norm]
                    # Heading que casa com capítulo: marca início,
                    # não duplica o título no conteúdo do capítulo
                    continue
                # Heading que não casa — se já temos cap_atual, segue
                # incluindo o heading como subseção; caso contrário,
                # mantém como conteúdo solto.

            destino = cap_atual
            if destino is None and id_capitulo_destino:
                destino = CapituloDocumento.query.get(id_capitulo_destino)
            if destino is None and capitulos:
                destino = capitulos[0]
            if destino is None:
                continue  # sem capítulos: nada a fazer

            chave = destino.id_capitulo_documento
            segmentos_por_cap.setdefault(chave, []).append(
                {
                    "tipo": "paragrafo",
                    "estilo": estilo,
                    "nivel": nivel,
                    "texto": texto,
                }
            )

        # Tabelas vão integralmente para o capítulo ativo no momento
        # da leitura — como python-docx não preserva ordenação mista
        # entre paragraphs/tables sem iterar pelo body, fazemos a
        # aproximação: cada tabela do DOCX vai para o último cap_atual
        # (ou destino solicitado).
        cap_destino_tabelas = (
            cap_atual
            or (CapituloDocumento.query.get(id_capitulo_destino) if id_capitulo_destino else None)
            or (capitulos[0] if capitulos else None)
        )
        if cap_destino_tabelas is not None:
            for table in doc.tables:
                segmentos_por_cap.setdefault(cap_destino_tabelas.id_capitulo_documento, []).append(
                    {
                        "tipo": "tabela",
                        "linhas": [[c.text for c in row.cells] for row in table.rows],
                    }
                )

        if forcar_destino:
            # Compatibilidade: já feito acima através de cap_atual.
            pass

        # Gerar prévia HTML por capítulo destino
        for id_cap, segmentos in segmentos_por_cap.items():
            cap = CapituloDocumento.query.get(id_cap)
            if cap is None:
                continue
            html_parts = [
                f'<section class="ew__previa-cap" data-cap="{id_cap}">',
                (f'<h2>{cap.indice_capitulo or ""} ' f"{cap.titulo_capitulo}</h2>"),
            ]
            for seg in segmentos:
                if seg["tipo"] == "paragrafo":
                    nivel = seg["nivel"]
                    texto_html = (
                        seg["texto"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    )
                    if nivel and nivel >= 1:
                        tag = f"h{min(max(nivel + 1, 2), 6)}"
                        html_parts.append(f"<{tag}>{texto_html}</{tag}>")
                    else:
                        html_parts.append(f"<p>{texto_html}</p>")
                elif seg["tipo"] == "tabela":
                    html_parts.append('<table class="ew__previa-tbl">')
                    for linha in seg["linhas"]:
                        html_parts.append("<tr>")
                        for celula in linha:
                            celula_html = (
                                celula.replace("&", "&amp;")
                                .replace("<", "&lt;")
                                .replace(">", "&gt;")
                            )
                            html_parts.append(f"<td>{celula_html}</td>")
                        html_parts.append("</tr>")
                    html_parts.append("</table>")
            html_parts.append("</section>")

            prev = PrevisualizacaoConteudo(
                id_envio_conteudo=envio.id_envio_conteudo,
                tipo_previsualizacao="parcial",
                resultado_html="\n".join(html_parts),
                caminho_saida=str(id_cap),
            )
            db.session.add(prev)

        # Caso nada tenha sido classificado, registrar uma prévia geral
        if not segmentos_por_cap:
            prev = PrevisualizacaoConteudo(
                id_envio_conteudo=envio.id_envio_conteudo,
                tipo_previsualizacao="vazio",
                resultado_html=(
                    '<div class="ew__erro">Nenhum conteúdo identificável '
                    "foi extraído do DOCX (verifique se há texto e "
                    "cabeçalhos compatíveis com a estrutura do relatório)."
                    "</div>"
                ),
            )
            db.session.add(prev)

    @classmethod
    def confirmar(cls, *, envio, acao):
        """Aplica a decisão do autor sobre o envio.

        - acao='importar': mescla o DOCX do autor IN-PLACE no DOCX
          em produção (caminho_template do relatório), substituindo
          o conteúdo do capítulo destino preservando o heading.
        - acao='rejeitar': marca como rejeitado e não altera DOCX.

        Implementação: delega para `servico_merge_docx.substituir_capitulo`,
        que usa docxcompose para preservar imagens, estilos e numeração.
        """
        if acao == "rejeitar":
            envio.status_envio = "rejeitado"
            db.session.commit()
            return {"ok": True, "acao": "rejeitado"}

        if acao != "importar":
            return {"ok": False, "erro": "Ação inválida"}

        if not envio.id_capitulo_destino:
            return {
                "ok": False,
                "erro": (
                    "Envio sem capítulo destino — não é possível "
                    "identificar onde mesclar o conteúdo."
                ),
            }

        cap_destino = CapituloDocumento.query.get(envio.id_capitulo_destino)
        if not cap_destino:
            return {"ok": False, "erro": "Capítulo destino não encontrado."}

        from app.models.relatorio_producao import RelatorioProducao
        from app.services.servico_merge_docx import (
            substituir_capitulo,
            sincronizar_subcapitulos,
        )

        rel = RelatorioProducao.query.get(envio.id_relatorio)
        # Gate de bloqueio: relatório finalizado não aceita merge.
        from app.services import servico_relatorio_core as relatorio_core

        if relatorio_core.esta_bloqueado(rel):
            return {
                "ok": False,
                "erro": (
                    "Relatório finalizado/bloqueado — não é possível "
                    "mesclar novos conteúdos. Crie uma nova versão para "
                    "continuar a edição."
                ),
            }
        if not rel or not rel.caminho_template:
            return {
                "ok": False,
                "erro": (
                    "Relatório de produção sem DOCX em "
                    "caminho_template — não é possível mesclar."
                ),
            }

        if not os.path.exists(rel.caminho_template):
            return {
                "ok": False,
                "erro": (f"DOCX de produção não encontrado em " f"{rel.caminho_template}"),
            }

        try:
            segmento_bytes = gerar_docx_segmento(envio, cap_destino)
            caminho_temporario = None
            caminho_autor = cls.caminho_docx_sugerido(envio) or envio.caminho_arquivo
            if segmento_bytes:
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as arquivo_tmp:
                    arquivo_tmp.write(segmento_bytes)
                    caminho_temporario = arquivo_tmp.name
                caminho_autor = caminho_temporario
            try:
                ok = substituir_capitulo(
                    caminho_master=rel.caminho_template,
                    capitulo=cap_destino,
                    caminho_autor=caminho_autor,
                    preservar_heading=True,
                )
            finally:
                if caminho_temporario:
                    try:
                        os.unlink(caminho_temporario)
                    except OSError:
                        pass
        except (OSError, ValueError, RuntimeError) as e:
            return {
                "ok": False,
                "erro": f"Falha ao mesclar no DOCX em produção: {e}",
            }

        if not ok:
            return {
                "ok": False,
                "erro": (
                    f'Capítulo "{cap_destino.titulo_capitulo}" não foi '
                    f"localizado no DOCX em produção. Verifique se o "
                    f"heading correspondente existe no arquivo."
                ),
            }

        cap_destino.status_capitulo = "aguardando_aprovacao"
        # Espelha o novo status do capitulo no proprio envio para que
        # a tabela `envios_conteudo` reflita o estado atual sem
        # depender de JOIN. O listener do modelo so dispara quando o
        # `id_capitulo_destino` e (re)atribuido — aqui ja temos o
        # id, entao copiamos diretamente.
        envio.status_capitulo_id = cap_destino.status_capitulo_id
        envio.status_envio = "importado"

        # Regra: só 1 envio por (relatório, capítulo). Descarta os
        # anteriores (qualquer status) para garantir o constraint
        # UNIQUE total. Em principio nao existem anteriores nesta
        # fase porque processar_upload ja descarta tudo, mas mantemos
        # a chamada como defesa em profundidade.
        cls._descartar_envios_anteriores(
            id_relatorio=envio.id_relatorio,
            id_capitulo_destino=envio.id_capitulo_destino,
            id_envio_atual=envio.id_envio_conteudo,
        )

        # Aplicar renomeações de capítulos:
        # 1) Automáticas (nível >= 3) — aplicam sempre na importação
        # 2) Pendentes (nível 1 e 2) — ficam guardadas; coordenador
        #    aprova/rejeita ao revisar o capítulo.
        # Atualizamos tanto banco quanto DOCX em produção para que
        # cabeçalho e árvore de capítulos fiquem consistentes.
        renom_auto = cls._aplicar_renomeacoes_automaticas(envio, rel.caminho_template)
        renomeacoes_aplicadas = renom_auto or []

        # Sincronizar subcapítulos no banco a partir dos subheadings
        # que o autor enviou no DOCX. Isso garante que a árvore na UI
        # reflita a estrutura recém-mesclada (cada Heading 2/3/4 do
        # upload vira CapituloDocumento filho de cap_destino).
        try:
            sync = sincronizar_subcapitulos(db.session, cap_destino, rel.caminho_template)
        except (OSError, ValueError, RuntimeError) as e:
            # Merge já foi escrito em disco — falha de sincronização
            # não deve reverter o conteúdo. Logamos e seguimos.
            sync = {"erro": str(e)}

        # Fase 2 — Captioning + cross-references:
        # 1) reindexar_captions: numera figuras/tabelas/equações
        #    hierarquicamente e devolve mapa_labels.
        # 2) substituir_referencias: troca {{fig:x}}, {{tab:x}},
        #    {{eq:x}}, {{ref:x}} no corpo pelos números.
        captions = {}
        cross_refs = {}
        try:
            from app.services.servico_captioning import reindexar_captions
            from app.services.servico_cross_refs import substituir_referencias
            from app.services.servico_perfil_formatacao import (
                PerfilFormatacao,
            )

            perfil = PerfilFormatacao.de_relatorio(rel)
            captions = reindexar_captions(rel.caminho_template, perfil=perfil)
            mapa = captions.get("mapa_labels", {}) if isinstance(captions, dict) else {}
            cross_refs = substituir_referencias(rel.caminho_template, mapa)
        except (OSError, ValueError, RuntimeError) as e:
            captions = captions or {"erro": str(e)}
            cross_refs = {"erro": str(e)}

        db.session.commit()
        return {
            "ok": True,
            "acao": "importado",
            "capitulos_atualizados": 1,
            "capitulo_destino_id": cap_destino.id_capitulo_documento,
            "subcapitulos_sync": sync,
            "captions": captions,
            "cross_refs": cross_refs,
            "renomeacoes_aplicadas": renomeacoes_aplicadas,
        }

    @classmethod
    def aprovar_capitulo(cls, *, capitulo, coordenador, observacao=None):
        """Coordenador aprova um capítulo que está aguardando.

        Aplica renomeações pendentes (nível 1/2) que ainda não foram
        aplicadas — banco + DOCX em produção. Atualiza status do
        capítulo para 'aprovado' e notifica o autor.
        """
        from app.models.relatorio_producao import (  # noqa: C0415
            RelatorioProducao,
        )
        from app.models.notificacao import Notificacao  # noqa: C0415

        if capitulo.status_capitulo != "aguardando_aprovacao":
            return {
                "ok": False,
                "erro": (
                    "Capítulo não está aguardando aprovação "
                    f"(status atual: {capitulo.status_capitulo})."
                ),
            }

        rel = RelatorioProducao.query.get(capitulo.id_relatorio)
        if not rel or not rel.caminho_template:
            return {
                "ok": False,
                "erro": "Relatório não localizado ou sem DOCX.",
            }

        # Aplicar renomeações pendentes do envio mais recente do
        # capítulo (nível 1/2). Se houver mais de um envio importado,
        # processamos todos em ordem de criação.
        envios = (
            EnvioConteudo.query.filter_by(
                id_capitulo_destino=capitulo.id_capitulo_documento,
                status_envio="importado",
            )
            .order_by(EnvioConteudo.criado_em.asc())
            .all()
        )
        renomeacoes_aplicadas = []
        for envio in envios:
            aplicadas = cls._aplicar_renomeacoes_pendentes(
                envio, caminho_master=rel.caminho_template
            )
            if aplicadas:
                renomeacoes_aplicadas.extend(aplicadas)

        capitulo.status_capitulo = "aprovado"

        # Espelha o novo status do capitulo em todos os envios desse
        # capitulo para manter a tabela `envios_conteudo` consistente.
        for envio in envios:
            envio.status_capitulo_id = capitulo.status_capitulo_id

        # Notifica o autor (último a enviar) sobre a aprovação.
        if envios:
            ultimo = envios[-1]
            obs = observacao.strip() if observacao else ""
            mensagem = f"O coordenador aprovou o capítulo " f'"{capitulo.titulo_capitulo}".'
            if obs:
                mensagem += f" Observação: {obs}"
            db.session.add(
                Notificacao(
                    id_usuario_destino=ultimo.id_usuario,
                    tipo_notificacao="revisao",
                    mensagem=mensagem,
                    lida=False,
                )
            )

        db.session.commit()
        return {
            "ok": True,
            "acao": "aprovado",
            "capitulo_id": capitulo.id_capitulo_documento,
            "renomeacoes_aplicadas": renomeacoes_aplicadas,
        }

    @classmethod
    def rejeitar_capitulo(cls, *, capitulo, coordenador, observacao=None):
        """Coordenador rejeita um capítulo aguardando aprovação.

        Volta o status para 'em_edicao' e notifica o autor com a
        observação. NÃO reverte o conteúdo já mesclado no DOCX em
        produção — o autor pode fazer um novo envio para sobrescrever.
        """
        from app.models.notificacao import Notificacao  # noqa: C0415

        if capitulo.status_capitulo != "aguardando_aprovacao":
            return {
                "ok": False,
                "erro": (
                    "Capítulo não está aguardando aprovação "
                    f"(status atual: {capitulo.status_capitulo})."
                ),
            }

        capitulo.status_capitulo = "em_edicao"

        # Espelha o novo status nos envios do capitulo.
        envios_do_cap = EnvioConteudo.query.filter_by(
            id_capitulo_destino=capitulo.id_capitulo_documento,
        ).all()
        for envio in envios_do_cap:
            envio.status_capitulo_id = capitulo.status_capitulo_id

        ultimo = (
            EnvioConteudo.query.filter_by(
                id_capitulo_destino=capitulo.id_capitulo_documento,
                status_envio="importado",
            )
            .order_by(EnvioConteudo.criado_em.desc())
            .first()
        )
        if ultimo:
            obs = observacao.strip() if observacao else ""
            mensagem = (
                f"O coordenador solicitou ajustes no capítulo " f'"{capitulo.titulo_capitulo}".'
            )
            if obs:
                mensagem += f" Observação: {obs}"
            db.session.add(
                Notificacao(
                    id_usuario_destino=ultimo.id_usuario,
                    tipo_notificacao="revisao",
                    mensagem=mensagem,
                    lida=False,
                )
            )

        db.session.commit()
        return {
            "ok": True,
            "acao": "rejeitado",
            "capitulo_id": capitulo.id_capitulo_documento,
        }

    @classmethod
    def _extrair_estrutura_completa(cls, doc):
        """Extrai estrutura completa do DOCX usando ServicoExtracaoCanonica.

        Retorna dict com:
        - capitulos: árvore hierárquica de capítulos e subcapítulos
        - legendas: figuras e tabelas com legendas (agregados)
        - arvore_estrutural: árvore que mistura capítulos com figuras,
          tabelas e equações como nós individuais (cada um com seu
          índice, título e capítulo pai), pronta para a UI exibir
          como árvore só de estrutura.
        """
        from app.services.servico_extracao_canonica import (  # noqa: C0415
            ServicoExtracaoCanonica,
        )

        # Extrair árvore de capítulos
        capitulos_arvore = ServicoExtracaoCanonica.extrair_capitulos(doc)

        # Se não encontrou capítulos via Heading, tentar detecção por padrão
        if not capitulos_arvore:
            capitulos_arvore = cls._extrair_capitulos_por_padrao(doc)

        # Extrair legendas (figuras e tabelas) — agregados
        legendas = ServicoExtracaoCanonica.extrair_legendas(doc)

        # Construir árvore estrutural (capítulos + figuras + tabelas
        # + equações como nós individuais sob seu capítulo pai).
        arvore_estrutural = cls._construir_arvore_estrutural(doc, capitulos_arvore)

        # Organizar figuras e tabelas por capítulo
        estrutura = {
            "capitulos": capitulos_arvore,
            "legendas": legendas,
            "arvore_estrutural": arvore_estrutural,
        }

        return estrutura

    @classmethod
    def extrair_estrutura_completa(cls, doc):
        """Extrai estrutura completa do DOCX por meio da API publica."""
        return cls._extrair_estrutura_completa(doc)

    @classmethod
    def _construir_arvore_estrutural(cls, doc, capitulos_arvore):
        """Constrói árvore estrutural rica: capítulos + figuras +
        tabelas + equações, cada um como nó individual com índice
        próprio, ancorado sob o capítulo pai mais próximo no DOCX.

        Cada nó tem o formato:
            {
                'tipo': 'capitulo'|'figura'|'tabela'|'equacao',
                'indice': '1.2'|'Figura 1.1'|'Tabela 2.3'|'Equação 1',
                'titulo': '...',
                'nivel': 1..N (apenas para capitulos),
                'filhos': [...]   (apenas para capitulos)
            }
        """
        # 1. Localizar os parágrafos índice de cada heading (capítulo)
        #    para ancorar os elementos visuais.
        capitulos_planos = cls._achatar_capitulos(capitulos_arvore)

        # 2. Coletar elementos visuais (figuras, tabelas, equações)
        #    com índice de parágrafo e capítulo pai.
        elementos = cls._coletar_elementos_visuais(doc)

        # 3. Distribuir elementos sob seus capítulos pais.
        #    Para cada capítulo plano, anexamos os elementos cujo
        #    índice de parágrafo cai entre o início do capítulo e o
        #    início do próximo capítulo.
        cls._distribuir_elementos_em_capitulos(capitulos_planos, elementos)

        # 4. Re-aninhar os capítulos planos preservando filhos
        #    (capítulos_arvore mantém a hierarquia; capitulos_planos
        #    agora carrega filhos extras: figuras, tabelas, equações).
        return capitulos_arvore

    @classmethod
    def _aplicar_renomeacoes_automaticas(cls, envio, caminho_master):
        """Aplica as renomeações de subcapítulos (nível >= 3) que
        foram detectadas no upload e armazenadas em
        `renomeacoes_automaticas_pendentes`. Sempre aplica no
        momento da importação — não exige aprovação."""
        if not envio.sugestoes_json:
            return []
        try:
            estrutura = json.loads(envio.sugestoes_json)
        except (ValueError, TypeError):
            return []

        auto = estrutura.get("renomeacoes_automaticas_pendentes") or []
        if not auto:
            return []

        cls._aplicar_renomeacoes_imediatas(auto, caminho_master)

        aplicadas = [
            {
                "id_capitulo_documento": r.get("id_capitulo_documento"),
                "de": r.get("titulo_atual"),
                "para": r.get("titulo_proposto"),
                "indice": r.get("indice"),
                "nivel": r.get("nivel"),
            }
            for r in auto
        ]
        estrutura["renomeacoes_automaticas_pendentes"] = []
        estrutura["renomeacoes_aplicadas"] = estrutura.get("renomeacoes_aplicadas", []) + aplicadas
        envio.sugestoes_json = json.dumps(estrutura)
        return aplicadas

    @classmethod
    def _aplicar_renomeacoes_imediatas(cls, lista, caminho_master=None):
        """Aplica em `CapituloDocumento.titulo_capitulo` a lista
        de renomeações automáticas (subcapítulos nível >= 3) sem
        exigir aprovação. Quando `caminho_master` é informado,
        atualiza também o texto do heading no DOCX em produção
        (preservando estilo, numeração e bookmarks)."""
        from app.services.servico_merge_docx import (  # noqa: C0415
            atualizar_titulo_capitulo,
        )

        for r in lista or []:
            cap = CapituloDocumento.query.get(r.get("id_capitulo_documento"))
            if not cap:
                continue
            novo = (r.get("titulo_proposto") or "").strip()
            if not novo or novo == (cap.titulo_capitulo or "").strip():
                continue
            # Atualiza o DOCX antes de mudar o titulo no banco — assim
            # `localizar_range_capitulo` ainda casa pelo nome antigo.
            if caminho_master:
                try:
                    atualizar_titulo_capitulo(caminho_master, cap, novo)
                except (OSError, ValueError, RuntimeError):
                    pass
            cap.titulo_capitulo = novo

    @classmethod
    def _notificar_coordenadores_renomeacao(cls, envio, pendentes):
        """Notifica coordenadores quando há renomeações de capítulo
        nível 1/2 propostas pelo autor que precisam de aprovação."""
        from app.models.notificacao import Notificacao  # noqa: C0415
        from app.models.usuario import Usuario  # noqa: C0415
        from app.models.dominio import Dominio  # noqa: C0415

        if not pendentes:
            return

        try:
            perfil_coord = Dominio.query.filter_by(
                tipo="perfil_usuario", valor="coordenador"
            ).first()
            if not perfil_coord:
                return
            coordenadores = Usuario.query.filter_by(perfil_id=perfil_coord.id, ativo=True).all()
        except (OSError, ValueError, RuntimeError):
            return

        n = len(pendentes)
        plural = "ões" if n != 1 else "ão"
        mensagem = (
            f"O autor propôs {n} renomeaç{plural} de capítulo no "
            f'envio "{envio.nome_arquivo}" — requer aprovação.'
        )
        for coord in coordenadores:
            db.session.add(
                Notificacao(
                    id_usuario_destino=coord.id,
                    tipo_notificacao="renomeacao_pendente",
                    mensagem=mensagem,
                    lida=False,
                )
            )

    @classmethod
    def _aplicar_renomeacoes_pendentes(cls, envio, caminho_master=None):
        """Aplica renomeações de capítulos sugeridas durante o upload.

        Lê `envio.sugestoes_json -> renomeacoes_pendentes` e atualiza
        `CapituloDocumento.titulo_capitulo` para os capítulos não
        rejeitados explicitamente. Quando `caminho_master` é informado,
        atualiza também o texto do heading no DOCX em produção.

        Retorna lista de aplicadas:
            [{'id_capitulo_documento': X, 'de': '...', 'para': '...'}]
        """
        from app.services.servico_merge_docx import (  # noqa: C0415
            atualizar_titulo_capitulo,
        )

        if not envio.sugestoes_json:
            return []
        try:
            estrutura = json.loads(envio.sugestoes_json)
        except (ValueError, TypeError):
            return []

        pendentes = estrutura.get("renomeacoes_pendentes") or []
        aplicadas = []
        for r in pendentes:
            if r.get("aprovado") is False:
                continue
            cap = CapituloDocumento.query.get(r.get("id_capitulo_documento"))
            if not cap:
                continue
            de = (cap.titulo_capitulo or "").strip()
            para = (r.get("titulo_proposto") or "").strip()
            if not para or de == para:
                continue
            # Atualiza heading no DOCX em produção ANTES de trocar o
            # titulo no banco — `localizar_range_capitulo` casa por
            # `titulo_capitulo` (precisa ser o atual no DOCX).
            if caminho_master:
                try:
                    atualizar_titulo_capitulo(caminho_master, cap, para)
                except (OSError, ValueError, RuntimeError):
                    pass
            cap.titulo_capitulo = para
            aplicadas.append(
                {
                    "id_capitulo_documento": cap.id_capitulo_documento,
                    "de": de,
                    "para": para,
                }
            )

        if aplicadas:
            estrutura["renomeacoes_pendentes"] = []
            estrutura["renomeacoes_aplicadas"] = (
                estrutura.get("renomeacoes_aplicadas", []) + aplicadas
            )
            envio.sugestoes_json = json.dumps(estrutura)

        return aplicadas

    @classmethod
    def _detectar_renomeacoes(cls, arvore, capitulos_banco):
        """Compara capítulos da árvore do envio (com índice) contra
        capítulos do banco (mesmo índice) e devolve sugestões de
        renomeação quando o título difere.

        Devolve dict com duas listas:
            {
                'pendentes': [...],   # nível 1 e 2: requer aprovação
                'automaticas': [...], # nível >= 3: aplica direto
            }

        Cada item:
            {
                'id_capitulo_documento': <int>,
                'indice': '1.2',
                'titulo_atual': '...',
                'titulo_proposto': '...',
                'nivel': <int>,
            }
        """
        # Indexar capítulos do banco por índice
        por_indice = {}
        for cap in capitulos_banco:
            ind = (cap.indice_capitulo or "").strip()
            if ind:
                por_indice[ind] = cap

        # Achata árvore do envio (só capítulos, ignora figuras/tabelas)
        propostos = []

        def visitar(lst):
            for n in lst:
                if n.get("tipo") in (None, "capitulo"):
                    if n.get("indice") and n.get("titulo"):
                        propostos.append(
                            {
                                "indice": n["indice"].strip(),
                                "titulo": n["titulo"].strip(),
                                "nivel": n.get("nivel") or (n["indice"].count(".") + 1),
                            }
                        )
                if n.get("filhos"):
                    visitar(n["filhos"])

        visitar(arvore or [])

        pendentes = []
        automaticas = []
        for p in propostos:
            cap = por_indice.get(p["indice"])
            if not cap:
                continue
            atual = (cap.titulo_capitulo or "").strip()
            if _normalizar(atual) == _normalizar(p["titulo"]) or atual == p["titulo"]:
                continue
            item = {
                "id_capitulo_documento": cap.id_capitulo_documento,
                "indice": p["indice"],
                "titulo_atual": atual,
                "titulo_proposto": p["titulo"],
                "nivel": p["nivel"],
            }
            # Nível 1 e 2: precisa aprovação do coordenador.
            # Nível >= 3: aplica direto (subcapítulo).
            if p["nivel"] <= 2:
                pendentes.append(item)
            else:
                automaticas.append(item)
        return {
            "pendentes": pendentes,
            "automaticas": automaticas,
        }

    @staticmethod
    def _achatar_capitulos(arvore):
        """Achata árvore de capítulos preservando referência (mesmas
        instâncias dos dicts), para anexar elementos sem perder a
        hierarquia."""
        planos = []

        def visitar(lst):
            for cap in lst:
                planos.append(cap)
                if cap.get("filhos"):
                    visitar(cap["filhos"])

        visitar(arvore)
        return planos

    @staticmethod
    def _coletar_elementos_visuais(doc):
        """Percorre o DOCX em ordem e coleta figuras (com legenda),
        tabelas (com legenda) e equações.

        Retorna lista ordenada por índice de parágrafo:
            [{'tipo': 'figura', 'indice_paragrafo': 12,
              'rotulo': 'Figura 1.1', 'legenda': '...'}, ...]
        """
        elementos = []
        ns = {
            "w": ("http://schemas.openxmlformats.org" "/wordprocessingml/2006/main"),
            "m": ("http://schemas.openxmlformats.org/officeDocument" "/2006/math"),
        }

        # Padrões para detectar legendas/equações
        padrao_fig = re.compile(
            r"^(?:figura|fig\.)\s*(\d+(?:[\.\-]\d+)*)\s*[-–:]?\s*(.*)$",
            re.IGNORECASE,
        )
        padrao_tab = re.compile(
            r"^(?:tabela|tab\.|quadro)\s*(\d+(?:[\.\-]\d+)*)\s*[-–:]?\s*(.*)$",
            re.IGNORECASE,
        )
        padrao_eq = re.compile(
            r"^(?:equa[cç][aã]o|eq\.)\s*(\d+(?:[\.\-]\d+)*)\s*[-–:]?\s*(.*)$",
            re.IGNORECASE,
        )

        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            if not texto:
                # Equações OOXML (omml) podem aparecer em parágrafos
                # sem texto nos runs — detectar pelo elemento <m:oMath>
                tem_math = bool(
                    _xml_findall(para, ".//m:oMath", ns) or _xml_findall(para, ".//m:oMathPara", ns)
                )
                if tem_math:
                    elementos.append(
                        {
                            "tipo": "equacao",
                            "indice_paragrafo": i,
                            "rotulo": "Equação",
                            "legenda": "",
                        }
                    )
                continue

            # Figura: legenda começa com "Figura N"
            m = padrao_fig.match(texto)
            if m:
                elementos.append(
                    {
                        "tipo": "figura",
                        "indice_paragrafo": i,
                        "rotulo": f"Figura {m.group(1)}",
                        "legenda": (m.group(2) or "").strip(),
                    }
                )
                continue

            # Tabela: legenda começa com "Tabela N" ou "Quadro N"
            m = padrao_tab.match(texto)
            if m:
                elementos.append(
                    {
                        "tipo": "tabela",
                        "indice_paragrafo": i,
                        "rotulo": f"Tabela {m.group(1)}",
                        "legenda": (m.group(2) or "").strip(),
                    }
                )
                continue

            # Equação: legenda começa com "Equação N"
            m = padrao_eq.match(texto)
            if m:
                elementos.append(
                    {
                        "tipo": "equacao",
                        "indice_paragrafo": i,
                        "rotulo": f"Equação {m.group(1)}",
                        "legenda": (m.group(2) or "").strip(),
                    }
                )
                continue

            # Equação OOXML embutida em parágrafo com texto
            tem_math = bool(
                _xml_findall(para, ".//m:oMath", ns) or _xml_findall(para, ".//m:oMathPara", ns)
            )
            if tem_math:
                elementos.append(
                    {
                        "tipo": "equacao",
                        "indice_paragrafo": i,
                        "rotulo": "Equação",
                        "legenda": texto[:80],
                    }
                )

        # Numerar equações sem número explícito (rotulo == 'Equação')
        contador_eq = 0
        for el in elementos:
            if el["tipo"] == "equacao" and el["rotulo"] == "Equação":
                contador_eq += 1
                el["rotulo"] = f"Equação {contador_eq}"

        return elementos

    @staticmethod
    def _distribuir_elementos_em_capitulos(capitulos_planos, elementos):
        """Anexa cada elemento visual à lista de filhos do capítulo
        pai mais próximo (capítulo cuja faixa de parágrafos contém
        o elemento). Mutação in-place: cada capítulo recebe os
        elementos como filhos com tipo apropriado.
        """
        if not capitulos_planos or not elementos:
            return

        # Cada capítulo plano tem 'indice_paragrafo' (definido pelo
        # _extrair_capitulos do ServicoExtracaoCanonica). Se não
        # tiver, ignoramos e devolvemos a árvore como está.
        capitulos_com_indice = [
            c for c in capitulos_planos if isinstance(c.get("indice_paragrafo"), int)
        ]
        if not capitulos_com_indice:
            return

        # Ordenar por índice de parágrafo
        capitulos_com_indice.sort(key=lambda c: c["indice_paragrafo"])

        for el in elementos:
            ip = el["indice_paragrafo"]
            # Encontrar o último capítulo cujo indice_paragrafo <= ip.
            cap_pai = None
            for cap in capitulos_com_indice:
                if cap["indice_paragrafo"] <= ip:
                    cap_pai = cap
                else:
                    break
            if cap_pai is None:
                continue
            cap_pai.setdefault("filhos", []).append(
                {
                    "tipo": el["tipo"],
                    "indice": el["rotulo"],
                    "titulo": el.get("legenda") or "",
                    "nivel": (cap_pai.get("nivel") or 1) + 1,
                    "filhos": [],
                }
            )

    @staticmethod
    def _extrair_capitulos_por_padrao(doc):
        """Extrai capítulos baseados em padrões de numeração hierárquica.

        Detecta títulos que começam com numeração como "1.", "1.1", "1.1.1"
        mesmo sem estilo Heading.
        """
        # Padrão para numeração hierárquica: 1, 1.1, 1.1.1, etc.
        padrao_numeracao = re.compile(r"^\s*(\d+(?:\.\d+)*)\s+(.+)")

        capitulos = []
        for idx, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            if not texto:
                continue

            match = padrao_numeracao.match(texto)
            if match:
                indice = match.group(1)
                titulo = match.group(2)
                nivel = indice.count(".") + 1

                capitulos.append(
                    {
                        "titulo": titulo,
                        "indice": indice,
                        "nivel": nivel,
                        "estilo": para.style.name or "Normal",
                        "tipo_elemento": "textual",
                        "indice_paragrafo": idx,
                        "filhos": [],
                    }
                )

        # Montar árvore hierárquica
        raiz = []
        pilha = []

        for item in capitulos:
            nv = item["nivel"]
            while pilha and pilha[-1][0] >= nv:
                pilha.pop()

            destino = pilha[-1][1]["filhos"] if pilha else raiz
            destino.append(item)
            pilha.append((nv, item))

        return raiz

    @staticmethod
    def _extrair_sugestoes(doc):
        """Extrai sugestões do DOCX upado de forma inteligente.

        Detecta padrões que sugerem títulos, figuras e tabelas,
        mesmo que não estejam formatados perfeitamente.

        Retorna dict com:
        - titulos: lista de headings encontrados (texto, nivel, confianca)
        - figuras: lista de figuras com/sem legendas
        - tabelas: lista de tabelas com/sem legendas
        """
        sugestoes = {
            "titulos": [],
            "figuras": [],
            "tabelas": [],
        }

        # Padrões para detecção inteligente de títulos
        padrao_numeracao = re.compile(  # noqa: E501
            r"^\s*(\d+(?:\.\d+)*|[ivx]+|[a-z])[\.\)]\s+", re.IGNORECASE
        )
        padrao_caixa_alta = re.compile(r"^[A-ZÁÀÂÃÉÊÍÓÔÕÚÇ\s]{5,}$")

        # Extrair títulos (headings + padrões inteligentes)
        for para in doc.paragraphs:
            estilo = _nome_estilo_paragrafo(para)
            texto = para.text.strip()
            if not texto:
                continue

            nivel = _heading_nivel(estilo)
            confianca = "alta"

            # Detecção por estilo Heading
            if nivel is not None:
                sugestoes["titulos"].append(
                    {
                        "texto": texto,
                        "nivel": nivel,
                        "estilo": estilo,
                        "confianca": confianca,
                    }
                )
                continue

            # Detecção por padrões de formatação
            # 1. Numeração no início (1., 1.1, 2., etc.)
            if padrao_numeracao.match(texto):
                # Inferir nível pela profundidade da numeração
                partes = texto.split(".")[0].split()
                if partes:
                    try:
                        nivel_inferido = len(partes[0].split("."))
                    except Exception:  # noqa: E722
                        nivel_inferido = 1
                else:
                    nivel_inferido = 1
                sugestoes["titulos"].append(
                    {
                        "texto": texto,
                        "nivel": nivel_inferido,
                        "estilo": estilo,
                        "confianca": "media",
                    }
                )
                continue

            # 2. Texto em caixa alta (sugere título)
            if padrao_caixa_alta.match(texto) and len(texto) < 100:
                sugestoes["titulos"].append(
                    {
                        "texto": texto,
                        "nivel": 1,
                        "estilo": estilo,
                        "confianca": "baixa",
                    }
                )
                continue

            # 3. Texto em negrito e tamanho maior que o normal
            if para.runs:
                tem_negrito = any(run.bold for run in para.runs if run.bold)
                if tem_negrito and len(texto) < 80:
                    sugestoes["titulos"].append(
                        {
                            "texto": texto,
                            "nivel": 1,
                            "estilo": estilo,
                            "confianca": "baixa",
                        }
                    )

        # Extrair figuras (imagens inline e flutuantes)
        # python-docx não detecta facilmente imagens flutuantes,
        # então focamos em imagens inline em parágrafos
        for para in doc.paragraphs:
            if _xml_xpath(para, ".//pic:pic"):
                # Parágrafo contém imagem
                texto_legenda = para.text.strip()
                if texto_legenda:
                    sugestoes["figuras"].append(
                        {
                            "legenda": texto_legenda,
                            "tipo": "inline",
                            "tem_legenda": True,
                        }
                    )
                else:
                    # Imagem sem legenda - sugerir adicionar
                    sugestoes["figuras"].append(
                        {  # noqa: E501
                            "legenda": None,
                            "tipo": "inline",
                            "tem_legenda": False,
                            "sugestao": ("Adicione uma legenda descritiva " "para esta figura."),
                        }
                    )
            else:
                # Detectar parágrafos que mencionam figuras
                texto_lower = para.text.lower()
                if any(
                    palavra in texto_lower
                    for palavra in ["figura", "fig.", "imagem", "img."]  # noqa: E501
                ):
                    sugestoes["figuras"].append(
                        {
                            "legenda": para.text.strip(),
                            "tipo": "referencia_texto",
                            "tem_legenda": True,
                        }
                    )

        # Extrair tabelas e suas legendas
        for i, table in enumerate(doc.tables):
            # Tenta encontrar legenda no parágrafo anterior à tabela
            # ou no primeiro parágrafo após a tabela
            legenda = None
            # Busca parágrafo anterior
            for para in doc.paragraphs:
                if _elemento_xml(table) in _xml_xpath(
                    para,
                    "following-sibling::w:p",
                ):
                    if para.text.strip():
                        legenda = para.text.strip()
                        break
            if not legenda:
                # Busca parágrafo seguinte
                for para in doc.paragraphs:
                    if _elemento_xml(table) in _xml_xpath(
                        para,
                        "preceding-sibling::w:p",
                    ):
                        if para.text.strip():
                            legenda = para.text.strip()
                            break

            tabela_info = {
                "indice": i + 1,
                "linhas": len(table.rows),
                "colunas": len(table.columns) if table.rows else 0,
                "legenda": legenda,
                "tem_legenda": legenda is not None,
            }

            if not legenda:
                tabela_info["sugestao"] = (  # noqa: E501
                    "Adicione uma legenda descritiva " "para esta tabela."
                )

            sugestoes["tabelas"].append(tabela_info)

        # Detectar referências a tabelas no texto
        for para in doc.paragraphs:
            texto_lower = para.text.lower()
            if any(
                palavra in texto_lower for palavra in ["tabela", "tab.", "quadro"]  # noqa: E501
            ):
                # Verificar se não é uma tabela já detectada
                if not any(  # noqa: E501
                    t.get("legenda") == para.text.strip() for t in sugestoes["tabelas"]
                ):
                    sugestoes["tabelas"].append(
                        {
                            "indice": len(sugestoes["tabelas"]) + 1,
                            "linhas": 0,
                            "colunas": 0,
                            "legenda": para.text.strip(),
                            "tem_legenda": True,
                            "tipo": "referencia_texto",
                        }
                    )

        return sugestoes
