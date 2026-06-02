"""Inseridor de Sumario, Lista de Figuras e Lista de Tabelas.

Filosofia (revisada):
- O sistema PRE-CALCULA todo o conteudo no servidor. NADA depende de
  "Word recalcular ao abrir".
- O coordenador aciona explicitamente 3 botoes na UI; cada um chama
  uma das funcoes expostas aqui.
- Os blocos sao posicionados na regiao PRE-TEXTUAL do DOCX em
  producao (antes do primeiro Heading 1 do corpo).
- Re-execucao e idempotente: cada bloco e delimitado por bookmarks
  marcadores `_Sra_Bloco_*_Inicio`/`_Fim`. Antes de inserir, a versao
  anterior (se existir) e removida.
- Hyperlinks reais: cada entrada do sumario/lista aponta para o
  bookmark do destino (heading ou legenda).

Estrutura gerada (exemplo, sumario):

    [titulo: "Sumário" com estilo TOC Heading]
    [bookmarkStart name="_Sra_Bloco_Sumario_Inicio"]
    [paragrafo estilo "toc 1" com hyperlink para bookmark do heading 1]
    [paragrafo estilo "toc 2" com hyperlink para bookmark do heading 2]
    ...
    [bookmarkEnd para _Sra_Bloco_Sumario_Inicio]

Numeros de pagina NAO sao incluidos (impossivel calcular sem renderizar
o documento). Listas servem como navegacao com hyperlinks.

API:
    inserir_sumario(caminho, perfil) -> dict
    inserir_lista_figuras(caminho, perfil) -> dict
    inserir_lista_tabelas(caminho, perfil) -> dict
"""
from __future__ import annotations

import re
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree

from app.services.servico_perfil_formatacao import PerfilFormatacao
from app.services._ooxml_helpers import (
    GeradorIdsBookmark,
    aplicar_estilo_paragrafo,
    criar_bookmark_par,
    criar_run_texto,
    texto_paragrafo as _texto_paragrafo,
)

_element = getattr(etree, 'Element')
_sub_element = getattr(etree, 'SubElement')


W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SPACE_PRESERVE = (
    '{http://www.w3.org/XML/1998/namespace}space',
    'preserve',
)

# Marcadores idempotentes dos blocos que inserimos. Cada bloco tem
# `_Inicio` e `_Fim` para permitir remocao precisa.
MARCADOR_SUMARIO_INI = '_Sra_Bloco_Sumario_Inicio'
MARCADOR_SUMARIO_FIM = '_Sra_Bloco_Sumario_Fim'
MARCADOR_LISTA_FIG_INI = '_Sra_Bloco_ListaFiguras_Inicio'
MARCADOR_LISTA_FIG_FIM = '_Sra_Bloco_ListaFiguras_Fim'
MARCADOR_LISTA_TAB_INI = '_Sra_Bloco_ListaTabelas_Inicio'
MARCADOR_LISTA_TAB_FIM = '_Sra_Bloco_ListaTabelas_Fim'
MARCADOR_LISTA_EQ_INI = '_Sra_Bloco_ListaEquacoes_Inicio'
MARCADOR_LISTA_EQ_FIM = '_Sra_Bloco_ListaEquacoes_Fim'
MARCADOR_LISTA_SIG_INI = '_Sra_Bloco_ListaSiglas_Inicio'
MARCADOR_LISTA_SIG_FIM = '_Sra_Bloco_ListaSiglas_Fim'

# Prefixo dos bookmarks gerados em headings para servir de target
# de hyperlink no sumario.
PREFIXO_BOOKMARK_HEADING = '_Toc_sra_'


# =====================================================================
# Detecao de elementos no DOCX
# =====================================================================


def _nivel_heading(p_element) -> Optional[int]:
    """Retorna o nivel (1..9) se for heading, senao None.

    Reconhece styleIds com 'Heading'/'eading' (Word en-US),
    'Titulo'/'tulo' (Word pt-BR localizado) ou apenas digito.
    """
    p_pr = p_element.find(qn('w:pPr'))
    if p_pr is None:
        return None
    p_style = p_pr.find(qn('w:pStyle'))
    if p_style is None:
        return None
    val = p_style.get(qn('w:val'), '')
    m = re.search(r'(\d)', val)
    if not m:
        return None
    if not ('eading' in val
            or 'tulo' in val.lower()
            or val.isdigit()):
        return None
    n = int(m.group(1))
    if 1 <= n <= 9:
        return n
    return None


_TODOS_MARCADORES_INI = {
    MARCADOR_SUMARIO_INI,
    MARCADOR_LISTA_FIG_INI,
    MARCADOR_LISTA_TAB_INI,
    MARCADOR_LISTA_EQ_INI,
    MARCADOR_LISTA_SIG_INI,
}
_TODOS_MARCADORES_FIM = {
    MARCADOR_SUMARIO_FIM,
    MARCADOR_LISTA_FIG_FIM,
    MARCADOR_LISTA_TAB_FIM,
    MARCADOR_LISTA_EQ_FIM,
    MARCADOR_LISTA_SIG_FIM,
}


# =====================================================================
# Ordem ABNT NBR 14724 / NBR 10719 dos pre-textuais.
#
# Sequencia canonica (apos capa e folha de rosto):
#   1. Lista de Figuras (ou de Ilustracoes)
#   2. Lista de Tabelas
#   3. Lista de Equacoes
#   4. Lista de Siglas / Abreviaturas
#   5. Sumario  ← SEMPRE o ULTIMO pre-textual, logo antes do conteudo
#
# Cada bloco e posicionado APOS o `_Fim` do bloco precedente mais
# proximo que ja exista, e ANTES do `_Inicio` do bloco sucessor mais
# proximo. Se nada existir ainda, fallback = primeiro heading do
# corpo (antes do conteudo textual).
# =====================================================================

_PRIORIDADE_PARA_PRECEDENTES = {
    # prio 1 = Lista de Figuras: primeira da serie, sem precedentes
    # entre as listas. (Capa e Folha de Rosto vivem no inicio do
    # documento, antes da regiao gerenciada por este servico.)
    1: [],
    # prio 2 = Lista de Tabelas: apos Lista de Figuras
    2: [MARCADOR_LISTA_FIG_FIM],
    # prio 3 = Lista de Equacoes: apos Tabelas, ou Figuras
    3: [MARCADOR_LISTA_TAB_FIM, MARCADOR_LISTA_FIG_FIM],
    # prio 4 = Lista de Siglas/Abreviaturas: apos Equacoes
    4: [
        MARCADOR_LISTA_EQ_FIM,
        MARCADOR_LISTA_TAB_FIM,
        MARCADOR_LISTA_FIG_FIM,
    ],
    # prio 5 = Sumario: ULTIMO, apos todas as listas (NBR 14724 6.2.10)
    5: [
        MARCADOR_LISTA_SIG_FIM,
        MARCADOR_LISTA_EQ_FIM,
        MARCADOR_LISTA_TAB_FIM,
        MARCADOR_LISTA_FIG_FIM,
    ],
}

# Mapeamento prio -> marcador `_Inicio` (usado p/ achar sucessor)
_PRIORIDADE_PARA_MARCADOR_INI = {
    1: MARCADOR_LISTA_FIG_INI,
    2: MARCADOR_LISTA_TAB_INI,
    3: MARCADOR_LISTA_EQ_INI,
    4: MARCADOR_LISTA_SIG_INI,
    5: MARCADOR_SUMARIO_INI,
}


def _calcular_posicao_insercao(body, *, prioridade: int) -> int:
    """Calcula posicao correta para inserir um bloco respeitando a
    ordem ABNT NBR 14724 / NBR 10719 dos pre-textuais:

        Lista Figuras (1) -> Lista Tabelas (2) -> Lista Equacoes (3)
        -> Lista Siglas (4) -> Sumario (5) -> Conteudo textual

    Estrategia:
    1. Se ha bloco precedente (segundo a ordem em
       `_PRIORIDADE_PARA_PRECEDENTES`) ja inserido: posicao = APOS
       o `_Fim` desse bloco (o mais proximo encontrado).
    2. Caso contrario:
       a. Posicao = bloco SUCESSOR mais proximo (que deve vir DEPOIS).
       b. Se nao houver sucessor, posicao = primeiro heading do corpo.
    """
    bms_tag = qn('w:bookmarkStart')
    name_attr = qn('w:name')

    # 1. Procurar precedente mais proximo (primeiro da lista que
    # estiver presente no documento). A ordem em
    # _PRIORIDADE_PARA_PRECEDENTES ja prioriza o mais proximo.
    precedentes = _PRIORIDADE_PARA_PRECEDENTES.get(prioridade, [])
    for marcador_fim in precedentes:
        for i, child in enumerate(body):
            if child.tag != qn('w:p'):
                continue
            for bm in child.iter(bms_tag):
                if bm.get(name_attr) == marcador_fim:
                    return i + 1

    # 2. Procurar sucessor mais proximo: bloco com prioridade MAIOR
    # que a atual e que ja exista no documento. Posicao = onde
    # comeca o sucessor (devemos inserir ANTES dele).
    sucessores = [
        _PRIORIDADE_PARA_MARCADOR_INI[p]
        for p in sorted(_PRIORIDADE_PARA_MARCADOR_INI)
        if p > prioridade
    ]
    pos_sucessor = None
    for marcador_ini in sucessores:
        for i, child in enumerate(body):
            if child.tag != qn('w:p'):
                continue
            for bm in child.iter(bms_tag):
                if bm.get(name_attr) == marcador_ini:
                    if pos_sucessor is None or i < pos_sucessor:
                        pos_sucessor = i
    if pos_sucessor is not None:
        return pos_sucessor

    # 3. Fallback: primeiro heading do corpo (antes do conteudo)
    return _achar_indice_primeiro_heading_corpo(body)


def _achar_indice_primeiro_heading_corpo(body) -> int:
    """Retorna o indice do primeiro heading que NAO esteja DENTRO de
    nenhum bloco `_Sra_Bloco_*_Inicio`/`_Fim`.

    Isso garante que insercoes subsequentes (lista de figuras, lista
    de tabelas) sejam posicionadas APOS blocos previamente inseridos
    (ex: sumario), preservando a ordem natural:

        Sumario -> Lista de Figuras -> Lista de Tabelas -> Conteudo

    Retorna `len(body)` se nao houver heading de corpo.
    """
    bms_tag = qn('w:bookmarkStart')
    name_attr = qn('w:name')

    dentro_de_bloco = False
    for i, child in enumerate(body):
        if child.tag != qn('w:p'):
            continue
        # Detectar entrada/saida de bloco via bookmarks marcadores
        for bm in child.iter(bms_tag):
            nome = bm.get(name_attr)
            if nome in _TODOS_MARCADORES_INI:
                dentro_de_bloco = True
            elif nome in _TODOS_MARCADORES_FIM:
                dentro_de_bloco = False
        if dentro_de_bloco:
            continue
        if _nivel_heading(child) is not None:
            return i
    return len(body)


# =====================================================================
# Idempotencia: remover blocos antigos
# =====================================================================


def _remover_bloco_marcado(
    body, marcador_inicio: str, marcador_fim: str
) -> int:
    """Remove TODOS os elementos entre `<w:bookmarkStart name=marcador_inicio>`
    e o `<w:bookmarkEnd>` correspondente ao marcador `_Fim`.

    Estrategia:
    1. Achar bookmarkStart com name=marcador_inicio. Captura sua POSICAO
       no body (subindo na arvore ate encontrar o filho direto).
    2. Achar bookmarkStart com name=marcador_fim. Captura POSICAO.
    3. Remover todos os elementos entre as posicoes (inclusive os
       paragrafos que CONTEM os bookmarks marcadores).

    Retorna a quantidade de elementos removidos. Zero se o bloco nao
    existir (caso de primeira insercao).
    """
    bms_tag = qn('w:bookmarkStart')
    name_attr = qn('w:name')

    pos_inicio = None
    pos_fim = None
    for bm in body.iter(bms_tag):
        nome = bm.get(name_attr)
        # Subir ate encontrar filho direto do body
        ancestor = bm
        while ancestor is not None and ancestor.getparent() is not body:
            ancestor = ancestor.getparent()
        if ancestor is None:
            continue
        try:
            idx = list(body).index(ancestor)
        except ValueError:
            continue
        if nome == marcador_inicio and pos_inicio is None:
            pos_inicio = idx
        elif nome == marcador_fim and pos_fim is None:
            pos_fim = idx

    if pos_inicio is None or pos_fim is None:
        return 0
    if pos_fim < pos_inicio:
        return 0

    # Capturar referencias antes de remover (lxml nao gosta de
    # iteracao com remocao concorrente).
    a_remover = list(body)[pos_inicio:pos_fim + 1]
    for elem in a_remover:
        body.remove(elem)
    return len(a_remover)


# =====================================================================
# Garantir bookmarks `_Toc_sra_*` nos headings
# =====================================================================


def _heading_ja_tem_bookmark_toc(p_element) -> Optional[str]:
    """Se o heading ja contem bookmarkStart com prefixo `_Toc_sra_`,
    retorna o nome. Senao, None.
    """
    for bm in p_element.iter(qn('w:bookmarkStart')):
        nome = bm.get(qn('w:name')) or ''
        if nome.startswith(PREFIXO_BOOKMARK_HEADING):
            return nome
    return None


def _garantir_bookmarks_em_headings(
    body, id_gen: GeradorIdsBookmark
) -> list[dict]:
    """Percorre headings e garante que cada um tenha um bookmark
    `_Toc_sra_h<N>_n<seq>` envolvendo seu conteudo. Se ja tem, mantem.

    IGNORA paragrafos dentro de blocos `_Sra_Bloco_*_Inicio`/`_Fim`
    (titulos de Lista de Figuras/Tabelas que usam estilo Heading 1
    fallback nao devem aparecer como entradas do sumario).

    Retorna lista de dicts ordenados:
        [{'nivel': 1, 'texto': 'APRESENTACAO',
          'bookmark': '_Toc_sra_h1_n1', 'p_element': <ref>}]
    """
    entradas = []
    contadores = {}  # {nivel: seq}
    bms_tag = qn('w:bookmarkStart')
    name_attr = qn('w:name')

    dentro_de_bloco = False
    for child in list(body):
        if child.tag != qn('w:p'):
            continue
        # Detectar entrada/saida de bloco _Sra_Bloco_*
        for bm in child.iter(bms_tag):
            nome = bm.get(name_attr)
            if nome in _TODOS_MARCADORES_INI:
                dentro_de_bloco = True
            elif nome in _TODOS_MARCADORES_FIM:
                dentro_de_bloco = False
        if dentro_de_bloco:
            continue
        nivel = _nivel_heading(child)
        if nivel is None:
            continue

        # Texto puro do heading (sem campos/numeracao automatica)
        texto = _texto_paragrafo(child).strip()
        if not texto:
            continue

        nome_existente = _heading_ja_tem_bookmark_toc(child)
        if nome_existente:
            nome_bm = nome_existente
        else:
            contadores.setdefault(nivel, 0)
            contadores[nivel] += 1
            seq = contadores[nivel]
            nome_bm = f'{PREFIXO_BOOKMARK_HEADING}h{nivel}_n{seq}'
            # Inserir bookmarkStart no comeco do paragrafo (apos p_pr)
            # e bookmarkEnd no fim.
            bms, bme = criar_bookmark_par(nome_bm, id_gen)
            p_pr = child.find(qn('w:pPr'))
            if p_pr is not None:
                p_pr.addnext(bms)
            else:
                child.insert(0, bms)
            child.append(bme)

        entradas.append({
            'nivel': nivel,
            'texto': texto,
            'bookmark': nome_bm,
        })
    return entradas


def _remover_bookmarks_toc_sra(body) -> int:
    """Remove todos os bookmarks com prefixo `_Toc_sra_` (preparacao
    para regenerar — garante que IDs nao colidem entre execucoes).
    """
    bms_tag = qn('w:bookmarkStart')
    bme_tag = qn('w:bookmarkEnd')
    name_attr = qn('w:name')
    id_attr = qn('w:id')

    ids_a_remover: set = set()
    elementos: list = []
    for bm in list(body.iter(bms_tag)):
        nome = bm.get(name_attr) or ''
        if nome.startswith(PREFIXO_BOOKMARK_HEADING):
            bm_id = bm.get(id_attr)
            if bm_id:
                ids_a_remover.add(bm_id)
            elementos.append(bm)
    for bme in list(body.iter(bme_tag)):
        if bme.get(id_attr) in ids_a_remover:
            elementos.append(bme)

    removidos = 0
    for elem in elementos:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            removidos += 1
    return removidos


# =====================================================================
# Construcao de paragrafos do TOC/Listas
# =====================================================================


def _criar_paragrafo_titulo(texto: str, estilo: Optional[str]):
    """Cria <w:p> com estilo de titulo (default Heading 1)."""
    w = f'{{{W_NS}}}'
    p = _element(f'{w}p')
    aplicar_estilo_paragrafo(p, estilo)
    p.append(criar_run_texto(texto))
    return p


def _criar_paragrafo_entrada_lista(
    *,
    texto: str,
    bookmark_destino: Optional[str],
    estilo_paragrafo: Optional[str],
):
    """Cria um `<w:p>` representando uma entrada de lista (figura/
    tabela/equacao/sigla).

    Se `bookmark_destino` for nao-vazio, envolve o texto em
    `<w:hyperlink w:anchor="...">` (Word abre como link clicavel para
    o bookmark). Se for `None` ou vazio, gera entrada sem hyperlink
    (caso de equacoes inline e siglas, que nao tem bookmark).

    Estrutura com hyperlink:
        <w:p>
          <w:pPr><w:pStyle w:val="toc 1"/></w:pPr>
          <w:hyperlink w:anchor="<bookmark>" w:history="1">
            <w:r><w:t>texto</w:t></w:r>
          </w:hyperlink>
        </w:p>

    Estrutura sem hyperlink:
        <w:p>
          <w:pPr><w:pStyle w:val="toc 1"/></w:pPr>
          <w:r><w:t>texto</w:t></w:r>
        </w:p>
    """
    w = f'{{{W_NS}}}'
    p = _element(f'{w}p')
    aplicar_estilo_paragrafo(p, estilo_paragrafo)

    if bookmark_destino:
        container = _sub_element(p, f'{w}hyperlink')
        container.set(f'{w}anchor', bookmark_destino)
        container.set(f'{w}history', '1')
    else:
        container = p

    r = _sub_element(container, f'{w}r')
    t = _sub_element(r, f'{w}t')
    t.set(*XML_SPACE_PRESERVE)
    t.text = texto

    return p


def _criar_paragrafo_marcador(nome_marcador: str, id_gen):
    """Cria um <w:p> minimo contendo apenas um bookmarkStart marcador
    (sem run de texto). Usado como sentinela de inicio/fim de bloco.
    """
    w = f'{{{W_NS}}}'
    p = _element(f'{w}p')
    bms, bme = criar_bookmark_par(nome_marcador, id_gen)
    p.append(bms)
    p.append(bme)
    return p


# =====================================================================
# Coleta de legendas (figuras/tabelas) ja inseridas pelo captioning
# =====================================================================

# Bookmark inicia com isso quando vem do servico_captioning
PREFIXO_BOOKMARK_LEGENDA = '_Ref_sra_'


def _coletar_legendas_por_tipo(body, prefixo_tipo: str) -> list:
    """Coleta legendas marcadas com bookmark `_Ref_sra_<prefixo>_*`.

    `prefixo_tipo` e 'fig', 'tab' ou 'eq'. Retorna lista de dicts:
        [{'texto': 'Figura 1.1: Sala', 'bookmark': '_Ref_sra_fig_x'}, ...]
    em ordem de aparicao no documento.
    """
    bms_tag = qn('w:bookmarkStart')
    name_attr = qn('w:name')
    prefixo = f'{PREFIXO_BOOKMARK_LEGENDA}{prefixo_tipo}_'

    entradas = []
    for bm in body.iter(bms_tag):
        nome = bm.get(name_attr) or ''
        if not nome.startswith(prefixo):
            continue
        # Subir ate encontrar paragrafo
        p = bm
        while p is not None and p.tag != qn('w:p'):
            p = p.getparent()
        if p is None:
            continue
        texto = _texto_paragrafo(p).strip()
        if texto:
            entradas.append({'texto': texto, 'bookmark': nome})
    return entradas


# =====================================================================
# API publica
# =====================================================================


def _resolver_estilo(
    doc, estilo_desejado: str, fallback: Optional[str]
) -> Optional[str]:
    """Verifica se `estilo_desejado` existe em styles.xml. Senao,
    retorna `fallback`.
    """
    try:
        styles_part = doc.styles.element
    except AttributeError:
        return fallback
    nomes: set = set()
    for style in styles_part.findall(qn('w:style')):
        sid = style.get(qn('w:styleId'))
        if sid:
            nomes.add(sid)
        name_el = style.find(qn('w:name'))
        if name_el is not None:
            nm = name_el.get(qn('w:val'))
            if nm:
                nomes.add(nm)
    return estilo_desejado if estilo_desejado in nomes else fallback


def _resolver_perfil(perfil):
    if perfil is not None:
        return perfil
    return PerfilFormatacao()


def _inserir_bloco(
    body,
    pos: int,
    paragrafos: list,
    marcador_inicio: str,
    marcador_fim: str,
    id_gen: GeradorIdsBookmark,
):
    """Insere `paragrafos` em `body[pos]`, envolvendo-os com
    paragrafos-sentinela contendo bookmarks `marcador_inicio` e
    `marcador_fim`.

    Modifica body in-place. Retorna a posicao IMEDIATAMENTE APOS o
    bloco inserido.
    """
    p_inicio = _criar_paragrafo_marcador(marcador_inicio, id_gen)
    p_fim = _criar_paragrafo_marcador(marcador_fim, id_gen)
    bloco = [p_inicio] + list(paragrafos) + [p_fim]
    for offset, elem in enumerate(bloco):
        body.insert(pos + offset, elem)
    return pos + len(bloco)


def inserir_sumario(caminho_master: str, perfil=None) -> dict:
    """Coleta todos os headings do DOCX e insere um Sumario PRE-PREENCHIDO
    na regiao pre-textual.

    - Adiciona bookmarks `_Toc_sra_*` em todos os headings (idempotente)
    - Gera lista de paragrafos com estilo `toc 1`/`toc 2`/`toc 3` e
      hyperlink interno para o bookmark de cada heading
    - Posiciona antes do primeiro Heading 1 do corpo
    - Idempotente: remove bloco anterior se existir

    Salva em `caminho_master`. Retorna contadores.
    """
    perfil = _resolver_perfil(perfil)
    doc = Document(caminho_master)
    body = doc.element.body

    estilo_titulo = _resolver_estilo(
        doc, perfil.estilo_titulo_toc, fallback='Heading 1',
    )

    id_gen = GeradorIdsBookmark(inicio=200000)

    # 1. Limpeza idempotente
    _remover_bloco_marcado(
        body, MARCADOR_SUMARIO_INI, MARCADOR_SUMARIO_FIM,
    )
    _remover_bookmarks_toc_sra(body)

    # 2. Garantir bookmarks em todos os headings
    entradas = _garantir_bookmarks_em_headings(body, id_gen)

    # 3. Construir paragrafos do sumario
    paragrafos = [_criar_paragrafo_titulo('Sumário', estilo_titulo)]
    for entrada in entradas:
        nivel = entrada['nivel']
        # Estilo "toc 1", "toc 2"... (Word usa minusculas com espaco)
        estilo_entrada = f'toc {nivel}' if 1 <= nivel <= 9 else 'toc 1'
        # Fallback caso o estilo nao exista no DOCX
        estilo_entrada = _resolver_estilo(
            doc, estilo_entrada, fallback=None,
        )
        paragrafos.append(_criar_paragrafo_entrada_lista(
            texto=entrada['texto'],
            bookmark_destino=entrada['bookmark'],
            estilo_paragrafo=estilo_entrada,
        ))

    # 4. Inserir na pre-textual.
    # Sumario e SEMPRE o ULTIMO pre-textual (ABNT NBR 14724 6.2.10),
    # logo antes do primeiro Heading 1 do corpo. Prioridade 5 garante
    # posicionamento apos todas as listas (figuras, tabelas, equacoes,
    # siglas) se existirem.
    pos = _calcular_posicao_insercao(body, prioridade=5)
    _inserir_bloco(
        body, pos, paragrafos,
        MARCADOR_SUMARIO_INI, MARCADOR_SUMARIO_FIM, id_gen,
    )

    doc.save(caminho_master)
    return {
        'entradas': len(entradas),
        'estilo_titulo': estilo_titulo,
        'posicao_inserido': pos,
        'perfil_origem': perfil.origem,
    }


def _inserir_lista_legendas(
    caminho_master: str,
    *,
    prefixo_tipo: str,         # 'fig' ou 'tab'
    titulo: str,               # "Lista de Figuras"
    marcador_ini: str,
    marcador_fim: str,
    prioridade: int,           # 2 = Figuras, 3 = Tabelas
    perfil,
) -> dict:
    """Implementacao comum para Lista de Figuras e Lista de Tabelas.

    Coleta legendas via bookmark `_Ref_sra_<prefixo>_*` (gerados pelo
    `servico_captioning`) e monta lista pre-preenchida.
    """
    perfil = _resolver_perfil(perfil)
    doc = Document(caminho_master)
    body = doc.element.body

    estilo_titulo = _resolver_estilo(
        doc, perfil.estilo_titulo_toc, fallback='Heading 1',
    )
    estilo_entrada = _resolver_estilo(
        doc, 'table of figures', fallback=None,
    )

    id_gen = GeradorIdsBookmark(inicio=210000)

    # 1. Limpeza idempotente
    _remover_bloco_marcado(body, marcador_ini, marcador_fim)

    # 2. Coletar legendas
    legendas = _coletar_legendas_por_tipo(body, prefixo_tipo)

    # 3. Construir paragrafos
    paragrafos = [_criar_paragrafo_titulo(titulo, estilo_titulo)]
    for leg in legendas:
        paragrafos.append(_criar_paragrafo_entrada_lista(
            texto=leg['texto'],
            bookmark_destino=leg['bookmark'],
            estilo_paragrafo=estilo_entrada,
        ))

    # 4. Inserir na pre-textual respeitando ordem (sumario < figs < tabs)
    pos = _calcular_posicao_insercao(body, prioridade=prioridade)
    _inserir_bloco(
        body, pos, paragrafos,
        marcador_ini, marcador_fim, id_gen,
    )

    doc.save(caminho_master)
    return {
        'entradas': len(legendas),
        'estilo_titulo': estilo_titulo,
        'posicao_inserido': pos,
        'perfil_origem': perfil.origem,
    }


def inserir_lista_figuras(caminho_master: str, perfil=None) -> dict:
    """Insere Lista de Figuras pre-preenchida na pre-textual.

    Posicao ABNT NBR 14724: primeira das listas (apos folha de rosto
    e antes de Tabelas, Equacoes, Siglas e Sumario).

    Pre-condicao: o `servico_captioning.reindexar_captions` ja foi
    executado (legendas tem bookmarks `_Ref_sra_fig_*`). Caso contrario,
    a lista vai vazia.
    """
    return _inserir_lista_legendas(
        caminho_master,
        prefixo_tipo='fig',
        titulo='Lista de Figuras',
        marcador_ini=MARCADOR_LISTA_FIG_INI,
        marcador_fim=MARCADOR_LISTA_FIG_FIM,
        prioridade=1,
        perfil=perfil,
    )


def inserir_lista_tabelas(caminho_master: str, perfil=None) -> dict:
    """Insere Lista de Tabelas pre-preenchida na pre-textual.

    Posicao ABNT: apos Lista de Figuras, antes de Equacoes/Siglas/Sumario.
    """
    return _inserir_lista_legendas(
        caminho_master,
        prefixo_tipo='tab',
        titulo='Lista de Tabelas',
        marcador_ini=MARCADOR_LISTA_TAB_INI,
        marcador_fim=MARCADOR_LISTA_TAB_FIM,
        prioridade=2,
        perfil=perfil,
    )


def inserir_lista_equacoes(caminho_master: str, perfil=None) -> dict:
    """Insere Lista de Equacoes pre-preenchida na pre-textual.

    Posicao ABNT: apos Tabelas, antes de Siglas/Sumario.

    Diferente de Figuras/Tabelas, equacoes nao tem bookmark associado
    (a numeracao e inline no fim do paragrafo da equacao — ver
    `servico_captioning._anexar_numero_inline_equacao`). Esta lista
    e construida varrendo o body por paragrafos que contem
    `<m:oMath>` / `<m:oMathPara>` e extraindo o numero inline `(N.M)`.
    As entradas sao informativas (sem hyperlink), no formato:
        Equacao N.M  ........  (Capitulo X)
    """
    perfil = _resolver_perfil(perfil)
    doc = Document(caminho_master)
    body = doc.element.body

    estilo_titulo = _resolver_estilo(
        doc, perfil.estilo_titulo_toc, fallback='Heading 1',
    )
    estilo_entrada = _resolver_estilo(
        doc, 'table of figures', fallback=None,
    )

    id_gen = GeradorIdsBookmark(inicio=230000)

    # 1. Limpeza idempotente
    _remover_bloco_marcado(
        body, MARCADOR_LISTA_EQ_INI, MARCADOR_LISTA_EQ_FIM,
    )

    # 2. Coletar equacoes inline. Reutiliza heuristica do captioning.
    m_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    entradas = []
    for p in body.iter(qn('w:p')):
        if (p.find(f'.//{{{m_ns}}}oMath') is None
                and p.find(f'.//{{{m_ns}}}oMathPara') is None):
            continue
        # Texto do paragrafo: ultimo token entre parenteses
        # tipicamente e a numeracao "(N.M)" gerada pelo captioning.
        texto = _texto_paragrafo(p).strip()
        m = re.search(r'\(([\d.]+)\)\s*$', texto)
        numero = m.group(1) if m else '?'
        entradas.append({
            'texto': f'Equação {numero}',
        })

    # 3. Construir paragrafos
    paragrafos = [_criar_paragrafo_titulo(
        'Lista de Equações', estilo_titulo,
    )]
    for ent in entradas:
        paragrafos.append(_criar_paragrafo_entrada_lista(
            texto=ent['texto'],
            bookmark_destino=None,
            estilo_paragrafo=estilo_entrada,
        ))

    # 4. Inserir respeitando ordem ABNT
    pos = _calcular_posicao_insercao(body, prioridade=3)
    _inserir_bloco(
        body, pos, paragrafos,
        MARCADOR_LISTA_EQ_INI, MARCADOR_LISTA_EQ_FIM, id_gen,
    )

    doc.save(caminho_master)
    return {
        'entradas': len(entradas),
        'estilo_titulo': estilo_titulo,
        'posicao_inserido': pos,
        'perfil_origem': perfil.origem,
    }


def inserir_lista_siglas(caminho_master: str, perfil=None) -> dict:
    """Insere Lista de Siglas e Abreviaturas pre-textual.

    Posicao ABNT NBR 14724 5.10: penultimo pre-textual, logo antes
    do Sumario.

    Diferente das demais, a Lista de Siglas e construida a partir de
    SIGLAS DETECTADAS automaticamente no texto do documento:
        - Sequencias de 2-6 letras maiusculas (com possiveis digitos).
        - Filtradas para excluir cabecalhos de tabela, palavras curtas
          em CAIXA ALTA por estilo, e siglas ja conhecidas.
    Para cada sigla, gera-se uma entrada placeholder com texto:
        "SIGLA  —  (descricao a preencher)"
    O coordenador edita as descricoes manualmente no DOCX.

    Esta abordagem e didatica e segura: nao tenta inferir significados
    (o que seria propenso a erros); apenas garante que TODAS as siglas
    usadas no texto apareçam na lista para preenchimento manual.
    """
    perfil = _resolver_perfil(perfil)
    doc = Document(caminho_master)
    body = doc.element.body

    estilo_titulo = _resolver_estilo(
        doc, perfil.estilo_titulo_toc, fallback='Heading 1',
    )
    estilo_entrada = _resolver_estilo(
        doc, 'table of figures', fallback=None,
    )

    id_gen = GeradorIdsBookmark(inicio=240000)

    # 1. Limpeza idempotente PRIMEIRO (caso contrario o regex abaixo
    # captaria os placeholders da rodada anterior como "siglas novas")
    _remover_bloco_marcado(
        body, MARCADOR_LISTA_SIG_INI, MARCADOR_LISTA_SIG_FIM,
    )

    # 2. Detectar siglas no texto. Heuristica conservadora:
    #    - 2 a 6 letras maiusculas
    #    - Pode conter digitos no meio (ex: D20, PLI-SP, NBR14724)
    #    - Nao pode ser uma palavra inteira (ex: "BRASIL" e nome,
    #      nao sigla); para evitar isso, exigimos pelo menos uma
    #      transicao MAIUSCULA->minuscula no contexto, OU 2+ letras
    #      seguidas de digito/hifen.
    # Padrao simples robusto:
    #   - Palavra com >=2 letras maiusculas
    #   - Aceita hifen, digito ou ponto no meio (PLI-SP, D20, NBR-14724)
    padrao_sigla = re.compile(
        r'\b([A-ZÀ-Ý]{2,}(?:[-./][A-ZÀ-Ý0-9]+)*\d*)\b'
    )
    # Stopwords: palavras inteiras em maiusculas que NAO sao siglas
    # (titulos de capitulos, vocativos, etc.). Como heuristica final,
    # tambem ignoramos sequencias com >=7 letras (provavel palavra
    # em CAIXA ALTA estilistica).
    stopwords = {
        'APRESENTAÇÃO', 'HISTÓRICO', 'CONTRATO', 'RELAÇÃO',
        'PRODUTOS', 'VISÃO', 'GERAL', 'ATIVIDADES', 'EQUIPE',
        'APOIO', 'GESTÃO', 'RECURSOS', 'CRONOGRAMA', 'ANÁLISE',
        'PRÓXIMOS', 'PASSOS', 'PRODUTO', 'RESUMO', 'MEDIÇÃO',
        'ASSINATURAS', 'APÊNDICE', 'ANEXO', 'OBJETIVO', 'OBJETIVOS',
        'CONCLUSÃO', 'INTRODUÇÃO', 'METODOLOGIA', 'REFERÊNCIAS',
        'SUMÁRIO', 'CAPA',
    }

    siglas_unicas = {}  # sigla -> primeira ocorrencia (para ordenar)
    for i, p in enumerate(body.iter(qn('w:p'))):
        texto = _texto_paragrafo(p)
        for m in padrao_sigla.finditer(texto):
            sig = m.group(1)
            # Filtros adicionais
            if len(sig) > 6 and '-' not in sig and '.' not in sig:
                # Provavelmente palavra em CAIXA ALTA (titulo)
                continue
            if sig in stopwords:
                continue
            if sig.isdigit():
                continue
            if sig not in siglas_unicas:
                siglas_unicas[sig] = i

    # Ordenar alfabeticamente (convencao ABNT 5.10)
    siglas_ordenadas = sorted(siglas_unicas.keys())

    # 3. Construir paragrafos
    paragrafos = [_criar_paragrafo_titulo(
        'Lista de Siglas e Abreviaturas', estilo_titulo,
    )]
    for sig in siglas_ordenadas:
        paragrafos.append(_criar_paragrafo_entrada_lista(
            texto=f'{sig}  —  (descrição a preencher)',
            bookmark_destino=None,
            estilo_paragrafo=estilo_entrada,
        ))

    # 4. Inserir respeitando ordem ABNT (penultimo, antes do Sumario)
    pos = _calcular_posicao_insercao(body, prioridade=4)
    _inserir_bloco(
        body, pos, paragrafos,
        MARCADOR_LISTA_SIG_INI, MARCADOR_LISTA_SIG_FIM, id_gen,
    )

    doc.save(caminho_master)
    return {
        'entradas': len(siglas_ordenadas),
        'estilo_titulo': estilo_titulo,
        'posicao_inserido': pos,
        'perfil_origem': perfil.origem,
    }
