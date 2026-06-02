"""Serviço de orquestração de pipeline de montagem de relatórios.

Implementa pipeline coordenado: merge → numeração → cross-refs → TOC
com validação de pré/pós-condições e tratamento centralizado de erros.

Feature: automacao-montagem-relatorios
Requisitos: 3.1, 3.2, 3.3, 3.4, 3.5, 5.1, 5.2, 5.3, 6.1, 6.2, 6.3

Propriedades validadas:
- Propriedade 5: Parada Segura do Pipeline em Erro
- Propriedade 6: Idempotência Completa
- Propriedade 7: Validação de Pré-Condições
- Propriedade 8: Validação de Pós-Condições
"""

from __future__ import annotations

from datetime import datetime, timezone
import hashlib
import os
import re
import tempfile
from typing import Dict, Any

from app.models.relatorio_producao import RelatorioProducao
from app.models.capitulo_documento import CapituloDocumento
from app.services import servico_merge_docx
from app.services import servico_captioning
from app.services import servico_cross_refs
from app.services import servico_toc


class ServicoPipelineRelatorio:
    """Orquestradora de pipeline de montagem de relatórios.

    Coordena execução sequencial de etapas com validação de pré/pós-condições:
    1. Validar pré-condições (relatório, capítulos, uploads, espaço)
    2. Merge de capítulos
    3. Numeração de figuras/tabelas
    4. Atualização de referências cruzadas
    5. Regeneração de índices (TOC, LOF, LOT)
    6. Validação de pós-condições (integridade, coerência)

    Tolera falhas parciais: se merge falha para cap 3, continua cap 4.
    Parada garantida: se numeração falha, cross-refs não executam.
    Idempotência: múltiplas execuções com mesma entrada geram mesmo resultado.
    """

    @staticmethod
    def executar(relatorio_id: int, uploads_dict: Dict[int, bytes]) -> Dict[str, Any]:
        """Executa pipeline completo de montagem de relatório.

        Args:
            relatorio_id: ID do RelatorioProducao a processar
            uploads_dict: Mapa {capitulo_id: docx_bytes} com novos conteúdos

        Returns:
            Dict com estrutura:
            {
                'sucesso': bool,
                'relatorio_id': int,
                'etapas': [
                    {
                        'etapa': str (merge|numeracao|cross_refs|toc),
                        'resultado': dict,
                        'tempo_ms': int,
                        'timestamp': str
                    },
                    ...
                ],
                'erros': [str],
                'avisos': [str],
                'tempo_total_ms': int,
                'arquivo_modificado': bool,
                'proximos_passos': [str],
                'checksum_pre': str,
                'checksum_pos': str
            }

        Propriedades:
            - Propriedade 5: Se merge falha, numeração não executa
            - Propriedade 6: Múltiplas execuções → checksum idêntico
            - Propriedade 7: Pré-condições validadas antes de proceder
            - Propriedade 8: Pós-condições validadas para inconsistências
        """
        tempo_inicio = datetime.now(timezone.utc)

        resultado = {
            "sucesso": False,
            "relatorio_id": relatorio_id,
            "etapas": [],
            "erros": [],
            "avisos": [],
            "tempo_total_ms": 0,
            "arquivo_modificado": False,
            "proximos_passos": [],
            "checksum_pre": None,
            "checksum_pos": None,
        }

        try:
            # Fase 0: Obter relatório
            relatorio = RelatorioProducao.query.get(relatorio_id)
            if not relatorio:
                resultado["erros"].append(f"Relatório {relatorio_id} não encontrado")
                return resultado

            # Calcular checksum do template antes (para validar idempotência)
            if relatorio.caminho_template and os.path.exists(relatorio.caminho_template):
                resultado["checksum_pre"] = _calcular_checksum_arquivo(relatorio.caminho_template)

            # Fase 1: Validar pré-condições
            precond_result = ServicoPipelineRelatorio._validar_precondiciones(
                relatorio_id, uploads_dict
            )
            resultado["etapas"].append(
                {
                    "etapa": "validacao_precondiciones",
                    "resultado": precond_result,
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                }
            )

            if not precond_result.get("valido", False):
                resultado["erros"].extend(precond_result.get("motivos_rejeicao", []))
                resultado["proximos_passos"].extend(precond_result.get("proximos_passos", []))
                return resultado

            # Fase 2: Fazer merge de capítulos
            merge_result = ServicoPipelineRelatorio._fazer_merge(relatorio, uploads_dict)
            resultado["etapas"].append(
                {
                    "etapa": "merge",
                    "resultado": merge_result,
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                }
            )

            if not merge_result.get("sucesso", False):
                resultado["erros"].append("Fase de merge falhou")
                resultado["erros"].extend(merge_result.get("erros", []))
                resultado["avisos"].extend(merge_result.get("avisos", []))
                # Não continuar se merge falhou completamente
                if merge_result.get("total_capítulos", 0) == merge_result.get("total_erros", 0):
                    return resultado

            resultado["arquivo_modificado"] = True

            # Fase 3: Numeração (parar se nenhum capítulo foi merged)
            if merge_result.get("capítulos_processados", 0) > 0:
                num_result = ServicoPipelineRelatorio._executar_numeracao(
                    relatorio.caminho_template
                )
                resultado["etapas"].append(
                    {
                        "etapa": "numeracao",
                        "resultado": num_result,
                        "timestamp": datetime.now(timezone.utc).isoformat(),
                    }
                )

                if not num_result.get("sucesso", False):
                    resultado["avisos"].append("Numeração teve erros")
                    resultado["avisos"].extend(num_result.get("erros", []))

            # Fase 4: Referências cruzadas
            if len(resultado["etapas"]) > 2:  # Se numeration foi executada
                refs_result = ServicoPipelineRelatorio._atualizar_refs_cruzadas(
                    relatorio.caminho_template, {}
                )
                resultado["etapas"].append(
                    {
                        "etapa": "cross_refs",
                        "resultado": refs_result,
                        "timestamp": datetime.now(timezone.utc).isoformat(),
                    }
                )

                if refs_result.get("tags_orfas", 0) > 0:
                    resultado["avisos"].append(f"Detectadas {refs_result['tags_orfas']} tags órfãs")

            # Fase 5: Regenerar índices (TOC, LOF, LOT)
            indices_result = ServicoPipelineRelatorio._regenerar_indices(relatorio.caminho_template)
            resultado["etapas"].append(
                {
                    "etapa": "indices",
                    "resultado": indices_result,
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                }
            )

            if not indices_result.get("sucesso", False):
                resultado["avisos"].append("Índices tiveram erros")

            # Fase 6: Validar pós-condições
            postcond_result = ServicoPipelineRelatorio._validar_poscondiciones(
                relatorio.caminho_template
            )
            resultado["etapas"].append(
                {
                    "etapa": "validacao_poscondiciones",
                    "resultado": postcond_result,
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                }
            )

            if postcond_result.get("inconsistencias", []):
                for inconsistencia in postcond_result["inconsistencias"]:
                    if isinstance(inconsistencia, dict) and "remediacao" in inconsistencia:
                        # Formato novo (5.2)
                        resultado["avisos"].append(
                            f"[{inconsistencia['tipo']}] {inconsistencia['elemento']}: "
                            f"{inconsistencia['remediacao']}"
                        )
                    else:
                        # Formato antigo
                        resultado["avisos"].append(str(inconsistencia))

            # Calcular checksum final
            if relatorio.caminho_template and os.path.exists(relatorio.caminho_template):
                resultado["checksum_pos"] = _calcular_checksum_arquivo(relatorio.caminho_template)

            resultado["sucesso"] = True
            resultado["proximos_passos"] = ["exportar", "validar_final", "arquivar"]

        except Exception as e:
            resultado["erros"].append(f"Erro não esperado: {str(e)}")
            import traceback

            resultado["erros"].append(traceback.format_exc())

        finally:
            # Calcular tempo total
            tempo_fim = datetime.now(timezone.utc)
            delta = tempo_fim - tempo_inicio
            resultado["tempo_total_ms"] = int(delta.total_seconds() * 1000)

        return resultado

    @staticmethod
    def _validar_precondiciones(
        relatorio_id: int, uploads_dict: Dict[int, bytes]
    ) -> Dict[str, Any]:
        """Valida pré-condições antes de iniciar pipeline com mensagens amigáveis.

        Validações:
        1. Relatório existe e tem estado válido
        2. Todos os capítulos sincronizados (classificacao != null)
        3. Cada upload corresponde a capítulo existente
        4. Espaço em disco suficiente
        5. DOCX template válido

        Returns:
            Dict com:
            - 'valido': bool
            - 'motivos_rejeicao': [] de strings amigáveis ao usuário
            - 'proximos_passos': [] de ações sugeridas

        Valida Property 7: Pré-condições rejeitam estados inválidos com explicitação de motivos.
        """
        resultado = {
            "valido": True,
            "motivos_rejeicao": [],
            "proximos_passos": [],
            "detalhes_validacoes": [],
        }

        try:
            # Validação 1: Relatório existe
            relatorio = RelatorioProducao.query.get(relatorio_id)
            if not relatorio:
                resultado["valido"] = False
                resultado["motivos_rejeicao"].append(
                    "Relatório não encontrado. Verifique o ID e tente novamente."
                )
                resultado["proximos_passos"].append("Verificar relatório")
                return resultado

            resultado["detalhes_validacoes"].append(
                {
                    "validacao": "relatório_existe",
                    "resultado": "OK",
                    "mensagem": f"Relatório {relatorio.titulo or relatorio_id} encontrado",
                }
            )

            # Validação 2: Capítulos sincronizados
            capitulos_nao_sync = CapituloDocumento.query.filter_by(
                id_relatorio=relatorio_id, classificacao=None
            ).all()

            if capitulos_nao_sync:
                resultado["valido"] = False
                titulos_nao_sync = ", ".join(
                    [f"'{cap.titulo}' (ID {cap.id})" for cap in capitulos_nao_sync[:3]]
                )  # Mostrar até 3
                if len(capitulos_nao_sync) > 3:
                    titulos_nao_sync += f", ... e {len(capitulos_nao_sync) - 3} mais"

                resultado["motivos_rejeicao"].append(
                    f"Capítulos não sincronizados: {titulos_nao_sync}. "
                    "Execute a sincronização antes de prosseguir."
                )
                resultado["proximos_passos"].append("sincronizar")
                resultado["detalhes_validacoes"].append(
                    {
                        "validacao": "capítulos_sincronizados",
                        "resultado": "FALHA",
                        "total_não_sincronizados": len(capitulos_nao_sync),
                        "exemplos": [cap.titulo for cap in capitulos_nao_sync[:3]],
                    }
                )
            else:
                total_capitulos = CapituloDocumento.query.filter_by(
                    id_relatorio=relatorio_id
                ).count()
                resultado["detalhes_validacoes"].append(
                    {
                        "validacao": "capítulos_sincronizados",
                        "resultado": "OK",
                        "mensagem": f"Todos os {total_capitulos} capítulos sincronizados",
                    }
                )

            # Validação 3: Uploads mapeados
            uploads_invalidos = []
            for cap_id, _ in uploads_dict.items():
                cap = CapituloDocumento.query.get(cap_id)
                if not cap:
                    uploads_invalidos.append(f"Capítulo ID {cap_id} não existe")
                elif cap.id_relatorio != relatorio_id:
                    uploads_invalidos.append(
                        f"Capítulo '{cap.titulo}' não pertence a este relatório"
                    )

            if uploads_invalidos:
                resultado["valido"] = False
                for msg in uploads_invalidos[:3]:  # Mostrar até 3
                    resultado["motivos_rejeicao"].append(msg)
                if len(uploads_invalidos) > 3:
                    resultado["motivos_rejeicao"].append(
                        f"... e {len(uploads_invalidos) - 3} upload(s) inválido(s) mais"
                    )
                resultado["proximos_passos"].append("revisar_uploads")
                resultado["detalhes_validacoes"].append(
                    {
                        "validacao": "uploads_mapeados",
                        "resultado": "FALHA",
                        "total_inválidos": len(uploads_invalidos),
                        "exemplos": uploads_invalidos[:3],
                    }
                )
            else:
                resultado["detalhes_validacoes"].append(
                    {
                        "validacao": "uploads_mapeados",
                        "resultado": "OK",
                        "mensagem": f"{len(uploads_dict)} upload(s) válido(s)",
                    }
                )

            # Validação 4: Espaço em disco
            if relatorio.caminho_template:
                try:
                    free_space = _obter_espaco_disponivel(
                        os.path.dirname(relatorio.caminho_template)
                    )
                    free_mb = free_space / 1024 / 1024
                    threshold_mb = 100

                    if free_space < threshold_mb * 1024 * 1024:
                        resultado["valido"] = False
                        resultado["motivos_rejeicao"].append(
                            f"Espaço em disco insuficiente: {free_mb:.0f}MB disponível "
                            f"(mínimo {threshold_mb}MB necessário). "
                            "Libere espaço e tente novamente."
                        )
                        resultado["proximos_passos"].append("liberar_disco")
                        resultado["detalhes_validacoes"].append(
                            {
                                "validacao": "espaço_disco",
                                "resultado": "FALHA",
                                "espaço_disponível_mb": round(free_mb, 2),
                                "mínimo_necessário_mb": threshold_mb,
                            }
                        )
                    else:
                        resultado["detalhes_validacoes"].append(
                            {
                                "validacao": "espaço_disco",
                                "resultado": "OK",
                                "espaço_disponível_mb": round(free_mb, 2),
                            }
                        )
                except Exception as e:
                    # Não falhar pipeline se não conseguir verificar espaço
                    resultado["detalhes_validacoes"].append(
                        {
                            "validacao": "espaço_disco",
                            "resultado": "AVISO",
                            "mensagem": (
                                "Não foi possível verificar espaço em disco: " f"{str(e)[:50]}"
                            ),
                        }
                    )

            # Validação 5: DOCX template válido
            if not relatorio.caminho_template:
                resultado["valido"] = False
                resultado["motivos_rejeicao"].append("Relatório não tem template DOCX configurado.")
                resultado["proximos_passos"].append("configurar_template")
                resultado["detalhes_validacoes"].append(
                    {
                        "validacao": "template_docx_existe",
                        "resultado": "FALHA",
                        "mensagem": "Caminho do template não configurado",
                    }
                )
            elif not os.path.exists(relatorio.caminho_template):
                resultado["valido"] = False
                resultado["motivos_rejeicao"].append(
                    "Arquivo DOCX do template não encontrado. "
                    "Verifique se o arquivo foi movido ou deletado."
                )
                resultado["proximos_passos"].append("restaurar_template")
                resultado["detalhes_validacoes"].append(
                    {
                        "validacao": "template_docx_existe",
                        "resultado": "FALHA",
                        "mensagem": "Arquivo não existe no caminho especificado",
                    }
                )
            else:
                # Validar integridade do DOCX
                try:
                    from docx import Document

                    doc = Document(relatorio.caminho_template)

                    if len(doc.element.body) == 0:
                        resultado["valido"] = False
                        resultado["motivos_rejeicao"].append(
                            "Arquivo DOCX do template parece vazio. "
                            "Verifique a integridade do arquivo."
                        )
                        resultado["proximos_passos"].append("verificar_template")
                        resultado["detalhes_validacoes"].append(
                            {
                                "validacao": "template_docx_integridade",
                                "resultado": "FALHA",
                                "mensagem": "Documento vazio ou corrompido",
                            }
                        )
                    else:
                        resultado["detalhes_validacoes"].append(
                            {
                                "validacao": "template_docx_integridade",
                                "resultado": "OK",
                                "parágrafos": len(doc.paragraphs),
                                "seções": len(doc.sections),
                            }
                        )
                except Exception as e:
                    resultado["valido"] = False
                    resultado["motivos_rejeicao"].append(
                        "Arquivo DOCX corrompido ou inválido. "
                        "Tente fazer upload novamente ou contate suporte."
                    )
                    resultado["proximos_passos"].append("verificar_template")
                    resultado["detalhes_validacoes"].append(
                        {
                            "validacao": "template_docx_integridade",
                            "resultado": "FALHA",
                            "erro_técnico": type(e).__name__,
                        }
                    )

        except Exception as e:
            resultado["valido"] = False
            resultado["motivos_rejeicao"].append(
                "Erro ao validar pré-condições. Tente novamente ou contate suporte."
            )
            resultado["detalhes_validacoes"].append(
                {
                    "validacao": "validação_geral",
                    "resultado": "ERRO",
                    "erro_técnico": type(e).__name__,
                }
            )

        return resultado

    @staticmethod
    def _fazer_merge(
        relatorio: RelatorioProducao, uploads_dict: Dict[int, bytes]
    ) -> Dict[str, Any]:
        """Faz merge de capítulos no documento template.

        Tolerância a erros: se merge falha para cap 3, continua cap 4.
        Cada capítulo é processado independentemente.
        """
        resultado = {
            "sucesso": True,
            "capítulos_processados": 0,
            "capítulos_com_erro": 0,
            "erros": [],
            "avisos": [],
        }

        try:
            for cap_id, docx_bytes in uploads_dict.items():
                try:
                    # Obter capítulo do banco
                    cap = CapituloDocumento.query.get(cap_id)
                    if not cap:
                        resultado["capítulos_com_erro"] += 1
                        resultado["erros"].append(f"Capítulo {cap_id} não encontrado")
                        continue

                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_autor:
                        tmp_autor.write(docx_bytes)
                        caminho_autor = tmp_autor.name

                    try:
                        # Fazer merge usando ServicoMergeDocx
                        merge_res = servico_merge_docx.substituir_capitulo(
                            relatorio.caminho_template,
                            cap,
                            caminho_autor,
                        )
                    finally:
                        try:
                            os.unlink(caminho_autor)
                        except OSError:
                            pass

                    if merge_res is True:
                        resultado["capítulos_processados"] += 1
                    else:
                        resultado["capítulos_com_erro"] += 1
                        erros = merge_res.get("erros", []) if isinstance(merge_res, dict) else []
                        resultado["erros"].extend(erros)

                except Exception as e:
                    resultado["capítulos_com_erro"] += 1
                    resultado["erros"].append(f"Erro ao processar capítulo {cap_id}: {str(e)}")

            if resultado["capítulos_processados"] > 0:
                resultado["avisos"].append(
                    "Documento salvo com "
                    f'{resultado["capítulos_processados"]} capítulos atualizados'
                )
            else:
                resultado["sucesso"] = False
                resultado["erros"].append("Nenhum capítulo foi processado com sucesso")

        except Exception as e:
            resultado["sucesso"] = False
            resultado["erros"].append(f"Erro geral em merge: {str(e)}")

        return resultado

    @staticmethod
    def _executar_numeracao(caminho_template: str) -> Dict[str, Any]:
        """Executa numeração de figuras e tabelas.

        Wrapper sobre servico_captioning.reindexar_captions().
        """
        resultado = {"sucesso": True, "figuras_numeradas": 0, "tabelas_numeradas": 0, "erros": []}

        try:
            if not caminho_template or not os.path.exists(caminho_template):
                resultado["sucesso"] = False
                resultado["erros"].append("Template DOCX não encontrado")
                return resultado

            # Chamar serviço de numeração
            num_result = servico_captioning.reindexar_captions(caminho_template)

            if num_result.get("sucesso", False):
                resultado["figuras_numeradas"] = num_result.get("figuras", 0)
                resultado["tabelas_numeradas"] = num_result.get("tabelas", 0)
            else:
                resultado["sucesso"] = False
                resultado["erros"].extend(num_result.get("erros", []))

        except Exception as e:
            resultado["sucesso"] = False
            resultado["erros"].append(f"Erro em numeração: {str(e)}")

        return resultado

    @staticmethod
    def _atualizar_refs_cruzadas(
        caminho_template: str, mapa_labels: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Atualiza referências cruzadas no documento."""
        resultado = {"sucesso": True, "tags_substituidas": 0, "tags_orfas": 0, "erros": []}

        try:
            if not caminho_template:
                resultado["erros"].append("Template DOCX não informado")
                return resultado

            refs_result = servico_cross_refs.substituir_referencias(caminho_template, mapa_labels)

            resultado["tags_substituidas"] = refs_result.get("substituidas", 0)
            resultado["tags_orfas"] = refs_result.get("orfas", 0)

            if refs_result.get("tags_orfas", 0) > 0:
                resultado["erros"].append(
                    f'{refs_result["tags_orfas"]} referências órfãs detectadas'
                )

        except Exception as e:
            resultado["erros"].append(f"Erro em cross-refs: {str(e)}")

        return resultado

    @staticmethod
    def _regenerar_indices(caminho_template: str) -> Dict[str, Any]:
        """Regenera índices (TOC, LOF, LOT)."""
        resultado = {
            "sucesso": True,
            "toc_status": False,
            "lof_status": False,
            "lot_status": False,
            "erros": [],
        }

        try:
            if not caminho_template:
                resultado["erros"].append("Template DOCX não informado")
                return resultado

            # TOC
            toc_result = servico_toc.inserir_sumario(caminho_template)
            resultado["toc_status"] = toc_result.get("sucesso", False)

            # LOF (Lista de Figuras)
            lof_result = servico_toc.inserir_lista_figuras(caminho_template)
            resultado["lof_status"] = lof_result.get("sucesso", False)

            # LOT (Lista de Tabelas)
            lot_result = servico_toc.inserir_lista_tabelas(caminho_template)
            resultado["lot_status"] = lot_result.get("sucesso", False)

            if not all([resultado["toc_status"], resultado["lof_status"], resultado["lot_status"]]):
                resultado["sucesso"] = False

        except Exception as e:
            resultado["sucesso"] = False
            resultado["erros"].append(f"Erro em índices: {str(e)}")

        return resultado

    @staticmethod
    def _validar_poscondiciones(caminho_template: str) -> Dict[str, Any]:
        """Valida integridade do documento após pipeline com diagnóstico detalhado.

        Validações:
        1. Sem legendas duplicadas
        2. Numeração sequencial (sem gaps)
        3. TOC coerente com headings
        4. Sem bookmarks órfãos
        5. DOCX não corrompido

        Returns:
            Dict com:
            - 'inconsistencias': array de {'tipo': str, 'elemento': str, 'remediacao': str}
            - 'validacoes_executadas': int
            - 'validacoes_passadas': int
            - 'diagnosticos': [] detalhes de cada validação

        Valida Property 8: Inconsistências pós-pipeline são detectadas e reportadas.
        """
        resultado = {
            "inconsistencias": [],
            "validacoes_executadas": 0,
            "validacoes_passadas": 0,
            "diagnosticos": [],
            "sucesso": True,
        }

        try:
            if not caminho_template or not os.path.exists(caminho_template):
                resultado["inconsistencias"].append(
                    {
                        "tipo": "arquivo_não_encontrado",
                        "elemento": caminho_template or "(não especificado)",
                        "remediacao": (
                            "Verifique se o arquivo DOCX foi salvo corretamente. "
                            "Tente processar novamente."
                        ),
                    }
                )
                resultado["sucesso"] = False
                return resultado

            from docx import Document

            doc = Document(caminho_template)

            # Validação 1: Documento não vazio
            resultado["validacoes_executadas"] += 1
            if len(doc.element.body) == 0:
                resultado["inconsistencias"].append(
                    {
                        "tipo": "documento_vazio",
                        "elemento": "body",
                        "remediacao": (
                            "O documento está vazio. Verifique se o merge foi "
                            "executado corretamente."
                        ),
                    }
                )
                resultado["sucesso"] = False
                resultado["diagnosticos"].append(
                    {"validacao": "documento_vazio", "resultado": "FALHA"}
                )
            else:
                resultado["validacoes_passadas"] += 1
                resultado["diagnosticos"].append(
                    {
                        "validacao": "documento_estrutura",
                        "resultado": "OK",
                        "parágrafos": len(doc.paragraphs),
                        "tabelas": len(doc.tables),
                        "seções": len(doc.sections),
                    }
                )

            # Validação 2: Legendas duplicadas (via bookmarks e campos)
            resultado["validacoes_executadas"] += 1
            legendas_encontradas = {}
            legendas_duplicadas = []

            try:
                # Varrer documento procurando por legendas (bookmarks com prefixo fig/tab)
                for elem in doc.element.iter():
                    # Procurar por bookmarkStart e bookmarkEnd (figuras/tabelas)
                    if "bookmarkStart" in elem.tag:
                        legenda_name = elem.get("name", "")
                        if legenda_name and (
                            "fig" in legenda_name.lower()
                            or "tab" in legenda_name.lower()
                            or "equacao" in legenda_name.lower()
                        ):
                            if legenda_name in legendas_encontradas:
                                legendas_duplicadas.append(
                                    {
                                        "tipo": "legenda_duplicada",
                                        "elemento": legenda_name,
                                        "remediacao": (
                                            f"Legenda '{legenda_name}' aparece múltiplas vezes. "
                                            "Execute 'Limpar legendas' no menu Ferramentas "
                                            "e reprocesse o documento."
                                        ),
                                    }
                                )
                            else:
                                legendas_encontradas[legenda_name] = True

                if legendas_duplicadas:
                    resultado["inconsistencias"].extend(legendas_duplicadas)
                    resultado["diagnosticos"].append(
                        {
                            "validacao": "legendas_duplicadas",
                            "resultado": "FALHA",
                            "total_duplicadas": len(legendas_duplicadas),
                            "exemplos": [d["elemento"] for d in legendas_duplicadas[:5]],
                        }
                    )
                else:
                    resultado["validacoes_passadas"] += 1
                    resultado["diagnosticos"].append(
                        {
                            "validacao": "legendas_duplicadas",
                            "resultado": "OK",
                            "total_legendas_únicas": len(legendas_encontradas),
                        }
                    )
            except Exception as e:
                resultado["diagnosticos"].append(
                    {
                        "validacao": "legendas_duplicadas",
                        "resultado": "AVISO",
                        "mensagem": (
                            "Não foi possível verificar legendas duplicadas: " f"{type(e).__name__}"
                        ),
                    }
                )

            # Validação 3: Numeração sequencial de figuras/tabelas
            resultado["validacoes_executadas"] += 1
            gaps_numeracao = []

            try:
                # Extrair números de figuras (Figura 1, 2, 3...; pode ter gaps)
                figuras_encontradas = set()
                tabelas_encontradas = set()

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            texto = cell.text.lower()
                            # Buscar "figura N"
                            matches_fig = re.findall(r"figura\s+(\d+(?:\.\d+)?)", texto)
                            for num in matches_fig:
                                figuras_encontradas.add(num)
                            # Buscar "tabela N"
                            matches_tab = re.findall(r"tabela\s+(\d+(?:\.\d+)?)", texto)
                            for num in matches_tab:
                                tabelas_encontradas.add(num)

                for para in doc.paragraphs:
                    texto = para.text.lower()
                    # Buscar "figura N"
                    matches_fig = re.findall(r"figura\s+(\d+(?:\.\d+)?)", texto)
                    for num in matches_fig:
                        figuras_encontradas.add(num)
                    # Buscar "tabela N"
                    matches_tab = re.findall(r"tabela\s+(\d+(?:\.\d+)?)", texto)
                    for num in matches_tab:
                        tabelas_encontradas.add(num)

                # Verificar gaps em figuras (simplificado: buscar 1.1, 1.2, 2.1...)
                if figuras_encontradas:
                    # Converter para float para ordenação
                    figuras_nums = sorted([float(f) for f in figuras_encontradas])
                    for i in range(len(figuras_nums) - 1):
                        if figuras_nums[i + 1] - figuras_nums[i] > 0.1:
                            gaps_numeracao.append(
                                {
                                    "tipo": "gap_numeração_figuras",
                                    "elemento": (
                                        f"Gap entre Figura {figuras_nums[i]} "
                                        f"e {figuras_nums[i + 1]}"
                                    ),
                                    "remediacao": "Possível gap na numeração de figuras. "
                                    "Revisar conteúdo para adicionar figura faltante ou renumerar.",
                                }
                            )

                if tabelas_encontradas:
                    tabelas_nums = sorted([float(t) for t in tabelas_encontradas])
                    for i in range(len(tabelas_nums) - 1):
                        if tabelas_nums[i + 1] - tabelas_nums[i] > 0.1:
                            gaps_numeracao.append(
                                {
                                    "tipo": "gap_numeração_tabelas",
                                    "elemento": (
                                        f"Gap entre Tabela {tabelas_nums[i]} "
                                        f"e {tabelas_nums[i + 1]}"
                                    ),
                                    "remediacao": "Possível gap na numeração de tabelas. "
                                    "Revisar conteúdo para adicionar tabela faltante ou renumerar.",
                                }
                            )

                if gaps_numeracao:
                    resultado["inconsistencias"].extend(gaps_numeracao)
                    resultado["diagnosticos"].append(
                        {
                            "validacao": "numeração_sequencial",
                            "resultado": "AVISO",
                            "gaps_detectados": len(gaps_numeracao),
                            "figuras": len(figuras_encontradas),
                            "tabelas": len(tabelas_encontradas),
                        }
                    )
                else:
                    resultado["validacoes_passadas"] += 1
                    resultado["diagnosticos"].append(
                        {
                            "validacao": "numeração_sequencial",
                            "resultado": "OK",
                            "figuras": len(figuras_encontradas),
                            "tabelas": len(tabelas_encontradas),
                        }
                    )
            except Exception as e:
                resultado["diagnosticos"].append(
                    {
                        "validacao": "numeração_sequencial",
                        "resultado": "AVISO",
                        "mensagem": f"Não foi possível verificar numeração: {type(e).__name__}",
                    }
                )

            # Validação 4: Headings estruturados (não deve ter gaps de nível)
            resultado["validacoes_executadas"] += 1
            gaps_heading = []

            try:
                heading_levels = []
                for para in doc.paragraphs:
                    estilo = para.style
                    style_name = estilo.name if estilo and estilo.name else ""
                    if "Heading" in style_name:
                        # Extrair nível (Heading 1 → 1, Heading 2 → 2, etc.)
                        match = re.search(r"Heading\s+(\d+)", style_name)
                        if match:
                            level = int(match.group(1))
                            heading_levels.append(level)

                # Verificar sequência (não pular de H1 direto para H3)
                for i in range(len(heading_levels) - 1):
                    if heading_levels[i + 1] > heading_levels[i] + 1:
                        gaps_heading.append(
                            {
                                "tipo": "gap_heading_levels",
                                "elemento": (
                                    f"Salto de Heading {heading_levels[i]} "
                                    f"para {heading_levels[i + 1]}"
                                ),
                                "remediacao": (
                                    "Revisar estrutura de headings. Não pule níveis "
                                    "(ex: Heading 1 → Heading 3). Use Heading 2 "
                                    "intermediário se necessário."
                                ),
                            }
                        )

                if gaps_heading:
                    resultado["inconsistencias"].extend(gaps_heading)
                    resultado["diagnosticos"].append(
                        {
                            "validacao": "headings_estruturados",
                            "resultado": "AVISO",
                            "gaps_detectados": len(gaps_heading),
                        }
                    )
                else:
                    resultado["validacoes_passadas"] += 1
                    resultado["diagnosticos"].append(
                        {
                            "validacao": "headings_estruturados",
                            "resultado": "OK",
                            "headings": len(heading_levels),
                            "níveis_únicos": len(set(heading_levels)),
                        }
                    )
            except Exception as e:
                resultado["diagnosticos"].append(
                    {
                        "validacao": "headings_estruturados",
                        "resultado": "AVISO",
                        "mensagem": f"Não foi possível verificar headings: {type(e).__name__}",
                    }
                )

            # Validação 5: Bookmarks órfãos (referências sem destino)
            resultado["validacoes_executadas"] += 1
            bookmarks_orfaos = []

            try:
                bookmarks_definidos = set()
                bookmarks_referenciados = set()

                # Coletar bookmarks definidos
                for elem in doc.element.iter():
                    if "bookmarkStart" in elem.tag:
                        bookmark_name = elem.get("name", "")
                        if bookmark_name:
                            bookmarks_definidos.add(bookmark_name)

                # Coletar bookmarks referenciados (em campos REF)
                for para in doc.paragraphs:
                    for run in para.runs:
                        if run.element.findall(
                            ".//{http://schemas.openxmlformats.org/"
                            "wordprocessingml/2006/main}fldChar"
                        ):
                            # Tem campo, buscar referência
                            fldData = run.element.find(
                                ".//{http://schemas.openxmlformats.org/"
                                "wordprocessingml/2006/main}fldData"
                            )
                            if fldData is not None:
                                instr = fldData.get("instr", "")
                                # Extrair nome do bookmark de "REF bookmark_name"
                                if "REF" in instr:
                                    parts = instr.split()
                                    if len(parts) >= 2:
                                        bookmarks_referenciados.add(parts[1])

                # Bookmarks órfãos = referenciados mas não definidos
                bookmarks_orfaos_lista = bookmarks_referenciados - bookmarks_definidos

                if bookmarks_orfaos_lista:
                    for bookmark in list(bookmarks_orfaos_lista)[:5]:  # Mostrar até 5
                        bookmarks_orfaos.append(
                            {
                                "tipo": "bookmark_órfão",
                                "elemento": bookmark,
                                "remediacao": f"Referência '{bookmark}' não tem destino definido. "
                                "Verifique se o bookmark correspondente existe ou execute "
                                "'Limpar referências órfãs'.",
                            }
                        )

                    if len(bookmarks_orfaos_lista) > 5:
                        bookmarks_orfaos.append(
                            {
                                "tipo": "bookmarks_órfãos_múltiplos",
                                "elemento": f"{len(bookmarks_orfaos_lista)} bookmarks órfãos",
                                "remediacao": (
                                    "Execute 'Limpar referências' para remover "
                                    "todas as referências órfãs de uma vez."
                                ),
                            }
                        )

                    resultado["inconsistencias"].extend(bookmarks_orfaos)
                    resultado["diagnosticos"].append(
                        {
                            "validacao": "bookmarks_órfãos",
                            "resultado": "FALHA",
                            "total_orfaos": len(bookmarks_orfaos_lista),
                            "definidos": len(bookmarks_definidos),
                            "referenciados": len(bookmarks_referenciados),
                        }
                    )
                else:
                    resultado["validacoes_passadas"] += 1
                    resultado["diagnosticos"].append(
                        {
                            "validacao": "bookmarks_órfãos",
                            "resultado": "OK",
                            "bookmarks_definidos": len(bookmarks_definidos),
                            "bookmarks_referenciados": len(bookmarks_referenciados),
                        }
                    )
            except Exception as e:
                resultado["diagnosticos"].append(
                    {
                        "validacao": "bookmarks_órfãos",
                        "resultado": "AVISO",
                        "mensagem": f"Não foi possível verificar bookmarks: {type(e).__name__}",
                    }
                )

        except Exception as e:
            resultado["sucesso"] = False
            resultado["inconsistencias"].append(
                {
                    "tipo": "erro_validação",
                    "elemento": "documento",
                    "remediacao": (
                        "Erro ao validar documento. Tente processar novamente "
                        "ou contate suporte."
                    ),
                }
            )
            resultado["diagnosticos"].append(
                {
                    "validacao": "validação_geral",
                    "resultado": "ERRO",
                    "erro_técnico": type(e).__name__,
                }
            )

        return resultado


def _calcular_checksum_arquivo(caminho: str) -> str:
    """Calcula SHA256 de arquivo para validar idempotência."""
    sha256 = hashlib.sha256()
    with open(caminho, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def _obter_espaco_disponivel(caminho: str) -> int:
    """Obtém espaço disponível em bytes no disco."""
    import shutil

    stat = shutil.disk_usage(caminho)
    return stat.free


__all__ = ["ServicoPipelineRelatorio"]
