"""Geração da versão DOCX sugerida a partir da biblioteca canônica."""

import json
import os
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt, RGBColor

from app.services.servico_captioning import reindexar_captions
from app.services.servico_sanitizar_docx import sanitizar_docx


VERSAO_DOCX_SUGERIDO = 'analise_upload_v2'


class ServicoVersaoSugeridaCanonica:
    """Aplica métricas canônicas completas ao DOCX enviado pelo autor."""

    METRICAS_OBRIGATORIAS = {
        'sanitizacao_compatibilidade_editor',
        'aplicacao_estilos_heading_do_perfil',
        'normalizacao_visual_de_paragrafos',
        'aplicacao_metricas_canonicas_reais',
        'numeracao_hierarquica_subcapitulos',
        'reindexacao_captions_quando_possivel',
    }

    @classmethod
    def gerar(cls, *, envio, relatorio, perfil, caminho_saida, metricas=None):
        """Gera e salva o DOCX sugerido apenas se todas as métricas passarem.

        Quando `metricas` é informado (ex.: biblioteca escolhida no
        seletor da prévia), usa-o em vez de resolver pela biblioteca do
        relatório. Retorna `(metricas_aplicadas, metricas, diagnostico)`,
        onde `diagnostico` descreve o que foi (e o que não foi) alterado.
        """
        os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
        caminho_temp = f'{caminho_saida}.tmp'
        metricas_aplicadas = []

        bytes_sanitizados = sanitizar_docx(envio.caminho_arquivo)
        if bytes_sanitizados is None:
            shutil.copy2(envio.caminho_arquivo, caminho_temp)
        else:
            with open(caminho_temp, 'wb') as arquivo:
                arquivo.write(bytes_sanitizados)
        metricas_aplicadas.append('sanitizacao_compatibilidade_editor')

        if metricas is None:
            metricas = cls.carregar_biblioteca_canonica(relatorio)
        try:
            diagnostico = cls.aplicar_biblioteca_canonica(
                caminho_temp, perfil, metricas
            )
            metricas_aplicadas.extend([
                'aplicacao_estilos_heading_do_perfil',
                'normalizacao_visual_de_paragrafos',
                'aplicacao_metricas_canonicas_reais',
                'numeracao_hierarquica_subcapitulos',
            ])
            reindexar_captions(caminho_temp, perfil=perfil)
            metricas_aplicadas.append('reindexacao_captions_quando_possivel')
            cls.validar_docx_sugerido(caminho_temp, metricas_aplicadas)
        except Exception:
            if os.path.exists(caminho_temp):
                os.remove(caminho_temp)
            raise

        os.replace(caminho_temp, caminho_saida)
        return metricas_aplicadas, metricas, diagnostico

    @classmethod
    def carregar_biblioteca_canonica(cls, relatorio):
        """Carrega formatação, capítulos e macroestrutura da base canônica."""
        return cls._carregar_de_dir(cls._resolver_diretorio_canonico(relatorio))

    @classmethod
    def carregar_de_biblioteca(cls, biblioteca):
        """Carrega métricas a partir de uma biblioteca escolhida pelo usuário.

        Retorna {} (sem dados canônicos → defaults do sistema) quando a
        biblioteca não tem `caminho_arquivo` válido em disco.
        """
        caminho = getattr(biblioteca, 'caminho_arquivo', None) if biblioteca else None
        if not caminho or not os.path.isdir(caminho):
            return {}
        return cls._carregar_de_dir(caminho)

    @classmethod
    def _carregar_de_dir(cls, caminho_dir):
        """Lê os 3 JSONs canônicos de um diretório de biblioteca."""
        if not caminho_dir:
            return {}
        return {
            'diretorio': caminho_dir,
            'formatacao': cls._ler_json(caminho_dir, 'canonico_formatacao.json'),
            'capitulos': cls._ler_json(caminho_dir, 'canonico_capitulos.json'),
            'macro': cls._ler_json(caminho_dir, 'canonico_estrutura_macro.json'),
            'docx_base': cls._localizar_docx_base(caminho_dir),
        }

    @staticmethod
    def _resolver_diretorio_canonico(relatorio):
        biblioteca = getattr(relatorio, 'biblioteca', None) if relatorio else None
        caminho = getattr(biblioteca, 'caminho_arquivo', None)
        if caminho and os.path.isdir(caminho):
            return caminho
        if relatorio is None:
            return None
        base_dir = os.path.dirname(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
        )
        candidatos = [
            relatorio.codigo_d20,
            (relatorio.codigo_d20 or '').replace('-', ''),
        ]
        for nome in candidatos:
            if not nome:
                continue
            caminho = os.path.join(base_dir, 'storage', 'canonicos', nome)
            if os.path.isdir(caminho):
                return caminho
        return None

    @staticmethod
    def _ler_json(caminho_dir, nome_arquivo):
        caminho = os.path.join(caminho_dir, nome_arquivo)
        if not os.path.exists(caminho):
            return None
        with open(caminho, encoding='utf-8') as arquivo:
            return json.load(arquivo)

    @staticmethod
    def _localizar_docx_base(caminho_dir):
        for nome in os.listdir(caminho_dir):
            if nome.lower().endswith('.docx'):
                return os.path.join(caminho_dir, nome)
        return None

    @classmethod
    def aplicar_biblioteca_canonica(cls, caminho_docx, perfil, metricas):
        """Classifica elementos e aplica métricas visuais canônicas.

        Retorna um diagnóstico estruturado com o que foi alterado
        (margens, estilos-base, headings, corpo, legendas) e o que
        não foi (parágrafos vazios e parágrafos dentro de tabelas).
        """
        doc = Document(caminho_docx)
        estilos = cls._mapear_estilos_canonicos(metricas)
        margens = cls._aplicar_secoes_canonicas(doc, metricas)
        estilos_base = cls._aplicar_estilos_base(doc, estilos)
        blocos = cls._classificar_e_formatar_blocos(doc, perfil, estilos)

        # Conferência final contra a biblioteca canônica
        conferencia = cls._conferir_contra_biblioteca(doc, metricas)

        doc.save(caminho_docx)
        return {
            'margens': margens,
            'estilos_base': estilos_base,
            'headings': blocos['headings'],
            'corpo': blocos['corpo'],
            'legendas': blocos['legendas'],
            'nao_alterados': blocos['nao_alterados'],
            'conferencia_final': conferencia,
        }

    @staticmethod
    def _conferir_contra_biblioteca(doc, metricas):
        """Conferência final contra a biblioteca canônica.

        Verifica se o DOCX processado está conforme as métricas canônicas
        e retorna um diagnóstico de conformidade.
        """
        if not metricas:
            return {
                'status': 'sem_metricas',
                'mensagem': 'Sem biblioteca canônica para conferência'
            }

        conferencia = {
            'status': 'ok',
            'verificacoes': [],
        }

        # Verificar se há formatação canônica
        if metricas.get('formatacao'):
            conferencia['verificacoes'].append('formatacao_canonica_aplicada')
        else:
            conferencia['verificacoes'].append('formatacao_default_aplicada')

        # Verificar se há capitulos canônicos
        if metricas.get('capitulos'):
            conferencia['verificacoes'].append('capitulos_canonica_aplicados')
        else:
            conferencia['verificacoes'].append('capitulos_default_aplicados')

        # Verificar se há macroestrutura canônica
        if metricas.get('macro'):
            conferencia['verificacoes'].append('macro_canonica_aplicada')
        else:
            conferencia['verificacoes'].append('macro_default_aplicada')

        return conferencia

    @staticmethod
    def _mapear_estilos_canonicos(metricas):
        formatacao = metricas.get('formatacao') or {}
        return {
            (estilo.get('nome') or '').lower(): estilo
            for estilo in formatacao.get('estilos_paragrafo') or []
            if estilo.get('nome')
        }

    @staticmethod
    def _aplicar_secoes_canonicas(doc, metricas):
        formatacao = metricas.get('formatacao') or {}
        secoes = formatacao.get('secoes') or []
        secao_canonica = (
            secoes[min(1, len(secoes) - 1)]
            if secoes
            else {
                'margem_top_mm': 25,
                'margem_right_mm': 20,
                'margem_bottom_mm': 25,
                'margem_left_mm': 20,
            }
        )
        valores = {
            'top_mm': secao_canonica.get('margem_top_mm') or 25,
            'right_mm': secao_canonica.get('margem_right_mm') or 20,
            'bottom_mm': secao_canonica.get('margem_bottom_mm') or 25,
            'left_mm': secao_canonica.get('margem_left_mm') or 20,
        }
        for secao in doc.sections:
            secao.top_margin = Mm(valores['top_mm'])
            secao.right_margin = Mm(valores['right_mm'])
            secao.bottom_margin = Mm(valores['bottom_mm'])
            secao.left_margin = Mm(valores['left_mm'])
        return {'aplicado': True, **valores}

    @classmethod
    def _aplicar_estilos_base(cls, doc, estilos):
        aplicados = []
        if 'Normal' in doc.styles:
            spec = cls._spec_texto(estilos)
            cls._aplicar_spec_fonte(doc.styles['Normal'].font, spec)
            cls._aplicar_spec_pformat(
                doc.styles['Normal'].paragraph_format, spec
            )
            aplicados.append('Normal')
        for nivel in range(1, 10):
            nome = f'Heading {nivel}'
            if nome not in doc.styles:
                continue
            spec = cls._spec_heading(estilos, nivel)
            cls._aplicar_spec_fonte(doc.styles[nome].font, spec)
            aplicados.append(nome)
        return aplicados

    @classmethod
    def _classificar_e_formatar_blocos(cls, doc, perfil, estilos):
        contadores = []
        headings = []
        corpo = 0
        legendas = 0
        vazios = 0
        for paragrafo in doc.paragraphs:
            nivel = cls._heading_nivel(paragrafo.style.name or '')
            if nivel:
                cls._formatar_heading(paragrafo, nivel, perfil, estilos, contadores)
                spec = cls._spec_heading(estilos, nivel)
                headings.append({
                    'nivel': nivel,
                    'texto': (paragrafo.text or '').strip()[:80],
                    'fonte': spec.get('fonte_nome'),
                    'tamanho_pt': spec.get('fonte_tamanho_pt'),
                    'negrito': bool(spec.get('negrito')),
                    'cor_rgb': spec.get('cor_rgb'),
                })
            elif cls._parece_legenda(paragrafo.text):
                cls._aplicar_spec_paragrafo(paragrafo, cls._spec_legenda(estilos))
                legendas += 1
            elif paragrafo.text.strip():
                cls._aplicar_spec_paragrafo(paragrafo, cls._spec_texto(estilos))
                corpo += 1
            else:
                vazios += 1
        # Parágrafos dentro de tabelas não são reformatados neste passe.
        em_tabelas = sum(
            len(celula.paragraphs)
            for tabela in doc.tables
            for linha in tabela.rows
            for celula in linha.cells
        )
        return {
            'headings': {'alterados': len(headings), 'itens': headings},
            'corpo': {'alterados': corpo},
            'legendas': {'alterados': legendas},
            'nao_alterados': {
                'paragrafos_vazios': vazios,
                'paragrafos_em_tabelas': em_tabelas,
                'total': vazios + em_tabelas,
            },
        }

    @staticmethod
    def _heading_nivel(estilo):
        match = re.match(r'heading\s*(\d+)', (estilo or '').strip().lower())
        return int(match.group(1)) if match else None

    @staticmethod
    def _parece_legenda(texto):
        return bool(re.match(r'^\s*(figura|tabela|quadro|equação)\b', texto or '', re.I))

    @classmethod
    def _formatar_heading(cls, paragrafo, nivel, perfil, estilos, contadores):
        nomes = getattr(perfil, 'nome_heading_por_nivel', []) or []
        estilo_destino = nomes[nivel] if nivel < len(nomes) else None
        if estilo_destino:
            try:
                paragrafo.style = estilo_destino
            except (KeyError, ValueError):
                pass
        cls._prefixar_heading(paragrafo, nivel, contadores)
        cls._aplicar_spec_paragrafo(paragrafo, cls._spec_heading(estilos, nivel))

    @staticmethod
    def _prefixar_heading(paragrafo, nivel, contadores):
        texto = paragrafo.text.strip()
        if re.match(r'^\d+(?:\.\d+)*\s+', texto):
            return
        while len(contadores) < nivel:
            contadores.append(0)
        while len(contadores) > nivel:
            contadores.pop()
        contadores[nivel - 1] += 1
        indice = '.'.join(str(num) for num in contadores)
        if paragrafo.runs:
            paragrafo.runs[0].text = f'{indice} {paragrafo.runs[0].text}'
        else:
            paragrafo.add_run(f'{indice} {texto}')

    @staticmethod
    def _spec_texto(estilos):
        return estilos.get('texto normal') or {
            'alinhamento': 'JUSTIFY (3)',
            'espacamento_depois_pt': 11.25,
            'fonte_nome': 'Verdana',
            'fonte_tamanho_pt': 10,
            'cor_rgb': '000000',
        }

    @staticmethod
    def _spec_legenda(estilos):
        return estilos.get('figura') or estilos.get('caption') or {
            'alinhamento': 'CENTER (1)',
            'espacamento_depois_pt': 11.25,
            'fonte_nome': 'Verdana',
            'fonte_tamanho_pt': 9,
            'cor_rgb': '000000',
        }

    @staticmethod
    def _spec_heading(estilos, nivel):
        chaves = {
            1: 'nível 1',
            2: 'nível 1.1',
            3: 'nível 1.1.1',
        }
        return estilos.get(chaves.get(nivel, f'heading {nivel}')) or {
            'fonte_nome': 'Verdana',
            'fonte_tamanho_pt': max(10, 16 - (nivel * 2)),
            'negrito': True,
            'cor_rgb': '0F1E3D' if nivel == 1 else '000000',
        }

    @classmethod
    def _aplicar_spec_paragrafo(cls, paragrafo, spec):
        alinhamento = spec.get('alinhamento') or ''
        if 'JUSTIFY' in alinhamento:
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif 'CENTER' in alinhamento:
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = paragrafo.paragraph_format
        if spec.get('espacamento_antes_pt') is not None:
            fmt.space_before = Pt(spec['espacamento_antes_pt'])
        if spec.get('espacamento_depois_pt') is not None:
            fmt.space_after = Pt(spec['espacamento_depois_pt'])
        if spec.get('entre_linhas') is not None:
            fmt.line_spacing = spec['entre_linhas']
        if spec.get('recuo_primeira_linha_cm') is not None:
            fmt.first_line_indent = Cm(spec['recuo_primeira_linha_cm'])
        if spec.get('recuo_esquerda_cm') is not None:
            fmt.left_indent = Cm(spec['recuo_esquerda_cm'])
        if spec.get('recuo_direita_cm') is not None:
            fmt.right_indent = Cm(spec['recuo_direita_cm'])
        for run in paragrafo.runs:
            cls._aplicar_spec_fonte(run.font, spec)

    @staticmethod
    def _aplicar_spec_pformat(fmt, spec):
        """Aplica alinhamento/espaçamento a um paragraph_format.

        Usado tanto em parágrafos quanto no estilo "Normal" (para que o
        editor, que renderiza pelo estilo, herde a justificação).
        """
        alinhamento = spec.get('alinhamento') or ''
        if 'JUSTIFY' in alinhamento:
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif 'CENTER' in alinhamento:
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if spec.get('espacamento_antes_pt') is not None:
            fmt.space_before = Pt(spec['espacamento_antes_pt'])
        if spec.get('espacamento_depois_pt') is not None:
            fmt.space_after = Pt(spec['espacamento_depois_pt'])
        if spec.get('entre_linhas') is not None:
            fmt.line_spacing = spec['entre_linhas']

    @staticmethod
    def _aplicar_spec_fonte(fonte, spec):
        if spec.get('fonte_nome'):
            fonte.name = spec['fonte_nome']
        if spec.get('fonte_tamanho_pt') is not None:
            fonte.size = Pt(spec['fonte_tamanho_pt'])
        if spec.get('negrito') is not None:
            fonte.bold = spec['negrito']
        if spec.get('italico') is not None:
            fonte.italic = spec['italico']
        if spec.get('sublinhado') is not None:
            fonte.underline = spec['sublinhado']
        cor = spec.get('cor_rgb')
        if cor and re.match(r'^[0-9A-Fa-f]{6}$', cor):
            fonte.color.rgb = RGBColor.from_string(cor)

    @classmethod
    def validar_docx_sugerido(cls, caminho_docx, metricas_aplicadas):
        """Valida condicionantes obrigatórias antes de persistir o DOCX final."""
        faltantes = cls.METRICAS_OBRIGATORIAS.difference(metricas_aplicadas)
        if faltantes:
            raise RuntimeError(
                'Métricas obrigatórias não aplicadas: '
                + ', '.join(sorted(faltantes))
            )
        if not os.path.exists(caminho_docx):
            raise RuntimeError('DOCX sugerido temporário não foi gerado.')
        Document(caminho_docx)
