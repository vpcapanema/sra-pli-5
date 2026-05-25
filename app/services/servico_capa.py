"""Servico de manipulacao da CAPA, FOLHA DE ROSTO e CONTROLE DE
VERSOES dos relatorios DOCX.

================================================================
ESTRUTURA REAL DO TEMPLATE CANONICO
================================================================

A capa dos relatorios PLI-SP nao usa paragrafos de texto comum.
Ela e composta por:

  body[0]  paragrafo vazio
  body[1]  PARAGRAFO MONOLITICO contendo dois objetos flutuantes:
           - <wp:anchor> + <pic:pic>  -> imagem de pagina inteira
             (8.26" x 11.69") com logos e layout decorativo (PNG).
           - <wp:anchor> + <wps:wsp>  -> caixa de texto flutuante
             ("Caixa de Texto 2") com o titulo dinamico:
                "RELATORIO MENSAL - MES {N}"
                "Produtos D-XX e D-YY"
           Tudo termina com <w:sectPr> (quebra de secao).

  body[2]  TABELA "FOLHA DE ROSTO" (9 linhas x 1 coluna).
           Cada linha tem o formato "LABEL | VALOR" (tabs internos).
           Labels canonicos: CODIGO DO DOCUMENTO, TITULO, ELABORACAO,
           CONTRATO, CONTRATACAO, FINANCIAMENTO, OBSERVACOES.

  body[3]  paragrafo separador (vazio)

  body[4]  TABELA "CONTROLE DE VERSOES" (5 linhas x 3 colunas).
           Cabecalho: VERSAO | DATA | CONTEUDO DAS MODIFICACOES
           Linhas seguintes: uma por versao (R00, R01, ...).

  body[5+] Sumario, listas e primeiro Heading 1 do textual.

================================================================
API PUBLICA
================================================================

`extrair_estrutura_capa(doc)` -> dict
    Retorna inventario detalhado da capa: shapes, imagens, tabelas
    com labels. Usado pelo extrator canonico.

`atualizar_capa(caminho, rel, perfil=None)` -> dict
    Substitui o texto do shape "Caixa de Texto 2" pelos dados do
    `RelatorioProducao` atual.

`atualizar_folha_rosto(caminho, rel, perfil=None)` -> dict
    Atualiza valores da Folha de Rosto a partir do mapeamento
    label -> campo do RelatorioProducao.

`atualizar_controle_versoes(caminho, rel, perfil=None)` -> dict
    Adiciona/atualiza linha referente a `rel.versao_atual` na
    tabela de controle de versoes.

`aplicar_dados_completos(caminho, rel, perfil=None)` -> dict
    Conveniencia: executa as 3 atualizacoes em sequencia. Usado
    no fluxo de clonagem.

Todos os atualizadores sao DEFENSIVOS: se nao encontrarem o
elemento esperado, retornam `{'sucesso': False, 'aviso': ...}`
em vez de quebrar. Nunca deixam o DOCX em estado invalido.
"""

from __future__ import annotations

from datetime import date
from typing import Optional

from docx import Document
from lxml import etree


# Namespaces OOXML usados pelos shapes/anchors da capa.
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP_NS = (
    'http://schemas.openxmlformats.org/drawingml/2006/'
    'wordprocessingDrawing'
)
WPS_NS = (
    'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
)
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
R_NS = (
    'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
)


def _q(ns: str, local: str) -> str:
    return f'{{{ns}}}{local}'


# Nome canonico do shape de titulo da capa (apos inspecao do
# template D20-15). Aceita variantes em PT-BR e EN.
NOMES_SHAPE_CAPA = {
    'caixa de texto 2', 'caixa de texto 1',
    'text box 2', 'text box 1',
}

# Labels canonicos da Folha de Rosto. A chave e o LABEL exato
# (case-insensitive, sem acentos) e o valor e o atributo do
# `RelatorioProducao` que produz o valor (ou um callable).
# Veja `_resolver_valor_folha_rosto` para a logica de extracao.
LABELS_FOLHA_ROSTO = (
    'codigo do documento',
    'titulo',
    'elaboracao',
    'contrato',
    'contratacao',
    'financiamento',
    'observacoes',
)

# Cabecalho esperado na tabela de controle de versoes.
CABECALHOS_CONTROLE_VERSOES = ('versao', 'data', 'conteudo')


# =====================================================================
# Helpers internos
# =====================================================================


def _norm(texto: Optional[str]) -> str:
    """Lowercase + remove acentos basicos para comparacao."""
    if not texto:
        return ''
    s = texto.strip().lower()
    repl = {
        'á': 'a', 'à': 'a', 'â': 'a', 'ã': 'a',
        'é': 'e', 'ê': 'e',
        'í': 'i',
        'ó': 'o', 'ô': 'o', 'õ': 'o',
        'ú': 'u',
        'ç': 'c',
    }
    for k, v in repl.items():
        s = s.replace(k, v)
    return s


def _texto_de(elem) -> str:
    """Junta todos os <w:t> descendentes de um elemento."""
    return ''.join(t.text or '' for t in elem.iter(_q(W_NS, 't')))


def _formatar_data_br(d) -> str:
    """Converte date/datetime para 'dd/mm/aaaa'. Aceita None."""
    if d is None:
        return ''
    if isinstance(d, str):
        return d
    try:
        return d.strftime('%d/%m/%Y')
    except (AttributeError, ValueError):
        return str(d)


# =====================================================================
# EXTRACAO da estrutura da capa
# =====================================================================


def _extrair_shape_info(anchor) -> dict:
    """Extrai metadados de um wp:anchor que envolve um drawing.

    Retorna dict com:
      - modo: 'anchor' | 'inline'
      - tipo_conteudo: 'text_box' | 'imagem' | 'desconhecido'
      - nome, descr  (do <wp:docPr>)
      - largura_pol, altura_pol  (do <wp:extent>, em polegadas)
      - posicao: {h: '...', v: '...'} com unidades
      - texto_interno: str (apenas para text_box)
      - rels_embed: id da relacao (apenas para imagem)
    """
    info = {
        'modo': 'anchor' if anchor.tag.endswith('anchor') else 'inline',
        'tipo_conteudo': 'desconhecido',
    }

    docPr = anchor.find(_q(WP_NS, 'docPr'))
    if docPr is not None:
        info['nome'] = docPr.get('name', '')
        info['descr'] = docPr.get('descr', '')

    ext = anchor.find(_q(WP_NS, 'extent'))
    if ext is not None:
        try:
            cx = int(ext.get('cx', 0))
            cy = int(ext.get('cy', 0))
            # EMU -> polegadas (914400 EMU = 1 pol)
            info['largura_pol'] = round(cx / 914400, 3)
            info['altura_pol'] = round(cy / 914400, 3)
        except (ValueError, TypeError):
            pass

    posH = anchor.find(_q(WP_NS, 'positionH'))
    posV = anchor.find(_q(WP_NS, 'positionV'))
    pos = {}
    for nome, p in (('h', posH), ('v', posV)):
        if p is None:
            continue
        rel = p.get('relativeFrom', '?')
        offset = p.find(_q(WP_NS, 'posOffset'))
        align = p.find(_q(WP_NS, 'align'))
        if offset is not None:
            try:
                pos[nome] = {
                    'tipo': 'offset',
                    'valor_pol': round(int(offset.text) / 914400, 3),
                    'relativo_a': rel,
                }
            except (ValueError, TypeError):
                pass
        elif align is not None:
            pos[nome] = {
                'tipo': 'align',
                'valor': align.text,
                'relativo_a': rel,
            }
    if pos:
        info['posicao'] = pos

    # Identifica tipo de conteudo: text box vs imagem
    wsp = anchor.find(f'.//{_q(WPS_NS, "wsp")}')
    pic = anchor.find(f'.//{_q(PIC_NS, "pic")}')
    if wsp is not None:
        info['tipo_conteudo'] = 'text_box'
        info['texto_interno'] = _texto_de(wsp).strip()
    elif pic is not None:
        info['tipo_conteudo'] = 'imagem'
        blip = pic.find(f'.//{_q(A_NS, "blip")}')
        if blip is not None:
            info['rels_embed'] = blip.get(_q(R_NS, 'embed'), '')

    return info


def extrair_estrutura_capa(doc) -> dict:
    """Inventario detalhado da CAPA (parte do body antes do primeiro
    Heading 1).

    Retorna dict com chaves:
      - paragrafos_capa: lista de descritores curtos dos paragrafos
        que compoem a capa (indice, estilo, texto resumido, flags)
      - shapes: lista de text boxes flutuantes (com texto interno)
      - imagens: lista de imagens flutuantes
      - tabela_folha_rosto: {indice, linhas: [(label, valor), ...]}
        ou None se nao detectada
      - tabela_controle_versoes: {indice, cabecalho, linhas} ou None
      - sumario_nativo_word: bool (se ha <w:sdt> de TOC)
      - secoes_indices: indices dos paragrafos onde ha <w:sectPr>
    """
    body = doc.element.body

    estrutura = {
        'paragrafos_capa': [],
        'shapes': [],
        'imagens': [],
        'tabela_folha_rosto': None,
        'tabela_controle_versoes': None,
        'sumario_nativo_word': False,
        'secoes_indices': [],
    }

    # Encontrar indice do primeiro Heading 1 — limita o escopo da capa
    fim_capa_idx = None
    for i, child in enumerate(body):
        if child.tag != _q(W_NS, 'p'):
            continue
        pPr = child.find(_q(W_NS, 'pPr'))
        if pPr is None:
            continue
        pStyle = pPr.find(_q(W_NS, 'pStyle'))
        if pStyle is None:
            continue
        val = pStyle.get(_q(W_NS, 'val'), '').lower()
        if ('eading' in val or 'tulo' in val) and '1' in val:
            fim_capa_idx = i
            break

    if fim_capa_idx is None:
        fim_capa_idx = len(body)

    estrutura['indice_fim_capa'] = fim_capa_idx

    # Percorrer os elementos da capa
    for i, child in enumerate(body[:fim_capa_idx]):
        tag = etree.QName(child.tag).localname

        if tag == 'p':
            # Descritor curto
            texto = _texto_de(child).strip()
            estilo = '(default)'
            pPr = child.find(_q(W_NS, 'pPr'))
            if pPr is not None:
                ps = pPr.find(_q(W_NS, 'pStyle'))
                if ps is not None:
                    estilo = ps.get(_q(W_NS, 'val'), '?')
                if pPr.find(_q(W_NS, 'sectPr')) is not None:
                    estrutura['secoes_indices'].append(i)

            descr = {
                'indice': i,
                'estilo': estilo,
                'texto': (texto[:120] + '…') if len(texto) > 120 else texto,
            }

            # Shapes / imagens dentro deste paragrafo
            anchors = list(child.iter(_q(WP_NS, 'anchor')))
            inlines = list(child.iter(_q(WP_NS, 'inline')))
            if anchors or inlines:
                descr['flutuantes'] = len(anchors) + len(inlines)
                for anc in anchors + inlines:
                    sinfo = _extrair_shape_info(anc)
                    sinfo['paragrafo_indice'] = i
                    if sinfo['tipo_conteudo'] == 'text_box':
                        estrutura['shapes'].append(sinfo)
                    elif sinfo['tipo_conteudo'] == 'imagem':
                        estrutura['imagens'].append(sinfo)

            # SDT (Sumario nativo do Word)
            if child.find(f'.//{_q(W_NS, "sdt")}') is not None:
                estrutura['sumario_nativo_word'] = True
                descr['tem_sdt'] = True

            estrutura['paragrafos_capa'].append(descr)

        elif tag == 'tbl':
            # Identificar se e Folha de Rosto ou Controle de Versoes
            tabela_info = _classificar_tabela_capa(child, i)
            if tabela_info['tipo'] == 'folha_rosto':
                estrutura['tabela_folha_rosto'] = tabela_info
            elif tabela_info['tipo'] == 'controle_versoes':
                estrutura['tabela_controle_versoes'] = tabela_info
            estrutura['paragrafos_capa'].append({
                'indice': i, 'tipo_elem': 'tbl',
                'classificacao': tabela_info['tipo'],
                'dimensoes': tabela_info.get('dimensoes', '?'),
            })

        elif tag == 'sdt':
            estrutura['sumario_nativo_word'] = True
            estrutura['paragrafos_capa'].append({
                'indice': i, 'tipo_elem': 'sdt',
            })

    return estrutura


def _classificar_tabela_capa(tbl, indice_no_body: int) -> dict:
    """Decide se a tabela e Folha de Rosto ou Controle de Versoes
    olhando para o conteudo das primeiras celulas.

    Retorna dict com chaves:
      - tipo: 'folha_rosto' | 'controle_versoes' | 'desconhecido'
      - indice_no_body
      - dimensoes: 'NxM' (linhas x colunas)
      - linhas: lista de tuplas/dicts com conteudo extraido
    """
    rows = tbl.findall(_q(W_NS, 'tr'))
    n_rows = len(rows)
    n_cols = len(rows[0].findall(_q(W_NS, 'tc'))) if rows else 0

    info = {
        'indice_no_body': indice_no_body,
        'dimensoes': f'{n_rows}x{n_cols}',
        'tipo': 'desconhecido',
        'linhas': [],
    }

    if not rows:
        return info

    # Coletar conteudo bruto de cada celula
    linhas_brutas = []
    for tr in rows:
        cells_texts = []
        for tc in tr.findall(_q(W_NS, 'tc')):
            txt = ' | '.join(
                _texto_de(p).strip()
                for p in tc.findall(_q(W_NS, 'p'))
                if _texto_de(p).strip()
            )
            cells_texts.append(txt)
        linhas_brutas.append(cells_texts)

    # Heuristica de classificacao
    # - Controle de Versoes: cabecalho tem "VERSAO", "DATA", "CONTEUDO"
    # - Folha de Rosto: primeiras linhas tem labels canonicos
    cab = linhas_brutas[0] if linhas_brutas else []
    cab_concat = _norm(' '.join(cab))
    if all(p in cab_concat for p in CABECALHOS_CONTROLE_VERSOES):
        info['tipo'] = 'controle_versoes'
        info['cabecalho'] = cab
        info['linhas'] = [
            {
                'indice_linha': li,
                'versao': (cells[0] if len(cells) > 0 else ''),
                'data': (cells[1] if len(cells) > 1 else ''),
                'conteudo': (cells[2] if len(cells) > 2 else ''),
            }
            for li, cells in enumerate(linhas_brutas[1:], start=1)
        ]
        return info

    # Folha de Rosto: ao menos uma linha cuja primeira parte (antes
    # do "|" ou tab) bate com um label conhecido.
    rosto_linhas = []
    matches_rosto = 0
    for li, cells in enumerate(linhas_brutas):
        # Concat de todas as celulas porque a tabela do template tem
        # 1 coluna com tabs internos (labels e valores juntos).
        bruto = ' '.join(cells).strip()
        # Tentar separar por "|" (vindo dos tabs)
        if '|' in bruto:
            partes = [p.strip() for p in bruto.split('|') if p.strip()]
        else:
            partes = [bruto] if bruto else []

        label_raw = partes[0] if partes else ''
        valor_raw = ' | '.join(partes[1:]) if len(partes) > 1 else ''
        label_n = _norm(label_raw)

        eh_label_canonico = any(
            label_n == lbl or label_n.startswith(lbl)
            for lbl in LABELS_FOLHA_ROSTO
        )
        if eh_label_canonico:
            matches_rosto += 1

        rosto_linhas.append({
            'indice_linha': li,
            'label_bruto': label_raw,
            'label_normalizado': label_n,
            'valor': valor_raw,
            'eh_label_canonico': eh_label_canonico,
        })

    if matches_rosto >= 2:
        info['tipo'] = 'folha_rosto'
        info['linhas'] = rosto_linhas

    return info


# =====================================================================
# ATUALIZADORES
# =====================================================================


def atualizar_capa(
    caminho_docx: str, rel, perfil=None
) -> dict:
    """Atualiza o texto do shape "Caixa de Texto" da capa com dados
    do `RelatorioProducao` fornecido.

    Estrategia:
    1. Localiza o `<wps:wsp>` cujo `<wp:docPr name>` esta em
       `NOMES_SHAPE_CAPA` (case-insensitive, sem acento).
    2. Dentro dele, localiza `<w:txbxContent>` e substitui o texto
       dos paragrafos preservando o estilo do primeiro run de cada
       paragrafo (fonte, tamanho, cor).
    3. Conteudo gerado (2 linhas, ABNT):
         linha 1: "RELATÓRIO MENSAL – MÊS {numero_medicao}"
         linha 2: "Produtos {codigo_d20}{ano_referencia}"

    Retorna dict {'sucesso': bool, 'shape_atualizado': str | None,
                  'avisos': [str]}.
    """
    doc = Document(caminho_docx)
    body = doc.element.body
    avisos = []

    shape_alvo = None
    nome_encontrado = None
    for anchor in body.iter(_q(WP_NS, 'anchor')):
        docPr = anchor.find(_q(WP_NS, 'docPr'))
        if docPr is None:
            continue
        nome = (docPr.get('name') or '').strip()
        if _norm(nome) in NOMES_SHAPE_CAPA:
            wsp = anchor.find(f'.//{_q(WPS_NS, "wsp")}')
            if wsp is not None:
                shape_alvo = wsp
                nome_encontrado = nome
                break

    if shape_alvo is None:
        return {
            'sucesso': False,
            'shape_atualizado': None,
            'avisos': [
                'Shape de capa nao encontrado. Esperado <wps:wsp> '
                f'com docPr em {sorted(NOMES_SHAPE_CAPA)}. '
                'Capa permanece inalterada.'
            ],
        }

    # Linhas a inserir
    num = rel.numero_medicao or '?'
    codigo = rel.codigo_d20 or 'D-XX'
    linha1 = f'RELATÓRIO MENSAL – MÊS {num}'
    linha2 = f'Produtos {codigo}'

    # Substituir conteudo do txbxContent
    txbx = shape_alvo.find(f'.//{_q(W_NS, "txbxContent")}')
    if txbx is None:
        return {
            'sucesso': False,
            'shape_atualizado': nome_encontrado,
            'avisos': [
                'Shape encontrado mas sem <w:txbxContent>. Estrutura '
                'inesperada — capa inalterada.'
            ],
        }

    paragrafos_existentes = txbx.findall(_q(W_NS, 'p'))
    if not paragrafos_existentes:
        avisos.append('Shape vazio: txbxContent nao tinha paragrafos.')
        # Cria paragrafo zerado a partir do template minimo
        novo_p = etree.SubElement(txbx, _q(W_NS, 'p'))
        paragrafos_existentes = [novo_p]

    # Capturar pPr e rPr do primeiro paragrafo para preservar formatacao
    rPr_modelo = None
    pPr_modelo = paragrafos_existentes[0].find(_q(W_NS, 'pPr'))
    for r in paragrafos_existentes[0].findall(_q(W_NS, 'r')):
        rpr = r.find(_q(W_NS, 'rPr'))
        if rpr is not None:
            rPr_modelo = rpr
            break

    # Remover paragrafos antigos
    for p in paragrafos_existentes:
        txbx.remove(p)

    # Criar 2 novos paragrafos (linha1 e linha2)
    for linha_texto in (linha1, linha2):
        p_novo = etree.SubElement(txbx, _q(W_NS, 'p'))
        if pPr_modelo is not None:
            p_novo.append(etree.fromstring(etree.tostring(pPr_modelo)))
        r_novo = etree.SubElement(p_novo, _q(W_NS, 'r'))
        if rPr_modelo is not None:
            r_novo.append(etree.fromstring(etree.tostring(rPr_modelo)))
        t_novo = etree.SubElement(r_novo, _q(W_NS, 't'))
        t_novo.set(
            '{http://www.w3.org/XML/1998/namespace}space', 'preserve'
        )
        t_novo.text = linha_texto

    doc.save(caminho_docx)
    return {
        'sucesso': True,
        'shape_atualizado': nome_encontrado,
        'linhas': [linha1, linha2],
        'avisos': avisos,
    }


def atualizar_folha_rosto(
    caminho_docx: str, rel, perfil=None
) -> dict:
    """Atualiza a tabela de Folha de Rosto com dados do relatorio.

    Para cada linha cuja primeira celula bate com um label canonico
    (CODIGO DO DOCUMENTO, TITULO, ELABORACAO, ...), substitui o
    VALOR pela informacao atual do `RelatorioProducao`.

    Linhas com labels nao canonicos sao preservadas intactas
    (ex.: CONTRATACAO, FINANCIAMENTO costumam ser fixos do
    contrato e ja vem corretos no template).
    """
    doc = Document(caminho_docx)
    body = doc.element.body
    avisos = []

    # Localizar a tabela de folha de rosto (primeira tbl que classifica)
    tbl_alvo = None
    for child in body:
        if child.tag != _q(W_NS, 'tbl'):
            continue
        info = _classificar_tabela_capa(child, -1)
        if info['tipo'] == 'folha_rosto':
            tbl_alvo = child
            break

    if tbl_alvo is None:
        return {
            'sucesso': False,
            'avisos': ['Tabela de Folha de Rosto nao encontrada.'],
        }

    # Valores a aplicar (None significa nao mexer)
    valores = _construir_valores_folha_rosto(rel)
    linhas_alteradas = []

    rows = tbl_alvo.findall(_q(W_NS, 'tr'))
    for tr in rows:
        # Para o template canonico (1 coluna com tabs internos), todas
        # as celulas dessa linha contem o par "LABEL | VALOR". Mas pode
        # haver tabelas estruturadas em 2 colunas reais. Tratamos ambos.
        tcs = tr.findall(_q(W_NS, 'tc'))
        if not tcs:
            continue

        # Texto bruto da linha (concatenando celulas)
        texto_bruto = ' | '.join(
            _texto_de(p).strip()
            for tc in tcs
            for p in tc.findall(_q(W_NS, 'p'))
            if _texto_de(p).strip()
        )
        if '|' in texto_bruto:
            partes = [p.strip() for p in texto_bruto.split('|')]
        else:
            partes = [texto_bruto]

        if not partes:
            continue

        label_n = _norm(partes[0])
        # Buscar match
        chave_label = None
        for lbl in LABELS_FOLHA_ROSTO:
            if label_n == lbl or label_n.startswith(lbl):
                chave_label = lbl
                break
        if chave_label is None:
            continue
        if chave_label not in valores or valores[chave_label] is None:
            continue

        novo_valor = valores[chave_label]
        # Aplicar: substitui o texto da SEGUNDA celula (se houver) ou
        # reescreve a celula unica preservando o label.
        if len(tcs) >= 2:
            _substituir_texto_celula(tcs[1], novo_valor)
        else:
            # Celula unica com tab interno — recriar paragrafos com
            # label + tab + valor
            _substituir_celula_com_label(
                tcs[0], partes[0], novo_valor,
            )
        linhas_alteradas.append({
            'label': chave_label,
            'novo_valor': novo_valor,
        })

    doc.save(caminho_docx)
    return {
        'sucesso': True,
        'linhas_alteradas': linhas_alteradas,
        'avisos': avisos,
    }


def atualizar_controle_versoes(
    caminho_docx: str, rel, perfil=None,
    descricao_modificacoes: Optional[str] = None,
) -> dict:
    """Adiciona ou atualiza a linha referente a `rel.versao_atual` na
    tabela de Controle de Versoes.

    Estrategia:
    1. Procurar linha cuja primeira celula seja igual a
       `rel.versao_atual` (ex.: 'R00'). Se existir, atualiza DATA
       e CONTEUDO.
    2. Senao, procura a primeira linha VAZIA (apos o cabecalho) e
       preenche com versao/data/conteudo.
    3. Se nao houver linha vazia, adiciona uma nova linha clonando
       a ultima preenchida.

    `descricao_modificacoes` default = 'Versao inicial' se for R00,
    senao 'Atualizacao automatica'.
    """
    doc = Document(caminho_docx)
    body = doc.element.body
    avisos = []

    tbl_alvo = None
    for child in body:
        if child.tag != _q(W_NS, 'tbl'):
            continue
        info = _classificar_tabela_capa(child, -1)
        if info['tipo'] == 'controle_versoes':
            tbl_alvo = child
            break

    if tbl_alvo is None:
        return {
            'sucesso': False,
            'avisos': ['Tabela de Controle de Versoes nao encontrada.'],
        }

    versao = (rel.versao_atual or 'R00').strip()
    data_str = _formatar_data_br(
        getattr(rel, 'atualizado_em', None)
        or getattr(rel, 'criado_em', None)
        or date.today()
    )
    if descricao_modificacoes is None:
        descricao_modificacoes = (
            'Versão inicial'
            if versao.upper() == 'R00'
            else f'Atualização automática ({versao})'
        )

    rows = tbl_alvo.findall(_q(W_NS, 'tr'))
    # Linha 0 = cabecalho; demais sao versoes
    linha_existente = None
    primeira_vazia = None
    for li, tr in enumerate(rows[1:], start=1):
        tcs = tr.findall(_q(W_NS, 'tc'))
        if not tcs:
            continue
        c0 = _texto_de(tcs[0]).strip()
        c1 = _texto_de(tcs[1]).strip() if len(tcs) > 1 else ''
        c2 = _texto_de(tcs[2]).strip() if len(tcs) > 2 else ''
        if c0 == versao:
            linha_existente = (li, tr, tcs)
            break
        if primeira_vazia is None and not (c0 or c1 or c2):
            primeira_vazia = (li, tr, tcs)

    if linha_existente is not None:
        _, _, tcs = linha_existente
        _substituir_texto_celula(tcs[0], versao)
        if len(tcs) > 1:
            _substituir_texto_celula(tcs[1], data_str)
        if len(tcs) > 2:
            _substituir_texto_celula(
                tcs[2], descricao_modificacoes,
            )
        acao = f'linha {versao} atualizada'
    elif primeira_vazia is not None:
        _, _, tcs = primeira_vazia
        _substituir_texto_celula(tcs[0], versao)
        if len(tcs) > 1:
            _substituir_texto_celula(tcs[1], data_str)
        if len(tcs) > 2:
            _substituir_texto_celula(
                tcs[2], descricao_modificacoes,
            )
        acao = f'linha {versao} criada em slot vazio'
    else:
        # Clonar a ultima linha como template e preencher
        if len(rows) < 2:
            return {
                'sucesso': False,
                'avisos': [
                    'Tabela de versoes sem linhas de dados — '
                    'estrutura inesperada.'
                ],
            }
        ultima = rows[-1]
        nova = etree.fromstring(etree.tostring(ultima))
        ultima.addnext(nova)
        tcs = nova.findall(_q(W_NS, 'tc'))
        _substituir_texto_celula(tcs[0], versao)
        if len(tcs) > 1:
            _substituir_texto_celula(tcs[1], data_str)
        if len(tcs) > 2:
            _substituir_texto_celula(
                tcs[2], descricao_modificacoes,
            )
        acao = f'linha {versao} adicionada ao final'

    doc.save(caminho_docx)
    return {
        'sucesso': True,
        'acao': acao,
        'versao': versao,
        'data': data_str,
        'descricao': descricao_modificacoes,
        'avisos': avisos,
    }


def aplicar_dados_completos(
    caminho_docx: str, rel, perfil=None,
) -> dict:
    """Conveniencia: aplica capa + folha de rosto + controle de
    versoes em sequencia. Usado no fluxo de clonagem para
    personalizar imediatamente o DOCX recem-clonado.

    Retorna dict com os 3 sub-resultados.
    """
    return {
        'capa': atualizar_capa(caminho_docx, rel, perfil=perfil),
        'folha_rosto': atualizar_folha_rosto(
            caminho_docx, rel, perfil=perfil,
        ),
        'controle_versoes': atualizar_controle_versoes(
            caminho_docx, rel, perfil=perfil,
        ),
    }


# =====================================================================
# Helpers de mutacao DOCX
# =====================================================================


def _substituir_texto_celula(tc, novo_texto: str) -> None:
    """Substitui o texto de uma <w:tc>, preservando o pPr/rPr do
    primeiro paragrafo/run para nao quebrar formatacao.

    Estrategia: mantem o primeiro <w:p> da celula, limpa runs antigos,
    cria um run novo com o texto. Remove paragrafos extras.
    """
    paragrafos = tc.findall(_q(W_NS, 'p'))
    if not paragrafos:
        # Celula sem paragrafo — criar um
        p = etree.SubElement(tc, _q(W_NS, 'p'))
        paragrafos = [p]

    primeiro_p = paragrafos[0]
    # Capturar rPr modelo
    rPr_modelo = None
    for r in primeiro_p.findall(_q(W_NS, 'r')):
        rpr = r.find(_q(W_NS, 'rPr'))
        if rpr is not None:
            rPr_modelo = rpr
            break

    # Remover todos os runs e hyperlinks existentes do primeiro paragrafo
    for filho in list(primeiro_p):
        tag = etree.QName(filho.tag).localname
        if tag in ('r', 'hyperlink', 'fldSimple'):
            primeiro_p.remove(filho)

    # Criar run novo com o texto
    r_novo = etree.SubElement(primeiro_p, _q(W_NS, 'r'))
    if rPr_modelo is not None:
        r_novo.append(etree.fromstring(etree.tostring(rPr_modelo)))
    t_novo = etree.SubElement(r_novo, _q(W_NS, 't'))
    t_novo.set(
        '{http://www.w3.org/XML/1998/namespace}space', 'preserve'
    )
    t_novo.text = novo_texto

    # Remover paragrafos extras (mantemos so o primeiro)
    for p_extra in paragrafos[1:]:
        tc.remove(p_extra)


def _substituir_celula_com_label(
    tc, label_original: str, novo_valor: str,
) -> None:
    """Atualiza o VALOR em uma celula que combina label + valor.

    Templates canonicos podem usar duas convencoes diferentes:

    (A) **Paragrafos separados** (forma mais comum no D20-15):
        <w:tc>
          <w:p>CÓDIGO DO DOCUMENTO</w:p>
          <w:p>D20-15</w:p>     <- substituimos APENAS este
        </w:tc>

    (B) **Mesmo paragrafo com tab interno**:
        <w:tc>
          <w:p>
            <w:r><w:t>CÓDIGO...</w:t></w:r>
            <w:r><w:tab/></w:r>
            <w:r><w:t>D20-15</w:t></w:r>  <- substituimos APENAS este
          </w:p>
        </w:tc>

    A funcao detecta qual caso aplica olhando se ha 2+ paragrafos
    com texto na celula. Se sim, substitui o texto do paragrafo do
    VALOR (todos exceto o primeiro). Se nao, recria o run de valor
    apos o tab dentro do paragrafo unico.

    Em todos os casos preserva pPr e rPr do paragrafo/run original.
    """
    paragrafos_texto = [
        p for p in tc.findall(_q(W_NS, 'p'))
        if _texto_de(p).strip()
    ]

    # Caso A: label e valor em paragrafos separados
    if len(paragrafos_texto) >= 2:
        p_valor = paragrafos_texto[1]
        _substituir_runs_paragrafo(p_valor, novo_valor)
        return

    # Caso B: tudo num so paragrafo (raro). Reescreve preservando
    # label como primeiro run, tab no segundo, valor no terceiro.
    if not paragrafos_texto:
        return
    primeiro_p = paragrafos_texto[0]
    rPr_modelo = None
    for r in primeiro_p.findall(_q(W_NS, 'r')):
        rpr = r.find(_q(W_NS, 'rPr'))
        if rpr is not None:
            rPr_modelo = rpr
            break

    for filho in list(primeiro_p):
        tag = etree.QName(filho.tag).localname
        if tag in ('r', 'hyperlink', 'fldSimple'):
            primeiro_p.remove(filho)

    r1 = etree.SubElement(primeiro_p, _q(W_NS, 'r'))
    if rPr_modelo is not None:
        r1.append(etree.fromstring(etree.tostring(rPr_modelo)))
    t1 = etree.SubElement(r1, _q(W_NS, 't'))
    t1.set(
        '{http://www.w3.org/XML/1998/namespace}space', 'preserve'
    )
    t1.text = label_original

    r2 = etree.SubElement(primeiro_p, _q(W_NS, 'r'))
    if rPr_modelo is not None:
        r2.append(etree.fromstring(etree.tostring(rPr_modelo)))
    etree.SubElement(r2, _q(W_NS, 'tab'))
    t2 = etree.SubElement(r2, _q(W_NS, 't'))
    t2.set(
        '{http://www.w3.org/XML/1998/namespace}space', 'preserve'
    )
    t2.text = novo_valor


def _substituir_runs_paragrafo(p, novo_texto: str) -> None:
    """Substitui o texto de TODOS os runs de um <w:p> por `novo_texto`,
    preservando o rPr do primeiro run original (estilo).
    """
    rPr_modelo = None
    for r in p.findall(_q(W_NS, 'r')):
        rpr = r.find(_q(W_NS, 'rPr'))
        if rpr is not None:
            rPr_modelo = rpr
            break

    for filho in list(p):
        tag = etree.QName(filho.tag).localname
        if tag in ('r', 'hyperlink', 'fldSimple'):
            p.remove(filho)

    r_novo = etree.SubElement(p, _q(W_NS, 'r'))
    if rPr_modelo is not None:
        r_novo.append(etree.fromstring(etree.tostring(rPr_modelo)))
    t_novo = etree.SubElement(r_novo, _q(W_NS, 't'))
    t_novo.set(
        '{http://www.w3.org/XML/1998/namespace}space', 'preserve'
    )
    t_novo.text = novo_texto


# =====================================================================
# Mapeamento label -> valor
# =====================================================================


def _construir_valores_folha_rosto(rel) -> dict:
    """Constroi dicionario {label_canonico: valor_a_aplicar} a partir
    do `RelatorioProducao`.

    Labels nao listados aqui nao serao tocados (preservam o valor
    original do template — caso de CONTRATACAO, FINANCIAMENTO).
    """
    codigo = rel.codigo_d20 or 'D-XX'
    numero = rel.numero_medicao
    # Codigo final usado nas etiquetas: "D20-15"
    codigo_etiqueta = (
        f'{codigo}-{numero}' if numero else codigo
    )

    titulo = (
        rel.titulo_curto
        or f'RELATÓRIO MENSAL – MÊS {numero or "?"}'
    )

    valores = {
        'codigo do documento': codigo_etiqueta,
        'titulo': titulo,
        # 'elaboracao', 'contrato', 'contratacao', 'financiamento'
        # nao sao tocados — permanecem como no template.
    }

    obs = _construir_observacoes(rel)
    if obs:
        valores['observacoes'] = obs

    return valores


def _construir_observacoes(rel) -> str:
    """Texto padrao para o campo OBSERVACOES da Folha de Rosto."""
    codigo = rel.codigo_d20 or 'D-XX'
    numero = rel.numero_medicao
    if numero:
        return (
            f'Este Relatório corresponde ao entregável '
            f'{codigo}-{numero}.'
        )
    return f'Este Relatório corresponde ao entregável {codigo}.'
