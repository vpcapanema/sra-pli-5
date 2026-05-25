"""
API REST para o fluxo de edição colaborativa.
Consumida pelo editor React (docx-editor) embedado nas telas
de autor e coordenador.
"""

import os
import secrets
import time
from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document
from flask import Blueprint, abort, jsonify, request, send_file, session
from flask_login import login_required, current_user

from app import db
from app.models.usuario import Usuario
from app.models.relatorio_producao import RelatorioProducao
from app.models.capitulo_documento import CapituloDocumento
from app.models.biblioteca_formatacao import (
    BibliotecaFormatacaoCanonica
)
from app.models.notificacao import Notificacao
# servico_motor_renderizacao removido pos Fase 1: o DOCX em producao
# (caminho_template) e a fonte unica; nao ha mais reconstrucao a
# partir de capitulos + conteudo_docx.
from app.utils.logger import sra_log_handler

MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB

# Rate limiting simples em memória
_rate_limit_store = {}
RATE_LIMIT_MAX = 60  # requests
RATE_LIMIT_WINDOW = 60  # segundos

api_bp = Blueprint('api', __name__, url_prefix='/api')


@api_bp.before_request
def _rate_limit():
    """Rate limiting por IP — 60 req/min."""
    ip = request.remote_addr or '0.0.0.0'
    now = time.time()
    entry = _rate_limit_store.get(ip)
    if not entry or now - entry['start'] > RATE_LIMIT_WINDOW:
        _rate_limit_store[ip] = {'start': now, 'count': 1}
    else:
        entry['count'] += 1
        if entry['count'] > RATE_LIMIT_MAX:
            abort(429)


@api_bp.before_request
def _verificar_csrf():
    """Verifica CSRF token para requests mutantes (POST/PUT/PATCH/DELETE)."""
    if request.method in ('POST', 'PUT', 'PATCH', 'DELETE'):
        token = request.headers.get('X-CSRF-Token', '')
        if token != session.get('csrf_token', ''):
            if 'csrf_token' not in session:
                session['csrf_token'] = secrets.token_hex(32)
                return
            abort(403)


@api_bp.route('/csrf-token')
@login_required
def obter_csrf_token():
    """Retorna CSRF token para uso pelo frontend."""
    if 'csrf_token' not in session:
        session['csrf_token'] = secrets.token_hex(32)
    return jsonify({'token': session['csrf_token']})


# ==============================================================
# VERSÃO DE TRABALHO
# ==============================================================

@api_bp.route('/versoes-trabalho/<int:id_vt>')
@login_required
def obter_versao(id_vt):
    """Retorna dados da versão de trabalho com capítulos."""
    vt = RelatorioProducao.query.get_or_404(id_vt)
    caps = CapituloDocumento.query.filter_by(
        id_relatorio=id_vt,
        id_capitulo_pai=None,
        ativo=True
    ).order_by(CapituloDocumento.ordem_capitulo).all()

    return jsonify({
        'id': vt.id,
        'titulo': vt.titulo_curto,
        'status': vt.status.codigo if vt.status else None,
        'id_biblioteca': None,
        'capitulos': [_serializar_capitulo(c) for c in caps],
    })


# ==============================================================
# CAPÍTULOS — CRUD
# ==============================================================

@api_bp.route('/versoes-trabalho/<int:id_vt>/capitulos')
@login_required
def listar_capitulos(id_vt):
    """Lista capítulos (raízes) da versão de trabalho."""
    caps = CapituloDocumento.query.filter_by(
        id_relatorio=id_vt,
        id_capitulo_pai=None,
        ativo=True
    ).order_by(CapituloDocumento.ordem_capitulo).all()
    return jsonify([_serializar_capitulo(c) for c in caps])


@api_bp.route(
    '/versoes-trabalho/<int:id_vt>/capitulos',
    methods=['POST']
)
@login_required
def criar_capitulo(id_vt):
    """Cria um novo capítulo/subcapítulo."""
    RelatorioProducao.query.get_or_404(id_vt)
    data = request.get_json()
    cap = CapituloDocumento(
        id_relatorio=id_vt,
        titulo_capitulo=data['titulo'],
        ordem_capitulo=data.get('ordem', 0),
        nivel_capitulo=data.get('nivel', 1),
        id_capitulo_pai=data.get('id_pai'),
        id_usuario_responsavel=data.get('id_responsavel'),
        status_capitulo='em_edicao',
    )
    db.session.add(cap)
    db.session.commit()
    return jsonify(_serializar_capitulo(cap)), 201


@api_bp.route('/capitulos/<int:id_cap>', methods=['PATCH'])
@login_required
def atualizar_capitulo(id_cap):
    """Atualiza metadados de um capítulo (coordenador only)."""
    _exigir_perfil('coordenador')
    cap = CapituloDocumento.query.get_or_404(id_cap)
    data = request.get_json()

    if 'titulo' in data:
        cap.titulo_capitulo = data['titulo']
    if 'ordem' in data:
        cap.ordem_capitulo = data['ordem']
    if 'nivel' in data:
        cap.nivel_capitulo = data['nivel']
    if 'id_responsavel' in data:
        cap.id_usuario_responsavel = data['id_responsavel']

    db.session.commit()
    return jsonify(_serializar_capitulo(cap))


@api_bp.route('/capitulos/<int:id_cap>', methods=['DELETE'])
@login_required
def excluir_capitulo(id_cap):
    """Soft-delete de um capítulo."""
    cap = CapituloDocumento.query.get_or_404(id_cap)
    cap.ativo = False
    db.session.commit()
    return '', 204


# ==============================================================
# ENVIOS DO AUTOR — blob DOCX original e segmentado por capítulo
# ==============================================================

@api_bp.route('/envios/<int:id_envio>/docx')
@login_required
def baixar_envio_docx(id_envio):
    """Serve o DOCX original do envio para visualização no editor.

    Acessível para o próprio autor que enviou e para coordenadores.
    """
    from app.models.envio_conteudo import EnvioConteudo
    envio = EnvioConteudo.query.get_or_404(id_envio)
    if (envio.id_usuario != current_user.id
            and session.get('perfil_ativo') not in ('coordenador', 'admin')):
        return jsonify({'erro': 'Sem permissão'}), 403
    if not envio.caminho_arquivo or not os.path.exists(envio.caminho_arquivo):
        return ('', 204)
    return send_file(
        envio.caminho_arquivo,
        mimetype=(
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.document'
        ),
        as_attachment=False,
        download_name=envio.nome_arquivo or f'envio_{id_envio}.docx',
    )


@api_bp.route('/envios/<int:id_envio>/capitulos/<int:id_capitulo>/docx')
@login_required
def baixar_envio_segmento_docx(id_envio, id_capitulo):
    """Serve o DOCX segmentado por capítulo para um envio.

    Gerado em memória a partir do DOCX original do envio + classificação:
    parte do DOCX que foi atribuída ao capítulo informado.
    """
    from app.models.envio_conteudo import EnvioConteudo
    from app.services.servico_envio_autor import gerar_docx_segmento
    envio = EnvioConteudo.query.get_or_404(id_envio)
    if (envio.id_usuario != current_user.id
            and session.get('perfil_ativo') not in ('coordenador', 'admin')):
        return jsonify({'erro': 'Sem permissão'}), 403
    cap = CapituloDocumento.query.get_or_404(id_capitulo)
    if cap.id_relatorio != envio.id_relatorio:
        return jsonify({'erro': 'Capítulo não pertence ao envio'}), 400
    try:
        blob = gerar_docx_segmento(envio, cap)
    except (OSError, ValueError, RuntimeError) as e:
        return jsonify({'erro': f'Falha ao gerar segmento: {e}'}), 500
    if not blob:
        return ('', 204)
    return send_file(
        BytesIO(blob),
        mimetype=(
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.document'
        ),
        as_attachment=False,
        download_name=f'envio_{id_envio}_cap_{id_capitulo}.docx',
    )


@api_bp.route(
    '/envios/<int:id_envio>/capitulos/<int:id_capitulo>/docx',
    methods=['PUT']
)
@login_required
def salvar_envio_segmento_docx(id_envio, id_capitulo):
    """Salva o DOCX editado (pelo editor no browser) como o segmento
    confirmado daquele capítulo no envio.

    Aceita:
    - application/octet-stream: bytes DOCX
    - text/html: HTML do contenteditable; convertido para DOCX
    """
    from app.models.envio_conteudo import EnvioConteudo
    envio = EnvioConteudo.query.get_or_404(id_envio)
    if envio.id_usuario != current_user.id:
        return jsonify({'erro': 'Sem permissão'}), 403
    cap = CapituloDocumento.query.get_or_404(id_capitulo)
    if cap.id_relatorio != envio.id_relatorio:
        return jsonify({'erro': 'Capítulo não pertence ao envio'}), 400

    content_type = request.content_type or ''
    dados = request.get_data()
    if len(dados) > MAX_UPLOAD_BYTES:
        return jsonify({'erro': 'Arquivo excede 50 MB'}), 413

    if 'text/html' in content_type:
        try:
            soup = BeautifulSoup(dados.decode('utf-8'), 'html.parser')
            doc = Document()
            for el in soup.find_all(
                ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']
            ):
                txt = el.get_text(strip=True)
                if not txt:
                    continue
                if el.name.startswith('h'):
                    nivel = int(el.name[1])
                    doc.add_heading(txt, level=min(nivel, 9))
                else:
                    doc.add_paragraph(txt)
            buf = BytesIO()
            doc.save(buf)
            dados = buf.getvalue()
        except (ValueError, TypeError, AttributeError) as e:
            return jsonify(
                {'erro': f'Erro ao converter HTML para DOCX: {e}'}
            ), 400

    # DEPRECATED: pos-Fase 1, edicoes inline foram substituidas por
    # upload + merge in-place (servico_merge_docx). Este endpoint nao
    # persiste mais nada. Mantido por compatibilidade do frontend ate
    # que `editor_autor.js` e `editor_coordenador.js` migrem para o
    # fluxo de upload completo.
    return jsonify({
        'erro': (
            'Endpoint depreciado. Use o fluxo de upload de DOCX por '
            'capitulo (rota /relatorio/.../capitulo/<id>/upload).'
        ),
    }), 410


# ==============================================================
# CONTEÚDO DOCX — upload/download por capítulo
# ==============================================================

@api_bp.route('/capitulos/<int:id_cap>/conteudo', methods=['PUT'])
@login_required
def salvar_conteudo(id_cap):
    """DEPRECATED (pos-Fase 1).

    Salvar conteudo de capitulo via API foi substituido pelo fluxo de
    upload de DOCX completo + merge in-place. O frontend deve enviar
    um arquivo `.docx` para
    `POST /relatorio/versao-trabalho/<id>/capitulo/<id_cap>/upload`
    e confirmar na tela de previa.

    Mantemos a rota retornando 410 Gone para que o frontend antigo
    receba mensagem clara em vez de salvar em coluna que ja foi
    removida do banco.
    """
    _ = id_cap
    return jsonify({
        'erro': (
            'Endpoint depreciado. Use o fluxo de upload completo de '
            'DOCX (rota /relatorio/.../capitulo/<id>/upload).'
        ),
    }), 410


@api_bp.route('/capitulos/<int:id_cap>/conteudo')
@login_required
def obter_conteudo(id_cap):
    """Retorna o DOCX binario do capitulo isolado.

    Pos-Fase 1: extraido em tempo real do DOCX em producao
    (caminho_template do relatorio) via
    `servico_merge_docx.extrair_capitulo_como_docx`. Nao usa mais
    `cap.conteudo_docx` (coluna em vias de ser removida).
    """
    from app.services.servico_sanitizar_docx import sanitizar_docx_bytes
    from app.services.servico_merge_docx import (
        extrair_capitulo_como_docx,
    )

    cap = CapituloDocumento.query.get_or_404(id_cap)
    rel = RelatorioProducao.query.get(cap.id_relatorio)
    if not rel or not rel.caminho_template:
        return ('Relatorio sem DOCX em producao', 404)
    if not os.path.exists(rel.caminho_template):
        return ('Arquivo DOCX em producao nao encontrado', 404)

    try:
        conteudo = extrair_capitulo_como_docx(
            rel.caminho_template, cap
        )
    except (ValueError, OSError, RuntimeError) as e:
        return (f'Erro ao extrair capitulo: {e}', 500)
    if not conteudo:
        return ('Capitulo nao localizado no DOCX em producao', 404)

    bytes_sanitizados = sanitizar_docx_bytes(conteudo)
    saida = bytes_sanitizados or conteudo
    return send_file(
        BytesIO(saida),
        mimetype=(
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.document'
        ),
        as_attachment=False,
        download_name=f'capitulo_{id_cap}.docx',
    )


# ==============================================================
# WORKFLOW — finalizar, aprovar, reprovar
# ==============================================================

@api_bp.route(
    '/relatorios-producao/<int:id_rp>/finalizar',
    methods=['POST']
)
@login_required
def finalizar_relatorio(id_rp):
    """Coordenador finaliza o relatório de produção.

    Pos-Fase 1: delega para `servico_finalizar_relatorio.finalizar`,
    que cria snapshot do `caminho_template` (fonte unica) em
    `storage/relatorios_finalizados/`, calcula checksum, persiste
    `RelatorioFinalizado` e avanca status para 'finalizado'.

    Nao monta mais o DOCX a partir de `conteudo_docx` por capitulo.
    """
    from app.services.servico_finalizar_relatorio import (
        finalizar,
        FinalizacaoError,
    )
    _exigir_perfil('coordenador')
    try:
        rf = finalizar(id_relatorio=id_rp, id_usuario=current_user.id)
    except FinalizacaoError as e:
        return jsonify({'erro': str(e)}), 400
    except (OSError, RuntimeError) as e:
        return jsonify({'erro': f'Erro inesperado: {e}'}), 500
    return jsonify({
        'mensagem': 'Relatorio finalizado com sucesso',
        'id_finalizado': rf.id,
        'versao': rf.versao,
        'checksum': rf.checksum_docx,
        'nome_arquivo': rf.nome_arquivo,
    }), 201


@api_bp.route(
    '/capitulos/<int:id_cap>/finalizar', methods=['POST']
)
@login_required
def finalizar_capitulo(id_cap):
    """Autor marca capítulo como finalizado (pronto para revisão)."""
    _exigir_perfil('autor')
    cap = CapituloDocumento.query.get_or_404(id_cap)
    if cap.status_capitulo not in ('em_edicao', 'reprovado'):
        return jsonify({'erro': 'Status inválido para finalizar'}), 400
    # Só o responsável pode finalizar
    if (cap.id_usuario_responsavel
            and cap.id_usuario_responsavel
            != current_user.id):
        return jsonify({'erro': 'Sem permissão'}), 403
    cap.status_capitulo = 'finalizado'
    # Notificar coordenador(es)
    # NOTE: Implementar query de coordenadores usando dom_perfis_usuario
    # coordenadores = Usuario.query.filter(
    #     Usuario.perfil_id == 2,  # coordenador
    #     Usuario.ativo
    # ).all()
    # for coord in coordenadores:
    #     _notificar(
    #         coord.id,
    #         f'Capítulo "{cap.titulo_capitulo}" finalizado '
    #         f'pelo autor e aguarda revisão.'
    #     )
    db.session.commit()
    return jsonify(_serializar_capitulo(cap))


@api_bp.route(
    '/capitulos/<int:id_cap>/aprovar', methods=['POST']
)
@login_required
def aprovar_capitulo(id_cap):
    """Coordenador aprova o capítulo."""
    _exigir_perfil('coordenador')
    cap = CapituloDocumento.query.get_or_404(id_cap)
    if cap.status_capitulo != 'finalizado':
        return jsonify({'erro': 'Capítulo não está finalizado'}), 400
    cap.status_capitulo = 'aprovado'
    cap.observacao_coordenador = None
    # Notificar autor
    if cap.id_usuario_responsavel:
        _notificar(
            cap.id_usuario_responsavel,
            f'Capítulo "{cap.titulo_capitulo}" aprovado '
            f'pelo coordenador.'
        )
    db.session.commit()
    return jsonify(_serializar_capitulo(cap))


@api_bp.route(
    '/capitulos/<int:id_cap>/reprovar', methods=['POST']
)
@login_required
def reprovar_capitulo(id_cap):
    """Coordenador reprova o capítulo com justificativa."""
    _exigir_perfil('coordenador')
    cap = CapituloDocumento.query.get_or_404(id_cap)
    if cap.status_capitulo != 'finalizado':
        return jsonify({'erro': 'Capítulo não está finalizado'}), 400
    data = request.get_json() or {}
    cap.status_capitulo = 'reprovado'
    cap.observacao_coordenador = data.get('observacao', '')
    # Notificar autor
    if cap.id_usuario_responsavel:
        _notificar(
            cap.id_usuario_responsavel,
            f'Capítulo "{cap.titulo_capitulo}" reprovado. '
            f'Observação: {cap.observacao_coordenador}'
        )
    db.session.commit()
    return jsonify(_serializar_capitulo(cap))


# ==============================================================
# RENDERIZAÇÃO — serve o DOCX em producao direto
# ==============================================================
# Pos-Fase 1: o DOCX em `RelatorioProducao.caminho_template` JA E o
# documento final montado (autores fizeram merge in-place via
# `servico_merge_docx`). Nao ha mais reconstrucao do zero. As rotas
# legadas abaixo apenas servem esse arquivo (com sanitizacao para o
# eigenpal). `_gerar_docx_versao` e `MotorRenderizacao` foram
# removidos.


def _servir_docx_producao(vt, *, as_attachment: bool, nome: str):
    """Helper comum para servir o DOCX em producao com sanitizacao."""
    from app.services.servico_sanitizar_docx import sanitizar_docx

    if not vt.caminho_template or not os.path.exists(vt.caminho_template):
        return jsonify({'erro': 'DOCX em producao indisponivel'}), 404
    mimetype = (
        'application/vnd.openxmlformats-officedocument'
        '.wordprocessingml.document'
    )
    bytes_sanitizados = sanitizar_docx(vt.caminho_template)
    if bytes_sanitizados is None:
        return send_file(
            vt.caminho_template,
            as_attachment=as_attachment,
            download_name=nome,
            mimetype=mimetype,
        )
    return send_file(
        BytesIO(bytes_sanitizados),
        as_attachment=as_attachment,
        download_name=nome,
        mimetype=mimetype,
    )


@api_bp.route(
    '/versoes-trabalho/<int:id_vt>/renderizar', methods=['POST']
)
@login_required
def renderizar_versao(id_vt):
    """Serve o DOCX em producao (fonte unica) como download.

    Rota mantida por compatibilidade do frontend; pos-Fase 1 nao
    monta mais nada — simplesmente devolve `caminho_template`. Para
    finalizar (snapshot + bloqueio), use
    `POST /api/relatorios-producao/<id>/finalizar` ou a rota web
    `/relatorio/producao/<id>/gerar-final`.
    """
    vt = RelatorioProducao.query.get_or_404(id_vt)
    return _servir_docx_producao(
        vt,
        as_attachment=True,
        nome=f'relatorio_{vt.id}_{vt.versao_atual}.docx',
    )


@api_bp.route(
    '/versoes-trabalho/<int:id_vt>/preview-docx'
)
@login_required
def preview_versao(id_vt):
    """Preview inline do DOCX em producao (sem download)."""
    vt = RelatorioProducao.query.get_or_404(id_vt)
    return _servir_docx_producao(
        vt,
        as_attachment=False,
        nome=f'preview_{vt.id}.docx',
    )


# ==============================================================
# GESTÃO DA VERSÃO — vincular biblioteca, listar autores
# ==============================================================

@api_bp.route(
    '/versoes-trabalho/<int:id_vt>/vincular-biblioteca',
    methods=['POST']
)
@login_required
def vincular_biblioteca(id_vt):
    """Coordenador vincula uma biblioteca canônica à versão."""
    _exigir_perfil('coordenador')
    RelatorioProducao.query.get_or_404(id_vt)
    data = request.get_json()
    id_bib = data.get('id_biblioteca')
    if not id_bib:
        return jsonify({'erro': 'id_biblioteca obrigatório'}), 400
    BibliotecaFormatacaoCanonica.query.get_or_404(id_bib)
    # NOTE: Implementar vinculo de biblioteca no novo schema
    db.session.commit()
    return jsonify({'ok': True, 'id_biblioteca': id_bib})


@api_bp.route('/usuarios-autores')
@login_required
def listar_autores():
    """Lista usuários com perfil 'autor' ativos."""
    autores = Usuario.query.filter(
        Usuario.perfil_id == 1,  # NOTE: usar codigo do perfil
        Usuario.ativo
    ).order_by(Usuario.nome).all()
    return jsonify([
        {'id': u.id, 'nome': u.nome}
        for u in autores
    ])


@api_bp.route('/bibliotecas-formatacao')
@login_required
def listar_bibliotecas():
    """Lista bibliotecas de formatação ativas."""
    bibs = BibliotecaFormatacaoCanonica.query.filter_by(
        ativa=True, extraida=True
    ).order_by(
        BibliotecaFormatacaoCanonica.nome_biblioteca
    ).all()
    return jsonify([
        {
            'id': b.id_biblioteca_formatacao_canonica,
            'nome': b.nome_biblioteca,
        }
        for b in bibs
    ])


# ==============================================================
# Helpers
# ==============================================================

def _exigir_perfil(*perfis):
    """Aborta se o usuário não tem um dos perfis exigidos."""
    perfil_ativo = session.get('perfil_ativo', '')
    if perfil_ativo not in perfis:
        abort(403)


def _notificar(id_usuario, mensagem, tipo='workflow'):
    """Cria notificação para o usuário."""
    n = Notificacao(
        id_usuario_destino=id_usuario,
        tipo_notificacao=tipo,
        mensagem=mensagem
    )
    db.session.add(n)


def _serializar_capitulo(cap):
    """Serializa CapituloDocumento para JSON."""
    filhos = CapituloDocumento.query.filter_by(
        id_capitulo_pai=cap.id_capitulo_documento,
        ativo=True
    ).order_by(CapituloDocumento.ordem_capitulo).all()

    return {
        'id': cap.id_capitulo_documento,
        'titulo': cap.titulo_capitulo,
        'ordem': cap.ordem_capitulo,
        'nivel': cap.nivel_capitulo,
        'status': cap.status_capitulo,
        'id_responsavel': cap.id_usuario_responsavel,
        'responsavel_nome': (
            cap.responsavel.nome
            if cap.responsavel else None
        ),
        # Pos-Fase 1: 'tem_conteudo' deixa de espelhar a coluna
        # `conteudo_docx` (em vias de remocao). O conteudo agora vive
        # no DOCX em producao; o capitulo "tem conteudo" se for
        # localizavel naquele DOCX, mas verificar isso no serializer
        # seria caro. Devolvemos True para preservar o comportamento
        # do frontend antigo (que so usa este flag para escolher entre
        # "exibir preview" vs "exibir placeholder"). Pos-migracao
        # completa, este campo pode ser removido.
        'tem_conteudo': True,
        'observacao_coordenador': cap.observacao_coordenador,
        'filhos': [_serializar_capitulo(f) for f in filhos],
    }


@api_bp.route('/logs')
@login_required
def get_logs():
    """Retorna logs do sistema para exibição no navegador."""
    level = request.args.get('level')
    limit = request.args.get('limit', 100, type=int)
    logs = sra_log_handler.get_logs(level=level, limit=limit)
    return jsonify({'logs': logs})
