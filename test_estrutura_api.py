#!/usr/bin/env python
"""Script simples para testar API de estrutura de envio."""

import json
import os

from docx import Document

from app import create_app
from app.models.envio_conteudo import EnvioConteudo
from app.services.servico_extracao_canonica import ServicoExtracaoCanonica
from app.services.servico_envio_autor import ServicoEnvioAutor


def test_estrutura_envio(id_envio):
    """Testa extração de estrutura de um envio existente."""
    print(f'=== Testando estrutura do envio {id_envio} ===\n')

    app = create_app()

    with app.app_context():
        # Buscar envio
        envio = EnvioConteudo.query.get(id_envio)
        if not envio:
            print(f'✗ Envio {id_envio} não encontrado')
            return

        print(f'✓ Envio encontrado: {envio.nome_arquivo}')
        print(f'  Caminho: {envio.caminho_arquivo}')
        print(f'  Status: {envio.status_envio}')

        # Verificar se já tem sugestoes_json
        if envio.sugestoes_json:
            print('\n✓ Estrutura já processada (sugestoes_json existe)')
            estrutura = json.loads(envio.sugestoes_json)
            print(f'  Capítulos: {len(estrutura.get("capitulos", []))}')
            print(f'  Legendas: {list(estrutura.get("legendas", {}).keys())}')

            # Se capítulos vazios, tentar extrair novamente para debug
            if len(estrutura.get("capitulos", [])) == 0:
                print('\n⚠ Capítulos vazios, tentando extração manual...')
                if os.path.exists(envio.caminho_arquivo):
                    try:
                        doc = Document(envio.caminho_arquivo)
                        print('  DOCX carregado: {} parágrafos'.format(
                            len(doc.paragraphs)
                        ))

                        # Mostrar primeiros parágrafos
                        print('\n  Primeiros 10 parágrafos:')
                        for i, p in enumerate(doc.paragraphs[:10]):
                            estilo_obj = p.style
                            estilo = estilo_obj.name if estilo_obj else ''
                            texto = p.text.strip()[:50]
                            print('    {}. [{}] {}'.format(
                                i + 1, estilo, texto
                            ))

                        # Tentar extração por padrão
                        capitulos = (
                            ServicoEnvioAutor
                            ._extrair_capitulos_por_padrao(  # noqa: SLF001,W0212
                                doc
                            )
                        )
                        print(
                            '\n  Capítulos extraídos por padrão: {}'.format(
                                len(capitulos)
                            )
                        )

                        if capitulos:
                            print('\n  Capítulos encontrados:')
                            for i, cap in enumerate(capitulos[:5]):
                                indice = cap.get("indice", "")
                                titulo = cap.get("titulo", "")
                                nivel = cap.get("nivel", "")
                                print('    {}. {} {} (nível {})'.format(
                                    i + 1, indice, titulo, nivel
                                ))

                    except Exception as e:  # noqa: W0719,W0703,B904 - debug script
                        print(f'  ✗ Erro na extração: {e}')
        else:
            print('\n⚠ Estrutura não processada ainda')

            # Tentar extrair estrutura agora
            if os.path.exists(envio.caminho_arquivo):
                print('\nExtraindo estrutura do DOCX...')
                try:
                    doc = Document(
                        envio.caminho_arquivo
                    )
                    capitulos = (
                        ServicoExtracaoCanonica
                        ._extrair_capitulos(  # noqa: SLF001,W0212
                            doc
                        )
                    )
                    legendas = (
                        ServicoExtracaoCanonica
                        ._extrair_legendas(  # noqa: SLF001,W0212
                            doc
                        )
                    )

                    print('Extração realizada')
                    print('  Capítulos encontrados: {}'.format(len(capitulos)))
                    print(
                        '  Legendas encontradas: {}'.format(
                            list(legendas.keys())
                        )
                    )

                    # Mostrar alguns capítulos
                    if capitulos:
                        print('\n  Primeiros capítulos:')
                        for i, cap in enumerate(capitulos[:3]):
                            indice = cap.get("indice", "")
                            titulo = cap.get("titulo", "")
                            nivel = cap.get("nivel", "")
                            print('    {}. {} {} (nível {})'.format(
                                i + 1, indice, titulo, nivel
                            ))

                except Exception as e:  # noqa: W0719,W0703,B904 - debug script
                    print(f'✗ Erro na extração: {e}')
            else:
                print(
                    'Arquivo DOCX não encontrado: {}'.format(
                        envio.caminho_arquivo
                    )
                )


if __name__ == '__main__':
    # Testar com o envio mais recente (ID 7)
    test_estrutura_envio(7)
